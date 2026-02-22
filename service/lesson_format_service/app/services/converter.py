"""
格式转换服务 - HTML到各种格式的转换
"""
import html2text
import json
import re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
from app.models import LessonPlan
from typing import Dict, Tuple
import pdfplumber

async def convert_lesson(
    pdf_bytes: bytes, 
    lesson_content: str,
    force_qwen: bool = False,
    precision_mode: str = "standard"
) -> Tuple[str, str]:
    """
    混合转换策略
    
    Args:
        pdf_bytes: PDF模版文件
        lesson_content: 教案内容（JSON或文本）
        force_qwen: 是否强制使用Qwen
        precision_mode: 精确度模式 "standard" | "precise"
    
    Returns: (html, method_used)
    method_used: "pdf_parse" | "pdf_parse_precise" | "qwen" | "qwen_precise" | "fallback"
    """
    
    # 提取精确样式（如果是精确模式）
    precise_styles = None
    if precision_mode == "precise":
        try:
            from app.services.pdf_parser import extract_precise_styles
            print("提取PDF精确样式数据...")
            precise_styles = extract_precise_styles(pdf_bytes)
            if precise_styles and precise_styles.get("pages"):
                print(f"成功提取精确样式数据：{len(precise_styles['pages'])}页")
        except Exception as e:
            print(f"精确样式提取失败: {e}")
            import traceback
            traceback.print_exc()
    
    if not force_qwen:
        # 尝试方案1: PDF解析填充
        try:
            print(f"尝试PDF解析填充（{precision_mode}模式）...")
            html = await parse_and_fill_pdf(pdf_bytes, lesson_content, precision_mode)
            method = "pdf_parse_precise" if precision_mode == "precise" else "pdf_parse"
            return html, method
        except Exception as e:
            print(f"PDF解析失败，降级到Qwen: {e}")
            import traceback
            traceback.print_exc()
    
    # 方案2: Qwen AI（传入精确样式数据）
    try:
        print("使用Qwen AI...")
        from app.services.qwen_service import convert_with_qwen_pdf
        html = await convert_with_qwen_pdf(pdf_bytes, lesson_content, precise_styles)
        method = "qwen_precise" if precise_styles else "qwen"
        return html, method
    except Exception as e:
        print(f"Qwen调用失败，使用基础模版: {e}")
    
    # 方案3: 基础模版
    print("使用基础模版...")
    html = generate_basic_template(lesson_content)
    return html, "fallback"


async def parse_and_fill_pdf(
    pdf_bytes: bytes, 
    content: str,
    precision_mode: str = "standard"
) -> str:
    """
    方案1: 解析PDF并填充内容
    
    Args:
        pdf_bytes: PDF文件字节
        content: 教案内容（JSON或文本）
        precision_mode: "standard" 标准模式 | "precise" 精确模式
    """
    import json
    
    if precision_mode == "precise":
        # 精确模式：使用像素级精确提取和映射
        from app.services.pdf_parser import extract_precise_styles
        from app.services.precise_mapper import map_content_to_template
        
        print("使用精确模式进行PDF解析和内容映射...")
        
        # 1. 精确提取PDF样式
        styles = extract_precise_styles(pdf_bytes)
        
        # 2. 解析JSON内容
        try:
            json_content = json.loads(content)
        except:
            # 如果不是JSON，转换为简单结构
            json_content = {"mainContent": content}
        
        # 3. 精确映射并生成HTML
        html = map_content_to_template(styles, json_content, auto_paginate=True)
        
    else:
        # 标准模式：使用现有的方式
        from app.services.pdf_parser import (
            extract_pdf_styles, 
            parse_lesson_content, 
            build_html_with_style
        )
        
        print("使用标准模式进行PDF解析...")
        
        # 1. 提取PDF样式
        style_info = extract_pdf_styles(pdf_bytes)
        
        # 2. 解析教案内容章节
        sections = parse_lesson_content(content)
        
        # 3. 生成HTML（使用提取的样式）
        html = build_html_with_style(style_info, sections)
    
    return html


def generate_basic_template(content: str) -> str:
    """
    生成基础HTML模版（降级方案）
    """
    import json
    
    title = "教案"
    body_content = content
    
    # 尝试解析JSON
    try:
        data = json.loads(content)
        if isinstance(data, dict):
            title = data.get("courseTitle", "教案")
            body_content = f"""
            <h2>教學目標</h2>
            <p>{data.get("teachingObjectives", "")}</p>
            
            <h2>教學內容</h2>
            <p>{data.get("mainContent", "")}</p>
            
            <h2>教學方法</h2>
            <p>{data.get("teachingMethods", "")}</p>
            """
    except:
        # 纯文本，换行转<br>
        body_content = content.replace('\n', '<br>\n')
    
    return f"""<!DOCTYPE html>
<html lang="zh-TW">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{title}</title>
    <style>
        body {{
            font-family: 'Microsoft JhengHei', '微軟正黑體', 'Arial', sans-serif;
            font-size: 14px;
            line-height: 1.8;
            color: #333;
            padding: 40px;
            max-width: 800px;
            margin: 0 auto;
        }}
        h1 {{
            text-align: center;
            color: #1a1a1a;
            margin-bottom: 30px;
        }}
        h2 {{
            color: #2c3e50;
            margin-top: 25px;
            margin-bottom: 15px;
            border-left: 4px solid #3498db;
            padding-left: 12px;
        }}
        p {{
            margin-bottom: 12px;
            text-indent: 2em;
        }}
    </style>
</head>
<body>
    <h1>{title}</h1>
    {body_content}
</body>
</html>"""


async def convert_html_to_formats(html: str, lesson_plan: LessonPlan) -> Dict[str, any]:
    """
    将HTML转换为多种格式
    
    Returns: {
        "json": lesson_plan dict,
        "docx": bytes,
        "md": str,
        "txt": str,
        "pdf": bytes
    }
    """
    return {
        "json": lesson_plan.model_dump(),
        "docx": convert_to_docx(html, lesson_plan),
        "md": convert_to_markdown(html),
        "txt": convert_to_txt(html),
        "pdf": convert_to_pdf(html)
    }


def convert_to_markdown(html: str) -> str:
    """HTML转Markdown"""
    try:
        h = html2text.HTML2Text()
        h.ignore_links = False
        h.ignore_images = False
        h.ignore_emphasis = False
        h.body_width = 0  # 不自动换行
        markdown = h.handle(html)
        return markdown
    except Exception as e:
        print(f"转换Markdown失败: {e}")
        return f"# 转换失败\n\n{str(e)}"


def convert_to_txt(html: str) -> str:
    """HTML转纯文本"""
    try:
        # 去除HTML标签
        text = re.sub(r'<script[^>]*>.*?</script>', '', html, flags=re.DOTALL | re.IGNORECASE)
        text = re.sub(r'<style[^>]*>.*?</style>', '', text, flags=re.DOTALL | re.IGNORECASE)
        text = re.sub(r'<[^>]+>', '', text)
        # 清理多余空白
        text = re.sub(r'\n\s*\n', '\n\n', text)
        text = text.strip()
        return text
    except Exception as e:
        print(f"转换TXT失败: {e}")
        return f"转换失败: {str(e)}"


def convert_to_pdf(html: str) -> bytes:
    """HTML转PDF（使用WeasyPrint或降级）"""
    try:
        # 尝试使用WeasyPrint
        from weasyprint import HTML as WeasyHTML
        pdf_buffer = BytesIO()
        WeasyHTML(string=html).write_pdf(pdf_buffer)
        pdf_buffer.seek(0)
        return pdf_buffer.read()
    except ImportError:
        print("WeasyPrint未安装，PDF功能不可用")
        return b"PDF conversion not available - WeasyPrint not installed"
    except Exception as e:
        print(f"转换PDF失败: {e}")
        # 返回错误信息
        error_msg = f"PDF转换失败: {str(e)}\n请确保WeasyPrint已正确安装"
        return error_msg.encode('utf-8')


def convert_to_docx(html: str, lesson_plan: LessonPlan) -> bytes:
    """HTML转DOCX（使用python-docx）"""
    try:
        doc = Document()
        
        # 设置文档样式
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Microsoft JhengHei'
        font.size = Pt(12)
        
        # 添加标题
        title = doc.add_heading(lesson_plan.metadata.courseTitle, level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 简单的HTML解析（提取文本内容）
        # 移除script和style标签
        clean_html = re.sub(r'<script[^>]*>.*?</script>', '', html, flags=re.DOTALL | re.IGNORECASE)
        clean_html = re.sub(r'<style[^>]*>.*?</style>', '', clean_html, flags=re.DOTALL | re.IGNORECASE)
        
        # 提取文本，保留基本结构
        text = re.sub(r'<[^>]+>', '\n', clean_html)
        text = re.sub(r'\n\s*\n', '\n\n', text)
        text = text.strip()
        
        # 添加内容
        for line in text.split('\n'):
            if line.strip():
                doc.add_paragraph(line.strip())
        
        # 保存到BytesIO
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
        
    except Exception as e:
        print(f"转换DOCX失败: {e}")
        # 返回基础文档
        doc = Document()
        doc.add_heading('转换失败', 0)
        doc.add_paragraph(f'错误: {str(e)}')
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
