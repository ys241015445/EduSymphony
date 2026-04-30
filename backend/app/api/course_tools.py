"""Course content tools: outline, PPT, exercises, practice — all powered by Doubao Chat API."""
import asyncio
import uuid
import json
import os
import traceback
from typing import Optional
from io import BytesIO

from fastapi import APIRouter, Depends, HTTPException, Form, Query
from fastapi.responses import Response, StreamingResponse
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select
from urllib.parse import quote
from loguru import logger

from app.core.database import get_db, async_session_maker
from app.core.config import settings
from app.core.deps import get_current_active_user, require_not_limited, resolve_documents_owner
from app.models.user import User
from app.models.lesson import LessonPlan, LessonSeries
from app.models.course_tool import CourseToolResult

router = APIRouter(prefix="/course-tools", tags=["课程工具"], dependencies=[Depends(require_not_limited)])

DOUBAO_PROVIDER = "doubao"
ALLOWED_PROVIDERS = {"doubao", "qwen", "deepseek", "kimi", "spark", "openai"}


def _resolve_provider(provider: Optional[str], default: str = DOUBAO_PROVIDER) -> str:
    """Accept a case-insensitive provider hint from the client; fall back to default."""
    if not provider:
        return default
    p = provider.strip().lower()
    return p if p in ALLOWED_PROVIDERS else default


def _cd(title: str, ext: str) -> str:
    safe = "".join(c for c in (title or "") if ord(c) < 128 and (c.isalnum() or c in " _-")).strip() or "file"
    try:
        utf8 = quote(f"{title}.{ext}")
        return f'attachment; filename="{safe}.{ext}"; filename*=UTF-8\'\'{utf8}'
    except Exception:
        return f'attachment; filename="{safe}.{ext}"'


def _get_ai():
    from app.services.ai_service import AIService
    return AIService()


MAX_SOURCE_CHARS = 8000


def _lesson_content_by_mode(lesson: LessonPlan, mode: str = "auto") -> str:
    """Extract textual content from a LessonPlan by requested mode.

    mode:
      - "optimized": final_content.full_optimized
      - "draft":     final_content.full_draft
      - "original":  lesson.source_content
      - "auto"(default): optimized → draft → original
    """
    fc = lesson.final_content or {}
    if isinstance(fc, str):
        try:
            fc = json.loads(fc)
        except Exception:
            fc = {}
    optimized = (fc.get("full_optimized") or "") if isinstance(fc, dict) else ""
    draft = (fc.get("full_draft") or "") if isinstance(fc, dict) else ""
    original = lesson.source_content or ""

    if mode == "optimized":
        picked = optimized or draft or original
    elif mode == "draft":
        picked = draft or optimized or original
    elif mode == "original":
        picked = original or optimized or draft
    else:
        picked = optimized or draft or original
    return (picked or "")[:MAX_SOURCE_CHARS]


def _syllabus_to_text(series: LessonSeries) -> str:
    """Flatten LessonSeries into a readable plaintext block for AI prompts."""
    parts: list[str] = []
    header = f"《{series.title}》 学科：{series.subject} 年级：{series.grade_level}"
    parts.append(header)
    if series.objectives:
        parts.append(f"培养目标：{series.objectives}")
    if series.outline_text:
        parts.append(f"教学大纲：\n{series.outline_text}")
    if series.book_content:
        parts.append(f"教材内容：\n{series.book_content}")
    if series.schedule_text:
        parts.append(f"教学进度：\n{series.schedule_text}")

    syl = series.syllabus
    if isinstance(syl, str):
        try:
            syl = json.loads(syl)
        except Exception:
            syl = None
    if isinstance(syl, dict):
        lessons = syl.get("lessons") or syl.get("weeks") or []
        if isinstance(lessons, list) and lessons:
            parts.append("—— 课程安排 ——")
            for idx, item in enumerate(lessons, 1):
                if not isinstance(item, dict):
                    continue
                title = item.get("title") or item.get("name") or f"第{idx}节"
                dur = item.get("duration") or item.get("week")
                line = f"{idx}. {title}" + (f"（{dur}）" if dur else "")
                parts.append(line)
                for obj in (item.get("objectives") or [])[:4]:
                    parts.append(f"   - 目标：{obj}")
                for kp in (item.get("key_points") or [])[:4]:
                    parts.append(f"   - 要点：{kp}")
                desc = item.get("content") or item.get("description")
                if desc:
                    parts.append(f"   {str(desc)[:200]}")
        elif syl.get("title") or syl.get("sections"):
            parts.append(json.dumps(syl, ensure_ascii=False)[:3000])
    text_out = "\n".join(parts).strip()
    return text_out[:MAX_SOURCE_CHARS]


async def _fetch_tool_result(
    db: AsyncSession, subject_user: User, result_id: str, expected_type: Optional[str] = None,
) -> CourseToolResult:
    """Fetch a CourseToolResult row with ownership + optional type check."""
    res = await db.execute(select(CourseToolResult).where(CourseToolResult.id == result_id))
    ci = res.scalar_one_or_none()
    if not ci:
        raise HTTPException(404, "所选记录不存在")
    if ci.user_id != subject_user.id:
        raise HTTPException(403, "无权访问所选记录")
    if expected_type and ci.tool_type != expected_type:
        raise HTTPException(400, f"记录类型不符，期望 {expected_type}，实际 {ci.tool_type}")
    return ci


def _tool_result_to_context(ci: CourseToolResult) -> dict:
    """Serialize a CourseToolResult into {text,subject,grade_level,topic,title}."""
    data = ci.result or {}
    if not isinstance(data, dict):
        data = {"content": str(data)}
    params = ci.params or {}
    title = (data.get("title") or params.get("topic") or ci.tool_type or "").strip()
    text = json.dumps(data, ensure_ascii=False)
    return {
        "text": text[:MAX_SOURCE_CHARS],
        "subject": params.get("subject") or "",
        "grade_level": params.get("grade_level") or "",
        "region": params.get("region") or "mainland",
        "topic": title,
        "title": title,
        "source_type": ci.tool_type,
        "source_id": ci.id,
    }


async def _resolve_source_context(
    db: AsyncSession,
    current_user: User,
    *,
    lesson_id: Optional[str] = None,
    series_id: Optional[str] = None,
    outline_id: Optional[str] = None,
    ppt_id: Optional[str] = None,
    exercises_id: Optional[str] = None,
    practice_id: Optional[str] = None,
    source_mode: str = "auto",
) -> dict:
    """Unified source resolver for course tools.

    Priority: explicit tool result (outline → ppt → exercises → practice) → series → lesson.

    Returns: {
      text, subject, grade_level, region, topic,
      source_type, source_id, title
    }
    Raises HTTPException(403/404) on ownership mismatch / not found.
    """
    out = {
        "text": "",
        "subject": "",
        "grade_level": "",
        "region": "",
        "topic": "",
        "source_type": None,
        "source_id": None,
        "title": "",
    }
    mode = (source_mode or "auto").strip().lower()
    if mode not in {"auto", "optimized", "draft", "original"}:
        mode = "auto"

    # 1) Tool-result sources (outline / ppt / exercises / practice)
    tool_refs = [
        ("outline", outline_id),
        ("ppt", ppt_id),
        ("exercises", exercises_id),
        ("practice", practice_id),
    ]
    for expected_type, tid in tool_refs:
        if tid:
            ci = await _fetch_tool_result(db, current_user, tid, expected_type=expected_type)
            return _tool_result_to_context(ci)

    # 2) Series source
    if series_id:
        res = await db.execute(select(LessonSeries).where(LessonSeries.id == series_id))
        series = res.scalar_one_or_none()
        if not series:
            raise HTTPException(404, "所选系列不存在")
        if series.user_id != current_user.id:
            raise HTTPException(403, "无权访问所选系列")
        out["text"] = _syllabus_to_text(series)
        out["subject"] = series.subject or ""
        out["grade_level"] = series.grade_level or ""
        out["region"] = series.region or "mainland"
        out["topic"] = series.title or ""
        out["title"] = series.title or ""
        out["source_type"] = "series"
        out["source_id"] = series_id
        return out

    # 3) Lesson source
    if lesson_id:
        res = await db.execute(select(LessonPlan).where(LessonPlan.id == lesson_id))
        lesson = res.scalar_one_or_none()
        if not lesson:
            raise HTTPException(404, "所选教案不存在")
        if lesson.user_id != current_user.id:
            raise HTTPException(403, "无权访问所选教案")
        out["text"] = _lesson_content_by_mode(lesson, mode)
        out["subject"] = lesson.subject or ""
        out["grade_level"] = lesson.grade_level or ""
        out["region"] = lesson.region or "mainland"
        out["topic"] = lesson.topic or ""
        out["title"] = lesson.title or ""
        out["source_type"] = "lesson"
        out["source_id"] = lesson_id
        return out

    return out


def _lesson_context(lesson: LessonPlan) -> str:  # kept for backward-compat
    return _lesson_content_by_mode(lesson, "auto")


async def _save(db: AsyncSession, user_id: str, lesson_id: Optional[str],
                tool_type: str, params: dict, result: dict, file_path: str = None,
                source_meta: Optional[dict] = None) -> CourseToolResult:
    if source_meta and isinstance(result, dict):
        result = {**result, "_source": {
            "type": source_meta.get("source_type"),
            "id": source_meta.get("source_id"),
            "title": source_meta.get("title"),
            "mode": source_meta.get("mode"),
        }}
    item = CourseToolResult(
        id=str(uuid.uuid4()), user_id=user_id, lesson_id=lesson_id,
        tool_type=tool_type, params=params, result=result, file_path=file_path,
    )
    db.add(item)
    await db.commit()
    await db.refresh(item)
    return item


async def _get_result(result_id: str, user_id: str, db: AsyncSession) -> CourseToolResult:
    r = await db.execute(
        select(CourseToolResult).where(CourseToolResult.id == result_id, CourseToolResult.user_id == user_id)
    )
    item = r.scalar_one_or_none()
    if not item:
        raise HTTPException(404, "记录不存在")
    return item


async def _create_pending(
    db: AsyncSession,
    *,
    user_id: str,
    lesson_id: Optional[str],
    tool_type: str,
    params: dict,
    initial_result: Optional[dict] = None,
) -> str:
    """Insert a `pending` CourseToolResult row and return its id."""
    rid = str(uuid.uuid4())
    item = CourseToolResult(
        id=rid, user_id=user_id, lesson_id=lesson_id,
        tool_type=tool_type, params=params or {},
        result=initial_result or {}, file_path=None,
        status="pending", error_message=None,
    )
    db.add(item)
    await db.commit()
    return rid


async def _enqueue_tool(
    *,
    db: AsyncSession,
    user_id: str,
    kind: str,
    result_id: str,
    tool_type: str,
    title: str,
) -> dict:
    """Enqueue the async job; return the canonical response shape for 4 POST endpoints.

    Response shape:
      { id: result_id, result_id, tool_type, status, title, queued: bool }
    Note: `id` is kept for backward-compat with the old sync-shape callers.
    """
    from app.tasks.queue_manager import enqueue
    try:
        queued = await enqueue(
            target_id=result_id, user_id=user_id, kind=kind, max_attempts=1,
        )
    except Exception as e:
        # queue insertion failed → flip result row to failed so the user sees it
        r = await db.execute(select(CourseToolResult).where(CourseToolResult.id == result_id))
        ci = r.scalar_one_or_none()
        if ci is not None:
            ci.status = "failed"
            ci.error_message = f"enqueue failed: {e}"[:2000]
            await db.commit()
        raise HTTPException(500, f"任务入队失败: {e}")

    return {
        "id": result_id,
        "result_id": result_id,
        "tool_type": tool_type,
        "status": "queued",
        "queued": queued,
        "title": title,
    }


# ── helpers ──────────────────────────────────────────────────────────

OUTLINE_SYSTEM = """你是课程大纲设计专家。根据提供的学科、年级、地区和教学内容，生成结构化的课程大纲。
输出严格JSON格式（不要markdown代码块），结构：
{
  "title": "大纲标题",
  "scope": "semester 或 single_lesson",
  "sections": [
    {
      "title": "章节/环节标题",
      "duration": "建议时长",
      "objectives": ["目标1","目标2"],
      "key_points": ["要点1","要点2"],
      "activities": ["活动1","活动2"],
      "sub_sections": [...]
    }
  ]
}"""

PPT_STYLE_ANALYZER_SYSTEM = """你是资深视觉设计总监，专长教育 PPT 的配色分析。

任务：根据提供的课程主题、学段、学科、地区、用户偏好标签和补充说明，给出 3 个配色候选方案，供教师三选一。

硬性要求：
1. 输出严格的 JSON（不要 markdown 代码块），结构必须是：
{
  "candidates": [
    {
      "name": "方案名（4-6 个中文字符，概括情绪）",
      "mood": "情绪关键词描述（10-20 字）",
      "palette": {
        "bg":            "#RRGGBB",
        "title_color":   "#RRGGBB",
        "body_color":    "#RRGGBB",
        "accent":        "#RRGGBB",
        "section_bg":    "#RRGGBB",
        "bullet_color":  "#RRGGBB"
      },
      "rationale": "为什么这个配色适合该主题/学段（40-80 字）"
    }
  ]
}

2. 必须返回 **恰好 3 个** candidates。
3. 三个方案的情绪必须明显不同（比如：暖色活泼 vs 冷色学术 vs 中性极简），而不是同一色系的三个变体。
4. 所有颜色必须是 6 位大写 HEX，形如 "#FFF8E1"，绝对不能用颜色名或 rgb()。
5. 确保 title_color 和 bg 对比度足够，body_color 也要在 bg 上清晰可读。
6. 针对低龄学生（小学）倾向明亮暖色；中高年级倾向稳重色；大学/专业课倾向学术深色。
7. 若用户指定了偏好标签或文字描述，至少其中 1-2 个方案要契合该偏好，第 3 个方案可以给出差异化备选。"""


PPT_STYLE_ANALYZER_STREAM_SYSTEM = """你是资深视觉设计总监，专长教育 PPT 的整体风格设计（不仅是配色，还包含版式、字体、封面变体）。

输出格式（**严格遵守，顺序不能颠倒**）：

## 分析
用 3-5 句自然语言，结合课程主题、学段、地区、用户偏好标签/描述，分析合适的情绪基调、版式方向、字体气质、封面思路。
可以使用简单 Markdown（段落、加粗），但不要出现项目符号列表，不要出现代码块。

## 推荐模板
紧接着输出 **1 个** fenced JSON 代码块（```json ... ```），结构如下；必须返回恰好 3 个候选模板，三者在视觉语言上必须明显不同（至少 layout_style 或主色色相不同），而不是同一色系的三个小变体：

```json
{
  "candidates": [
    {
      "name":          "模板名（4-8 个中文字符）",
      "mood":          "情绪关键词（10-20 字）",
      "palette": {
        "bg":            "#RRGGBB",
        "title_color":   "#RRGGBB",
        "body_color":    "#RRGGBB",
        "accent":        "#RRGGBB",
        "section_bg":    "#RRGGBB",
        "bullet_color":  "#RRGGBB"
      },
      "layout_style":  "academic / modern / kawaii / tech / editorial / minimal / business / natural 任选其一",
      "typography":    "serif / sans_display / handwriting / mono 任选其一",
      "cover_style":   "centered / split / decorative 任选其一",
      "use_case":      "适用场景短句（10-20 字）",
      "rationale":     "为什么这个组合适合该主题/学段（40-80 字）"
    }
  ]
}
```

硬性要求：
1. 必须两段都输出：先「## 分析」自然语言，再「## 推荐模板」1 个 fenced JSON 代码块。
2. `candidates` 数组长度必须恰好为 3。
3. 三个候选的 layout_style / palette hue / cover_style 至少要有 **2 个维度不同**；不允许三个都是同一 layout_style。
4. 所有颜色必须是 6 位大写 HEX（形如 "#FFF8E1"），不能用颜色名或 rgb()；title_color 和 bg 对比度要高。
5. 小学段倾向 kawaii / natural + 明亮暖色；中高年级倾向 modern / editorial + 稳重色；大学/专业课倾向 academic / business + 深色或中性色；计算机/工程类可选 tech。
6. 若用户指定偏好标签或描述，至少 1-2 个候选须契合该偏好，第 3 个可做差异化备选。
7. JSON 代码块外不要再出现任何 ``` 符号或多余的 JSON。"""

PPT_SYSTEM = """你是专业 PPT 设计师。根据提供的课程内容 + 用户选中的【模板元数据】生成 PPT 的结构化 JSON 数据。
输出严格 JSON 格式（**不要** markdown 代码块，只输出 JSON 本体），结构：
{
  "title": "演示文稿标题",
  "subtitle": "副标题（可空）",
  "slides": [
    {
      "layout": "title_slide | section_header | agenda | content | two_column | comparison | timeline | process_steps | quote | callout | stats | closing",
      "title": "页面标题",
      "subtitle": "副标题（title_slide / section_header / quote 用）",
      "bullets":       ["要点1","要点2","要点3"],                 // content / agenda / closing
      "left_bullets":  ["左列要点1","左列要点2"],                  // two_column / comparison
      "right_bullets": ["右列要点1","右列要点2"],                  // two_column / comparison
      "left_title":    "左列小标题",                              // comparison
      "right_title":   "右列小标题",                              // comparison
      "steps":         [{"name":"阶段1","desc":"描述"}, ...],     // timeline / process_steps
      "stats":         [{"value":"95%","label":"正确率"}, ...],   // stats
      "quote":         "引用原文",                                // quote
      "quote_author":  "作者/出处",                               // quote
      "callout":       "一句话重点强调",                           // callout
      "image_prompt":  "如果这页适合配图，用 20-40 字描述画面",    // 可选
      "notes":         "演讲者备注（说课提示）"
    }
  ]
}

硬性要求：
1. 第一页必须是 `title_slide`（封面），最后一页必须是 `closing`。
2. 第二页建议是 `agenda`（目录）；每个主要章节前插 `section_header`；总页数 15-25 页。
3. **必须至少出现 5 种不同的 layout**（混合 content / two_column / comparison / timeline / process_steps / quote / callout / stats 等）以保证视觉节奏。
4. content 页每页 3-6 条 bullets；two_column / comparison 左右各 2-5 条；timeline / process_steps 的 steps 3-6 个；stats 3-5 个。
5. **所有 notes 字段都必须写**（说课提示，30-80 字），方便老师讲解。
6. 不要输出 markdown 代码块、不要输出解释说明、只输出 JSON 本体。"""

PPT_OUTLINE_SYSTEM = """你是资深 PPT 教学设计师，专门为中国中小学/高校老师设计课堂 PPT。
当前任务：**只输出整套 PPT 的「全局结构大纲」**，不要展开任何正文 —— 正文会由下一步逐页深度生成。
深度思考要求：先在脑内推演一遍这堂课的"教学逻辑链"——为什么这样开场、章节怎么递进、每页承担什么教学功能、
学生注意力曲线怎么走 —— 想清楚后再产出大纲。

输出严格 JSON（不要 markdown 代码块），结构：
{
  "title": "演示文稿主标题",
  "subtitle": "副标题(可空)",
  "theme_keywords": ["关键词1","关键词2","关键词3"],
  "total_pages": 18,
  "pages": [
    {
      "index": 1,
      "section": "导入" | "概念" | "案例" | "练习" | "总结" 等章节归属,
      "layout": "title_slide | section_header | agenda | content | two_column | comparison | timeline | process_steps | quote | callout | stats | closing",
      "page_title": "本页标题(8-20字)",
      "key_focus": "本页要讲清楚的核心点(20-40字, 写明这页要让学生带走什么)",
      "prev_link": "与上一页的衔接(0-20字, 第1页留空)",
      "next_link": "为下一页埋的伏笔(0-20字, 末页留空)"
    }
  ]
}

硬性约束：
1. total_pages ∈ [15, 25]，pages.length 必须等于 total_pages。
2. 第 1 页 layout=title_slide；第 2 页建议 agenda；每个新章节前插一页 section_header；最后一页 closing。
3. **整组 pages 至少出现 6 种不同 layout**，刻意避免连续 3 页都是 content。
4. 章节顺序遵循"导入→概念铺垫→核心讲解→案例/演练→巩固/总结"教学闭环。
5. 不要输出说明文字、不要输出 markdown、只输出 JSON。"""

PPT_PAGE_SYSTEM = """你是资深 PPT 教学设计师，正在为一节课的**单独一页**做深度内容创作。
深度思考要求：基于给定的全局上下文（学科/年级/主题/模板风格 + 本页在大纲里的位置 + 上下页衔接），
**先在脑内想清楚 4 件事再下笔**：
  (a) 这一页学生该带走什么知识点 / 能力？
  (b) 用什么例子 / 场景 / 数据 / 类比讲最有效？
  (c) 老师站在讲台上，这页要怎么"讲"出来 —— 开场一句话、关键术语怎么过渡、板书重点写什么？
  (d) 学生易错点 / 易混点 / 常见误解是什么？要不要现场提问？

然后按指定的 layout 把字段填到位，输出严格 JSON（不要 markdown 代码块），单页 schema：
{
  "layout": "(沿用入参指定的 layout)",
  "title": "页标题",
  "subtitle": "副标题(title_slide / section_header / quote 才用)",
  "bullets":       ["要点1(25-50字, 含具体例子或定义)", "..."],   // content / agenda / closing
  "left_bullets":  ["..."],                                     // two_column / comparison
  "right_bullets": ["..."],                                     // two_column / comparison
  "left_title":    "左列小标题",                                 // comparison
  "right_title":   "右列小标题",                                 // comparison
  "steps":         [{"name":"阶段名","desc":"30-60字描述,要有动作或例子"}],  // timeline / process_steps
  "stats":         [{"value":"95%","label":"指标说明"}],         // stats
  "quote":         "引用原文",                                   // quote
  "quote_author":  "作者/出处",                                  // quote
  "callout":       "一句话重点(20-40字)",                        // callout
  "image_prompt":  "若适合配图,用 20-40 字描述画面",             // 可选
  "notes":         "教师讲话说口 80-200 字, 必须包含: 开场过渡句 + 关键概念怎么讲 + 板书重点 + 易错点提醒 / 互动设计"
}

硬性约束：
1. **必须**沿用入参指定的 layout，不要擅自改。
2. **notes 字段长度 80-200 字**，必须有"教师怎么开口讲这页"的实操指引，不能空、不能敷衍。
3. content 类页面 bullets 给 3-5 条，每条 25-50 字，**禁止**只写名词短语。
4. timeline / process_steps 的 steps 给 3-6 个，每个 desc 30-60 字、要有动作。
5. 不要输出说明文字、不要输出 markdown、只输出 JSON。"""

EXERCISE_SYSTEM = """你是教育评估专家。根据提供的教学内容生成习题。
输出严格JSON格式（不要markdown代码块），结构：
{
  "title": "习题标题",
  "exercise_type": "daily_homework / quiz / exam",
  "total_score": 100,
  "exercises": [
    {
      "id": 1,
      "type": "choice / fill_blank / short_answer / essay / true_false",
      "question": "题目内容",
      "options": ["A. ...", "B. ...", "C. ...", "D. ..."],
      "answer": "正确答案",
      "explanation": "解析",
      "score": 5,
      "difficulty": "easy / medium / hard"
    }
  ]
}
要求：题型多样，难度梯度合理，答案解析详细。"""

PRACTICE_SYSTEM = """你是课堂教学设计专家。根据教学内容生成课上练习和实操方案。
输出严格JSON格式（不要markdown代码块），结构：
{
  "title": "课上练习标题",
  "theory_summary": "理论知识要点总结（200-400字）",
  "practices": [
    {
      "id": 1,
      "type": "individual / group / hands_on / discussion",
      "title": "练习标题",
      "description": "练习描述和步骤",
      "duration": "建议时长",
      "materials": "所需材料",
      "expected_outcome": "预期成果"
    }
  ],
  "assessment_criteria": "评价标准"
}"""


def _parse_json_response(text: str) -> dict:
    t = text.strip()
    if t.startswith("```"):
        first_nl = t.index("\n") if "\n" in t else 3
        t = t[first_nl + 1:]
    if t.endswith("```"):
        t = t[:-3]
    t = t.strip()
    return json.loads(t)


# ── PPT 两阶段深度生成辅助函数 ──────────────────────────────────────
_PPT_PAGE_CONCURRENCY = 8  # 单次任务里逐页深度生成的并发上限，平衡速度和豆包 QPS 限流


async def _generate_ppt_outline(
    *, ai, subject: str, grade_level: str, region: str,
    topic: str, source: str, template_meta: dict,
) -> dict:
    """Stage A: 让豆包出 15-25 页的元数据大纲（不展开正文）。"""
    style_hint = (
        f"模板：{template_meta.get('name') or '-'}（{template_meta.get('mood') or '-'}）；"
        f"版式倾向：{template_meta.get('layout_style') or 'modern'}；"
        f"封面风格：{template_meta.get('cover_style') or 'centered'}"
    )
    prompt = (
        f"学科：{subject or '-'}，年级：{grade_level or '-'}，地区：{region or 'mainland'}\n"
        f"主题：{topic or '-'}\n"
        f"{style_hint}\n"
    )
    if source:
        prompt += f"\n参考教学内容（节选）：\n{source[:6000]}\n"
    raw = await ai.generate(
        prompt, provider_name=DOUBAO_PROVIDER,
        max_tokens=2500, system_message=PPT_OUTLINE_SYSTEM,
    )
    return _parse_json_response(raw)


async def _generate_ppt_page(
    *, ai, page_meta: dict, outline: dict, neighbors: dict,
    subject: str, grade_level: str, topic: str, source_excerpt: str,
    template_meta: dict, sem: asyncio.Semaphore,
) -> dict:
    """Stage B: 并发深度生成单页，吃 outline 上下文 + 前后页提示，确保连贯。

    单页失败时返回 layout/title 沿用元数据的 stub，避免一页坏整组挂掉。
    """
    async with sem:
        ctx_lines = [
            f"全局主题：{outline.get('title','')}" + (
                f" / 副标题：{outline.get('subtitle','')}" if outline.get('subtitle') else ""
            ),
            f"学科：{subject or '-'}，年级：{grade_level or '-'}，主题：{topic or '-'}",
            f"本页位置：第 {page_meta.get('index')}/{outline.get('total_pages') or len(outline.get('pages') or [])} 页"
            f"，章节：{page_meta.get('section') or '-'}",
            f"指定 layout：{page_meta.get('layout') or 'content'}（必须沿用，不要改）",
            f"页标题：{page_meta.get('page_title') or ''}",
            f"本页要讲清楚：{page_meta.get('key_focus') or ''}",
            f"上一页讲了：{neighbors.get('prev') or '无'}",
            f"下一页要讲：{neighbors.get('next') or '无'}",
            f"模板风格提示：{template_meta.get('mood') or template_meta.get('name') or 'modern'}",
        ]
        ctx = "\n".join(ctx_lines)
        if source_excerpt:
            ctx += f"\n\n相关参考内容片段：\n{source_excerpt[:1500]}\n"

        try:
            raw = await ai.generate(
                ctx, provider_name=DOUBAO_PROVIDER,
                max_tokens=1500, system_message=PPT_PAGE_SYSTEM,
            )
            page = _parse_json_response(raw)
        except Exception as e:
            logger.warning(
                f"[ppt] page#{page_meta.get('index')} deep-gen failed: {e}; "
                f"using stub layout={page_meta.get('layout')}"
            )
            page = {
                "layout": page_meta.get("layout") or "content",
                "title": page_meta.get("page_title") or "",
                "bullets": [page_meta.get("key_focus") or ""],
                "notes": "",
            }
        page.setdefault("layout", page_meta.get("layout") or "content")
        page.setdefault("title", page_meta.get("page_title") or "")
        return page


# ── Async pipeline helpers (called by queue worker) ──────────────────

async def _emit_tool_event(event: str, user_id: str, payload: dict) -> None:
    """Emit a socket.io event into the per-user room so any open page
    (CourseTools / Library / Header badge) can react immediately."""
    try:
        from app.main import sio
        if not sio or not user_id:
            return
        await sio.emit(event, payload, room=f"user_{user_id}")
    except Exception as e:
        logger.warning(f"[course-tool] socket emit failed event={event} err={e}")


async def _load_tool_result_or_none(db: AsyncSession, result_id: str) -> Optional[CourseToolResult]:
    r = await db.execute(
        select(CourseToolResult).where(CourseToolResult.id == result_id)
    )
    return r.scalar_one_or_none()


async def _mark_running(_db: Optional[AsyncSession], ci: CourseToolResult) -> None:
    """Fresh-session 标 running。

    长任务 worker 持有的旧 session 在 AI 调用期间会被 Supabase Pooler 单方面踢掉
    (ConnectionDoesNotExistError)，因此所有写入必须开新 session 走 pool_pre_ping
    重新 checkout，旧 session 的死活完全无关紧要。`_db` 仅为兼容旧签名保留。
    """
    async with async_session_maker() as s:
        row = (await s.execute(
            select(CourseToolResult).where(CourseToolResult.id == ci.id)
        )).scalar_one_or_none()
        if row is None:
            return
        row.status = "running"
        await s.commit()
        ci.status = row.status


async def _mark_completed(
    _db: Optional[AsyncSession], ci: CourseToolResult, data: dict,
    file_path: Optional[str] = None, extra_params: Optional[dict] = None,
) -> None:
    """Fresh-session 标 completed，参见 `_mark_running` 注释。"""
    async with async_session_maker() as s:
        row = (await s.execute(
            select(CourseToolResult).where(CourseToolResult.id == ci.id)
        )).scalar_one_or_none()
        if row is None:
            return
        row.result = data or {}
        if file_path is not None:
            row.file_path = file_path
        if extra_params:
            row.params = {**(row.params or {}), **extra_params}
        row.status = "completed"
        row.error_message = None
        await s.commit()
        ci.result = row.result
        ci.status = row.status
        ci.error_message = row.error_message
        if file_path is not None:
            ci.file_path = file_path

    await _emit_tool_event(
        "course_tool_completed",
        ci.user_id,
        {
            "result_id": ci.id,
            "tool_type": ci.tool_type,
            "title": (ci.result or {}).get("title", ""),
            "status": "completed",
        },
    )


async def _mark_failed(_db: Optional[AsyncSession], ci: CourseToolResult, err: str) -> None:
    """Fresh-session 标 failed，参见 `_mark_running` 注释。"""
    async with async_session_maker() as s:
        row = (await s.execute(
            select(CourseToolResult).where(CourseToolResult.id == ci.id)
        )).scalar_one_or_none()
        if row is None:
            return
        row.status = "failed"
        row.error_message = (err or "")[:2000]
        await s.commit()
        ci.status = row.status
        ci.error_message = row.error_message

    await _emit_tool_event(
        "course_tool_failed",
        ci.user_id,
        {
            "result_id": ci.id,
            "tool_type": ci.tool_type,
            "status": "failed",
            "error": ci.error_message,
        },
    )


def _merge_source_meta(result: dict, source_meta: Optional[dict]) -> dict:
    """Mirror the previous _save() behavior: attach _source metadata inside the result."""
    if source_meta and isinstance(result, dict):
        return {**result, "_source": {
            "type": source_meta.get("source_type"),
            "id": source_meta.get("source_id"),
            "title": source_meta.get("title"),
            "mode": source_meta.get("mode"),
        }}
    return result


async def _do_outline(db: AsyncSession, ci: CourseToolResult) -> None:
    p = ci.params or {}
    subject = p.get("subject") or ""
    grade_level = p.get("grade_level") or ""
    region = p.get("region") or "mainland"
    topic = p.get("topic") or ""
    scope = p.get("scope") or "single_lesson"
    source = p.get("_resolved_source") or ""
    source_meta = p.get("_source_meta") or {}

    scope_label = "整学期课程大纲" if scope == "semester" else "单节课教学大纲"
    prompt = (
        f"请为以下课程生成{scope_label}。\n"
        f"学科：{subject}，年级：{grade_level}，地区：{region}，主题：{topic}\n"
        f"需要分析该年龄段学生的认知特点，结合地区教学要求来设计。\n\n"
    )
    if source:
        prompt += f"参考教学内容：\n{source}\n"

    ai = _get_ai()
    raw = await ai.generate(prompt, provider_name=DOUBAO_PROVIDER, max_tokens=6000, system_message=OUTLINE_SYSTEM)
    data = _parse_json_response(raw)
    data = _merge_source_meta(data, source_meta)
    await _mark_completed(db, ci, data)


async def _do_ppt(db: AsyncSession, ci: CourseToolResult) -> None:
    """本地豆包两阶段深度生成 PPT。

    Stage A: 豆包按 PPT_OUTLINE_SYSTEM 出 15-25 页全局结构大纲。
    Stage B: asyncio.gather 并发（≤8 路），按 PPT_PAGE_SYSTEM 对每页深度展开。
    Stage C: python-pptx 用 12 种 layout 渲染输出 .pptx。

    Coze + aippt 引流路径已被实测无法返回真二进制（aippt 给的是营销网页 URL），
    这里彻底放弃，所有内容由豆包本地深度思考产出，老师拿到的是真实可用的课堂 PPT。

    任一阶段失败都退回到 last-resort：旧 PPT_SYSTEM 单轮 8000 token 一次性生成。
    """
    p = ci.params or {}
    subject       = p.get("subject") or ""
    grade_level   = p.get("grade_level") or ""
    region        = p.get("region") or "mainland"
    topic         = p.get("topic") or ""
    source        = p.get("_resolved_source") or ""
    style         = p.get("style") or "modern"

    template_meta: dict = dict(p.get("template") or {})
    palette_dict = template_meta.get("palette") or p.get("palette")
    if palette_dict and not template_meta.get("palette"):
        template_meta["palette"] = palette_dict
    template_meta.setdefault("layout_style", style)
    template_meta.setdefault("typography", "sans_display")
    template_meta.setdefault("cover_style", "centered")
    template_meta.setdefault("name", p.get("palette_name") or style)
    source_meta = p.get("_source_meta") or {}

    ai = _get_ai()
    slide_data: dict
    engine = "doubao_two_stage"

    # ── Stage A: 全局大纲 ─────────────────────────────────────────────
    pages_meta: list[dict] = []
    outline: dict = {}
    try:
        outline = await _generate_ppt_outline(
            ai=ai, subject=subject, grade_level=grade_level, region=region,
            topic=topic, source=source, template_meta=template_meta,
        )
        pages_meta = list(outline.get("pages") or [])
        logger.info(
            f"[ppt] outline ok: title={outline.get('title','')[:30]!r} "
            f"pages={len(pages_meta)} layouts={sorted({m.get('layout') for m in pages_meta})}"
        )
    except Exception as e:
        logger.warning(f"[ppt] outline stage failed: {e}; will fallback to single-shot")
        pages_meta = []

    if pages_meta:
        # ── Stage B: 并发逐页深度生成 ────────────────────────────────
        sem = asyncio.Semaphore(_PPT_PAGE_CONCURRENCY)

        async def _one(i: int, meta: dict) -> dict:
            prev_t = pages_meta[i - 1].get("page_title") if i > 0 else ""
            next_t = pages_meta[i + 1].get("page_title") if i + 1 < len(pages_meta) else ""
            return await _generate_ppt_page(
                ai=ai, page_meta=meta, outline=outline,
                neighbors={"prev": prev_t, "next": next_t},
                subject=subject, grade_level=grade_level, topic=topic,
                source_excerpt=source, template_meta=template_meta, sem=sem,
            )

        pages = await asyncio.gather(
            *[_one(i, m) for i, m in enumerate(pages_meta)],
            return_exceptions=False,
        )
        slide_data = {
            "title":    outline.get("title") or topic or "课程演示",
            "subtitle": outline.get("subtitle") or "",
            "slides":   pages,
        }
    else:
        # ── Last-resort: 旧单轮兜底 ──────────────────────────────────
        engine = "doubao_single_shot"
        style_hint = (
            f"模板名：{template_meta.get('name') or '-'}（{template_meta.get('mood') or '-'}）\n"
            f"版式倾向：{template_meta.get('layout_style') or 'modern'}\n"
            f"字体倾向：{template_meta.get('typography') or 'sans_display'}\n"
            f"封面样式：{template_meta.get('cover_style') or 'centered'}\n"
        )
        prompt = (
            f"请为以下课程生成详细的 PPT 演示文稿结构化 JSON。\n"
            f"学科：{subject}，年级：{grade_level}，地区：{region}，主题：{topic}\n"
            f"【用户选中的模板元数据】\n{style_hint}"
            f"要求：遵循 system 指令，总页数 15-25，至少出现 5 种不同 layout，每页都要有 notes。\n\n"
        )
        if source:
            prompt += f"参考内容：\n{source[:8000]}\n"
        raw = await ai.generate(
            prompt, provider_name=DOUBAO_PROVIDER,
            max_tokens=8000, system_message=PPT_SYSTEM,
        )
        slide_data = _parse_json_response(raw)

    # ── Stage C: python-pptx 渲染 ────────────────────────────────────
    from app.services.ppt_service import build_pptx
    pptx_bytes = build_pptx(
        slide_data, style=style, palette=palette_dict, template=template_meta,
    )

    fname = f"ppt_{uuid.uuid4().hex[:8]}.pptx"
    fpath = os.path.join(settings.FILES_DIR, fname)
    with open(fpath, "wb") as f:
        f.write(pptx_bytes)

    result_data = dict(slide_data or {})
    result_data["_engine"] = engine
    result_data = _merge_source_meta(result_data, source_meta)
    await _mark_completed(db, ci, result_data, file_path=fpath)


async def _do_exercises(db: AsyncSession, ci: CourseToolResult) -> None:
    p = ci.params or {}
    subject = p.get("subject") or ""
    grade_level = p.get("grade_level") or ""
    source = p.get("_resolved_source") or ""
    exercise_type = p.get("exercise_type") or "daily_homework"
    difficulty = p.get("difficulty") or "medium"
    count = int(p.get("count") or 10)
    chosen_provider = p.get("provider") or DOUBAO_PROVIDER
    source_meta = p.get("_source_meta") or {}

    type_labels = {"daily_homework": "每日作业", "quiz": "随堂测验", "exam": "考试试卷"}
    diff_labels = {"easy": "简单", "medium": "中等", "hard": "困难"}
    prompt = (
        f"请生成{type_labels.get(exercise_type, exercise_type)}。\n"
        f"学科：{subject}，年级：{grade_level}\n"
        f"难度：{diff_labels.get(difficulty, difficulty)}，题目数量：{count}\n"
        f"需要根据学生年龄段设计合适的题目。\n\n"
        f"教学内容：\n{(source or '')[:6000]}\n"
    )

    ai = _get_ai()
    raw = await ai.generate(prompt, provider_name=chosen_provider, max_tokens=8000, system_message=EXERCISE_SYSTEM)
    data = _parse_json_response(raw)
    data = _merge_source_meta(data, source_meta)
    await _mark_completed(db, ci, data)


async def _do_practice(db: AsyncSession, ci: CourseToolResult) -> None:
    p = ci.params or {}
    subject = p.get("subject") or ""
    grade_level = p.get("grade_level") or ""
    source = p.get("_resolved_source") or ""
    include_theory = bool(p.get("include_theory", True))
    chosen_provider = p.get("provider") or DOUBAO_PROVIDER
    source_meta = p.get("_source_meta") or {}

    theory_hint = "请先总结理论要点，再设计对应的课上练习和实操环节。" if include_theory else "直接设计课上练习和实操环节。"
    prompt = (
        f"请为以下课程设计课上练习和实操方案。\n"
        f"学科：{subject}，年级：{grade_level}\n"
        f"{theory_hint}\n\n"
        f"教学内容：\n{(source or '')[:6000]}\n"
    )

    ai = _get_ai()
    raw = await ai.generate(prompt, provider_name=chosen_provider, max_tokens=6000, system_message=PRACTICE_SYSTEM)
    data = _parse_json_response(raw)
    data = _merge_source_meta(data, source_meta)
    await _mark_completed(db, ci, data)


# ── 1. Outline ───────────────────────────────────────────────────────

@router.post("/outline")
async def generate_outline(
    scope: str = Form("single_lesson"),
    subject: str = Form(""),
    grade_level: str = Form(""),
    region: str = Form("mainland"),
    topic: str = Form(""),
    content: str = Form(""),
    lesson_id: Optional[str] = Form(None),
    series_id: Optional[str] = Form(None),
    outline_id: Optional[str] = Form(None),
    ppt_id: Optional[str] = Form(None),
    exercises_id: Optional[str] = Form(None),
    practice_id: Optional[str] = Form(None),
    source_mode: str = Form("auto"),
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_active_user),
):
    src = await _resolve_source_context(
        db, current_user,
        lesson_id=lesson_id, series_id=series_id, outline_id=outline_id,
        ppt_id=ppt_id, exercises_id=exercises_id, practice_id=practice_id,
        source_mode=source_mode,
    )
    subject = subject or src["subject"]
    grade_level = grade_level or src["grade_level"]
    region = region or src["region"] or "mainland"
    topic = topic or src["topic"]

    source = content or src["text"]
    if not source and not topic:
        raise HTTPException(400, "请提供教案内容、手动输入或主题")

    params = {
        "scope": scope,
        "subject": subject,
        "grade_level": grade_level,
        "region": region,
        "topic": topic,
        "series_id": series_id,
        "outline_id": outline_id,
        "ppt_id": ppt_id,
        "exercises_id": exercises_id,
        "practice_id": practice_id,
        "source_mode": source_mode,
        "_resolved_source": source,
        "_source_meta": {**src, "mode": source_mode},
    }
    placeholder_title = topic or src.get("title") or "课程大纲"
    result_id = await _create_pending(
        db, user_id=current_user.id, lesson_id=lesson_id,
        tool_type="outline", params=params,
        initial_result={"title": placeholder_title},
    )
    return await _enqueue_tool(
        db=db, user_id=current_user.id,
        kind="tool_outline", result_id=result_id, tool_type="outline",
        title=placeholder_title,
    )


@router.get("/outline/{result_id}/download")
async def download_outline(
    result_id: str,
    for_user_id: Optional[str] = Query(None, description="管理员：下载指定用户的记录"),
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_active_user),
):
    owner = await resolve_documents_owner(db, current_user, for_user_id)
    item = await _get_result(result_id, owner.id, db)
    data = item.result or {}

    from docx import Document
    from docx.shared import Pt
    doc = Document()
    style = doc.styles["Normal"]
    style.font.size = Pt(11)
    style.font.name = "SimSun"

    doc.add_heading(data.get("title", "课程大纲"), level=0)

    for sec in data.get("sections", []):
        doc.add_heading(sec.get("title", ""), level=1)
        if sec.get("duration"):
            doc.add_paragraph(f"建议时长：{sec['duration']}")
        for obj in sec.get("objectives", []):
            doc.add_paragraph(f"• {obj}")
        if sec.get("key_points"):
            doc.add_heading("要点", level=2)
            for kp in sec["key_points"]:
                doc.add_paragraph(f"  - {kp}")
        if sec.get("activities"):
            doc.add_heading("活动", level=2)
            for act in sec["activities"]:
                doc.add_paragraph(f"  - {act}")
        for sub in sec.get("sub_sections", []):
            doc.add_heading(sub.get("title", ""), level=2)
            for kp in sub.get("key_points", []):
                doc.add_paragraph(f"  - {kp}")

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return Response(content=buf.getvalue(),
                    media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    headers={"Content-Disposition": _cd(data.get("title", "大纲"), "docx")})


# ── 2. PPT ───────────────────────────────────────────────────────────

TAG_DESCRIPTIONS = {
    "childish": "童趣活泼（适合低年级，明亮暖色）",
    "academic": "学术严谨（深蓝/深灰主调，专业感）",
    "business": "商务简洁（克制中性色，品牌感）",
    "minimal": "极简留白（高对比度，大面积留白）",
    "tech": "科技未来（冷色调 + 强荧光点缀）",
    "natural": "自然温和（大地色/植物绿）",
    "artistic": "艺术个性（高饱和强对比，不拘一格）",
}


def _validate_palette(p: dict) -> dict:
    """Ensure palette has all 6 required keys and each is a valid hex color."""
    required = ["bg", "title_color", "body_color", "accent", "section_bg", "bullet_color"]
    out = {}
    import re as _re
    hex_re = _re.compile(r"^#[0-9A-Fa-f]{6}$")
    for k in required:
        v = str(p.get(k, "")).strip()
        if not hex_re.match(v):
            raise ValueError(f"palette.{k} 不是合法的 6 位 HEX 颜色: {v!r}")
        out[k] = v.upper()
    return out


def _extract_last_json_block(text: str) -> dict:
    """Try hard to extract a JSON object from a free-form AI response.

    Priority:
    1. last fenced ```json ... ``` block
    2. last fenced ``` ... ``` block whose content parses as JSON
    3. widest balanced {...} region in the string
    """
    import re as _re

    if not text:
        raise ValueError("empty stream buffer")

    # 1) last ``` ... ``` fenced block (prefer json-tagged), parse each as JSON
    fenced = _re.findall(r"```(?:json)?\s*(.+?)\s*```", text, _re.DOTALL | _re.IGNORECASE)
    for block in reversed(fenced):
        candidate = block.strip()
        if not candidate.startswith("{"):
            # block might contain surrounding prose; take the widest {...}
            lb, rb = candidate.find("{"), candidate.rfind("}")
            if lb >= 0 and rb > lb:
                candidate = candidate[lb:rb + 1]
        try:
            return json.loads(candidate)
        except Exception:
            continue

    # 2) widest balanced {...} region
    depth = 0
    start = -1
    best: Optional[tuple[int, int]] = None
    for i, ch in enumerate(text):
        if ch == "{":
            if depth == 0:
                start = i
            depth += 1
        elif ch == "}":
            if depth > 0:
                depth -= 1
                if depth == 0 and start >= 0:
                    span = (start, i + 1)
                    if best is None or (span[1] - span[0]) > (best[1] - best[0]):
                        best = span
                    start = -1
    if best:
        candidate = text[best[0]:best[1]]
        try:
            return json.loads(candidate)
        except Exception:
            pass

    raise ValueError("未找到合法的 JSON 方案块")


@router.post("/ppt/analyze-style/stream")
async def analyze_ppt_style_stream(
    subject: str = Form(""),
    grade_level: str = Form(""),
    region: str = Form("mainland"),
    topic: str = Form(""),
    style_tags: str = Form(""),
    style_description: str = Form(""),
    lesson_id: Optional[str] = Form(None),
    series_id: Optional[str] = Form(None),
    outline_id: Optional[str] = Form(None),
    ppt_id: Optional[str] = Form(None),
    exercises_id: Optional[str] = Form(None),
    practice_id: Optional[str] = Form(None),
    source_mode: str = Form("auto"),
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_active_user),
):
    """Stream a Doubao style analysis + 3 recommended templates via SSE.

    Events (text/event-stream):
      - ``event: start``      – signals the channel is open
      - ``event: delta``      – ``{"text": "..."}`` token-by-token analysis
      - ``event: final``      – ``{"candidates":[{name,mood,palette,layout_style,typography,cover_style,use_case,rationale}, ...]}``
      - ``event: error``      – ``{"message": "..."}`` if anything fails
    """
    src = await _resolve_source_context(
        db, current_user,
        lesson_id=lesson_id, series_id=series_id, outline_id=outline_id,
        ppt_id=ppt_id, exercises_id=exercises_id, practice_id=practice_id,
        source_mode=source_mode,
    )
    subject_r = subject or src["subject"]
    grade_r = grade_level or src["grade_level"]
    region_r = region or src["region"] or "mainland"
    topic_r = topic or src["topic"]

    if not topic_r and not subject_r:
        raise HTTPException(400, "请至少提供主题或学科")

    tags_list = [t.strip() for t in style_tags.split(",") if t.strip()]
    tag_lines = [f"- {t}: {TAG_DESCRIPTIONS.get(t, '')}" for t in tags_list if t in TAG_DESCRIPTIONS]
    tag_block = ("\n【用户勾选的情绪标签】\n" + "\n".join(tag_lines)) if tag_lines else "\n（用户未勾选任何情绪标签）"
    desc_block = f"\n【用户补充的风格描述】\n{style_description.strip()}" if style_description.strip() else ""

    prompt = (
        f"请为以下课程分析 PPT 视觉方向，并给出 3 个风格明显不同的模板候选。\n\n"
        f"【课程信息】\n"
        f"- 学科：{subject_r or '（未指定）'}\n"
        f"- 学段/年级：{grade_r or '（未指定）'}\n"
        f"- 地区：{region_r}\n"
        f"- 主题：{topic_r or '（未指定）'}"
        f"{tag_block}{desc_block}\n\n"
        f"按 system 指令规定的两段格式输出：先「## 分析」自然语言，再「## 推荐模板」一个 JSON 代码块（candidates 必须恰好 3 个）。"
    )

    def _sse(event: str, data: dict) -> str:
        return f"event: {event}\ndata: {json.dumps(data, ensure_ascii=False)}\n\n"

    _LAYOUT_STYLES = {"academic", "modern", "kawaii", "tech", "editorial", "minimal", "business", "natural"}
    _TYPOGRAPHIES = {"serif", "sans_display", "handwriting", "mono"}
    _COVER_STYLES = {"centered", "split", "decorative"}

    def _clean_candidate(c: dict) -> Optional[dict]:
        """Validate one candidate; return cleaned dict or None if palette fails."""
        try:
            palette = _validate_palette(c.get("palette", {}))
        except ValueError as ve:
            logger.warning(f"[analyze-style/stream] skip candidate, palette invalid: {ve}  raw={c}")
            return None
        layout_style = str(c.get("layout_style", "")).strip().lower() or "modern"
        if layout_style not in _LAYOUT_STYLES:
            layout_style = "modern"
        typography = str(c.get("typography", "")).strip().lower() or "sans_display"
        if typography not in _TYPOGRAPHIES:
            typography = "sans_display"
        cover_style = str(c.get("cover_style", "")).strip().lower() or "centered"
        if cover_style not in _COVER_STYLES:
            cover_style = "centered"
        return {
            "name": (str(c.get("name", ""))[:40] or "候选模板"),
            "mood": str(c.get("mood", ""))[:80],
            "palette": palette,
            "layout_style": layout_style,
            "typography": typography,
            "cover_style": cover_style,
            "use_case": str(c.get("use_case", ""))[:120],
            "rationale": str(c.get("rationale", ""))[:300],
        }

    async def event_gen():
        ai = _get_ai()
        yield _sse("start", {})
        buf = ""
        try:
            async for chunk in ai.generate_stream(
                prompt,
                provider_name=DOUBAO_PROVIDER,
                system_message=PPT_STYLE_ANALYZER_STREAM_SYSTEM,
                temperature=0.7,
                max_tokens=2600,
            ):
                if not chunk:
                    continue
                buf += chunk
                yield _sse("delta", {"text": chunk})
        except Exception as e:
            logger.error(f"PPT style stream AI call failed: {e}\n{traceback.format_exc()}")
            yield _sse("error", {"message": f"AI 流式调用失败: {e}"})
            return

        try:
            parsed = _extract_last_json_block(buf)
            raw_candidates = parsed.get("candidates") if isinstance(parsed, dict) else None
            # Backward compat: if the model returned a single object (old prompt), wrap it.
            if raw_candidates is None and isinstance(parsed, dict) and parsed.get("palette"):
                raw_candidates = [parsed]
            if not isinstance(raw_candidates, list) or not raw_candidates:
                raise ValueError("candidates 字段缺失或为空")

            cleaned: list[dict] = []
            for c in raw_candidates[:3]:
                if isinstance(c, dict):
                    item = _clean_candidate(c)
                    if item:
                        cleaned.append(item)
            if not cleaned:
                raise ValueError("所有候选模板的配色都未通过校验")
        except Exception as e:
            logger.warning(f"PPT style stream final parse failed: {e}  buf_tail={buf[-600:]!r}")
            yield _sse("error", {"message": f"AI 返回内容未通过校验: {e}"})
            return

        yield _sse("final", {"candidates": cleaned})

    return StreamingResponse(
        event_gen(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "X-Accel-Buffering": "no",
            "Connection": "keep-alive",
        },
    )


@router.post("/ppt/analyze-style")
async def analyze_ppt_style(
    subject: str = Form(""),
    grade_level: str = Form(""),
    region: str = Form("mainland"),
    topic: str = Form(""),
    style_tags: str = Form(""),  # comma-separated: e.g. "childish,natural"
    style_description: str = Form(""),
    lesson_id: Optional[str] = Form(None),
    series_id: Optional[str] = Form(None),
    outline_id: Optional[str] = Form(None),
    ppt_id: Optional[str] = Form(None),
    exercises_id: Optional[str] = Form(None),
    practice_id: Optional[str] = Form(None),
    source_mode: str = Form("auto"),
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_active_user),
):
    """Let Doubao propose 3 color palette candidates based on topic + tags + free-form hint."""
    src = await _resolve_source_context(
        db, current_user,
        lesson_id=lesson_id, series_id=series_id, outline_id=outline_id,
        ppt_id=ppt_id, exercises_id=exercises_id, practice_id=practice_id,
        source_mode=source_mode,
    )
    subject = subject or src["subject"]
    grade_level = grade_level or src["grade_level"]
    region = region or src["region"] or "mainland"
    topic = topic or src["topic"]

    if not topic and not subject:
        raise HTTPException(400, "请至少提供主题或学科")

    tags_list = [t.strip() for t in style_tags.split(",") if t.strip()]
    tag_lines = [f"- {t}: {TAG_DESCRIPTIONS.get(t, '')}" for t in tags_list if t in TAG_DESCRIPTIONS]
    tag_block = ("\n【用户勾选的情绪标签】\n" + "\n".join(tag_lines)) if tag_lines else "\n（用户未勾选任何情绪标签）"
    desc_block = f"\n【用户补充的风格描述】\n{style_description.strip()}" if style_description.strip() else ""

    prompt = (
        f"请为以下课程分析合适的 PPT 配色，并给出 3 个候选方案。\n\n"
        f"【课程信息】\n"
        f"- 学科：{subject or '（未指定）'}\n"
        f"- 学段/年级：{grade_level or '（未指定）'}\n"
        f"- 地区：{region}\n"
        f"- 主题：{topic or '（未指定）'}"
        f"{tag_block}{desc_block}\n\n"
        f"请严格按 system 指令要求的 JSON 结构输出，仅输出 JSON 本体，不要任何额外文字。"
    )

    ai = _get_ai()
    try:
        raw = await ai.generate(
            prompt,
            provider_name=DOUBAO_PROVIDER,
            max_tokens=2500,
            system_message=PPT_STYLE_ANALYZER_SYSTEM,
            temperature=0.6,
        )
        data = _parse_json_response(raw)
    except Exception as e:
        logger.error(f"PPT style analyze failed: {e}\n{traceback.format_exc()}")
        raise HTTPException(500, f"风格分析失败: {e}")

    candidates = data.get("candidates") if isinstance(data, dict) else None
    if not isinstance(candidates, list) or len(candidates) < 1:
        raise HTTPException(500, "AI 未返回合法的候选方案")

    clean: list = []
    for c in candidates[:3]:
        try:
            palette = _validate_palette(c.get("palette", {}))
            clean.append({
                "name": str(c.get("name", ""))[:40] or "候选方案",
                "mood": str(c.get("mood", ""))[:80],
                "palette": palette,
                "rationale": str(c.get("rationale", ""))[:300],
            })
        except ValueError as ve:
            logger.warning(f"Skip invalid candidate: {ve}  raw={c}")

    if not clean:
        raise HTTPException(500, "AI 返回的候选方案颜色格式不合法")

    return {"candidates": clean}


@router.post("/ppt")
async def generate_ppt(
    style: str = Form("modern"),
    palette: str = Form(""),
    palette_name: str = Form(""),
    template: str = Form(""),
    subject: str = Form(""),
    grade_level: str = Form(""),
    region: str = Form("mainland"),
    topic: str = Form(""),
    content: str = Form(""),
    lesson_id: Optional[str] = Form(None),
    outline_id: Optional[str] = Form(None),
    series_id: Optional[str] = Form(None),
    ppt_id: Optional[str] = Form(None),
    exercises_id: Optional[str] = Form(None),
    practice_id: Optional[str] = Form(None),
    source_mode: str = Form("auto"),
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_active_user),
):
    src = await _resolve_source_context(
        db, current_user,
        lesson_id=lesson_id, series_id=series_id, outline_id=outline_id,
        ppt_id=ppt_id, exercises_id=exercises_id, practice_id=practice_id,
        source_mode=source_mode,
    )
    subject = subject or src["subject"]
    grade_level = grade_level or src["grade_level"]
    region = region or src["region"] or "mainland"
    topic = topic or src["topic"]
    source = content or src["text"]

    if not source and not topic:
        raise HTTPException(400, "请提供内容来源")

    # --- palette (backward compat) -----------------------------------
    palette_dict: Optional[dict] = None
    if palette:
        try:
            palette_dict = _validate_palette(json.loads(palette))
        except Exception as e:
            raise HTTPException(400, f"palette 参数无效: {e}")

    # --- template (new, full metadata) -------------------------------
    _ALLOWED_LAYOUT_STYLES = {"academic", "modern", "kawaii", "tech", "editorial", "minimal", "business", "natural"}
    _ALLOWED_TYPO = {"serif", "sans_display", "handwriting", "mono"}
    _ALLOWED_COVER = {"centered", "split", "decorative"}

    template_dict: Optional[dict] = None
    if template:
        try:
            raw_tpl = json.loads(template)
            if not isinstance(raw_tpl, dict):
                raise ValueError("template 必须是 JSON 对象")
            tpl_palette = raw_tpl.get("palette")
            if tpl_palette:
                tpl_palette = _validate_palette(tpl_palette)
            else:
                tpl_palette = palette_dict
            layout_style = str(raw_tpl.get("layout_style", "")).strip().lower() or style or "modern"
            if layout_style not in _ALLOWED_LAYOUT_STYLES:
                layout_style = "modern"
            typography = str(raw_tpl.get("typography", "")).strip().lower() or "sans_display"
            if typography not in _ALLOWED_TYPO:
                typography = "sans_display"
            cover_style = str(raw_tpl.get("cover_style", "")).strip().lower() or "centered"
            if cover_style not in _ALLOWED_COVER:
                cover_style = "centered"
            template_dict = {
                "name": str(raw_tpl.get("name", ""))[:60] or (palette_name or layout_style),
                "mood": str(raw_tpl.get("mood", ""))[:120],
                "palette": tpl_palette,
                "layout_style": layout_style,
                "typography": typography,
                "cover_style": cover_style,
                "use_case": str(raw_tpl.get("use_case", ""))[:200],
                "rationale": str(raw_tpl.get("rationale", ""))[:400],
            }
        except Exception as e:
            raise HTTPException(400, f"template 参数无效: {e}")

    params = {
        "style": style,
        "subject": subject,
        "grade_level": grade_level,
        "region": region,
        "topic": topic,
        "series_id": series_id,
        "outline_id": outline_id,
        "ppt_id": ppt_id,
        "exercises_id": exercises_id,
        "practice_id": practice_id,
        "source_mode": source_mode,
        "_resolved_source": source,
        "_source_meta": {**src, "mode": source_mode},
    }
    if palette_dict:
        params["palette"] = palette_dict
        params["palette_name"] = palette_name
    if template_dict:
        params["template"] = template_dict
        # mirror palette to top-level for older readers
        if template_dict.get("palette") and "palette" not in params:
            params["palette"] = template_dict["palette"]
        if template_dict.get("name") and not params.get("palette_name"):
            params["palette_name"] = template_dict["name"]

    placeholder_title = topic or src.get("title") or "PPT 演示文稿"
    result_id = await _create_pending(
        db, user_id=current_user.id, lesson_id=lesson_id,
        tool_type="ppt", params=params,
        initial_result={"title": placeholder_title},
    )
    return await _enqueue_tool(
        db=db, user_id=current_user.id,
        kind="tool_ppt", result_id=result_id, tool_type="ppt",
        title=placeholder_title,
    )


@router.get("/ppt/{result_id}/download")
async def download_ppt(
    result_id: str,
    for_user_id: Optional[str] = Query(None, description="管理员：下载指定用户的记录"),
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_active_user),
):
    owner = await resolve_documents_owner(db, current_user, for_user_id)
    item = await _get_result(result_id, owner.id, db)
    if not item.file_path or not os.path.isfile(item.file_path):
        raise HTTPException(404, "PPT文件不存在")
    with open(item.file_path, "rb") as f:
        data = f.read()
    title = (item.result or {}).get("title", "PPT")
    return Response(content=data,
                    media_type="application/vnd.openxmlformats-officedocument.presentationml.presentation",
                    headers={"Content-Disposition": _cd(title, "pptx")})


# ── 3. Exercises ─────────────────────────────────────────────────────

@router.post("/exercises")
async def generate_exercises(
    exercise_type: str = Form("daily_homework"),
    difficulty: str = Form("medium"),
    count: int = Form(10),
    subject: str = Form(""),
    grade_level: str = Form(""),
    content: str = Form(""),
    lesson_id: Optional[str] = Form(None),
    ppt_id: Optional[str] = Form(None),
    outline_id: Optional[str] = Form(None),
    exercises_id: Optional[str] = Form(None),
    practice_id: Optional[str] = Form(None),
    series_id: Optional[str] = Form(None),
    source_mode: str = Form("auto"),
    provider: Optional[str] = Form(None),
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_active_user),
):
    src = await _resolve_source_context(
        db, current_user,
        lesson_id=lesson_id, series_id=series_id, outline_id=outline_id,
        ppt_id=ppt_id, exercises_id=exercises_id, practice_id=practice_id,
        source_mode=source_mode,
    )
    subject = subject or src["subject"]
    grade_level = grade_level or src["grade_level"]

    source = content or src["text"]
    if not source:
        raise HTTPException(400, "请提供内容来源")

    chosen_provider = _resolve_provider(provider)
    params = {
        "exercise_type": exercise_type,
        "difficulty": difficulty,
        "count": count,
        "provider": chosen_provider,
        "subject": subject,
        "grade_level": grade_level,
        "series_id": series_id,
        "outline_id": outline_id,
        "ppt_id": ppt_id,
        "exercises_id": exercises_id,
        "practice_id": practice_id,
        "source_mode": source_mode,
        "_resolved_source": source,
        "_source_meta": {**src, "mode": source_mode},
    }
    placeholder_title = src.get("title") or f"{subject} 习题" if subject else "习题"
    result_id = await _create_pending(
        db, user_id=current_user.id, lesson_id=lesson_id,
        tool_type="exercises", params=params,
        initial_result={"title": placeholder_title},
    )
    payload = await _enqueue_tool(
        db=db, user_id=current_user.id,
        kind="tool_exercises", result_id=result_id, tool_type="exercises",
        title=placeholder_title,
    )
    payload["provider"] = chosen_provider
    return payload


@router.get("/exercises/{result_id}/download")
async def download_exercises(
    result_id: str,
    version: str = "student",
    for_user_id: Optional[str] = Query(None, description="管理员：下载指定用户的记录"),
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_active_user),
):
    owner = await resolve_documents_owner(db, current_user, for_user_id)
    item = await _get_result(result_id, owner.id, db)
    data = item.result or {}

    from docx import Document
    from docx.shared import Pt
    doc = Document()
    style = doc.styles["Normal"]
    style.font.size = Pt(11)
    style.font.name = "SimSun"

    is_teacher = version == "teacher"
    suffix = "（教师版）" if is_teacher else "（学生版）"
    title = data.get("title", "习题") + suffix
    doc.add_heading(title, level=0)

    for ex in data.get("exercises", []):
        q = f"{ex.get('id', '')}. {ex.get('question', '')}"
        doc.add_paragraph(q).bold = True
        for opt in ex.get("options", []):
            doc.add_paragraph(f"    {opt}")
        if is_teacher:
            doc.add_paragraph(f"  答案：{ex.get('answer', '')}")
            doc.add_paragraph(f"  解析：{ex.get('explanation', '')}")
        doc.add_paragraph("")

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return Response(content=buf.getvalue(),
                    media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    headers={"Content-Disposition": _cd(title, "docx")})


# ── 4. Practice ──────────────────────────────────────────────────────

@router.post("/practice")
async def generate_practice(
    subject: str = Form(""),
    grade_level: str = Form(""),
    content: str = Form(""),
    include_theory: bool = Form(True),
    lesson_id: Optional[str] = Form(None),
    outline_id: Optional[str] = Form(None),
    ppt_id: Optional[str] = Form(None),
    exercises_id: Optional[str] = Form(None),
    practice_id: Optional[str] = Form(None),
    series_id: Optional[str] = Form(None),
    source_mode: str = Form("auto"),
    provider: Optional[str] = Form(None),
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_active_user),
):
    src = await _resolve_source_context(
        db, current_user,
        lesson_id=lesson_id, series_id=series_id, outline_id=outline_id,
        ppt_id=ppt_id, exercises_id=exercises_id, practice_id=practice_id,
        source_mode=source_mode,
    )
    subject = subject or src["subject"]
    grade_level = grade_level or src["grade_level"]

    source = content or src["text"]
    if not source:
        raise HTTPException(400, "请提供内容来源")

    chosen_provider = _resolve_provider(provider)
    params = {
        "subject": subject,
        "grade_level": grade_level,
        "include_theory": include_theory,
        "provider": chosen_provider,
        "series_id": series_id,
        "outline_id": outline_id,
        "ppt_id": ppt_id,
        "exercises_id": exercises_id,
        "practice_id": practice_id,
        "source_mode": source_mode,
        "_resolved_source": source,
        "_source_meta": {**src, "mode": source_mode},
    }
    placeholder_title = src.get("title") or (f"{subject} 课上练习" if subject else "课上练习")
    result_id = await _create_pending(
        db, user_id=current_user.id, lesson_id=lesson_id,
        tool_type="practice", params=params,
        initial_result={"title": placeholder_title},
    )
    payload = await _enqueue_tool(
        db=db, user_id=current_user.id,
        kind="tool_practice", result_id=result_id, tool_type="practice",
        title=placeholder_title,
    )
    payload["provider"] = chosen_provider
    payload["merge_available"] = True
    return payload


@router.post("/practice/{practice_id}/merge-ppt")
async def merge_practice_to_ppt(
    practice_id: str,
    ppt_id: str = Form(...),
    for_user_id: Optional[str] = Query(None, description="管理员：以指定用户身份合并并保存结果"),
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_active_user),
):
    owner = await resolve_documents_owner(db, current_user, for_user_id)
    practice = await _get_result(practice_id, owner.id, db)
    ppt_item = await _get_result(ppt_id, owner.id, db)

    if not ppt_item.file_path or not os.path.isfile(ppt_item.file_path):
        raise HTTPException(404, "原始PPT文件不存在")

    practice_data = practice.result or {}
    practice_slides = []
    practice_slides.append({
        "layout": "section_header",
        "title": "课上练习与实操",
        "subtitle": practice_data.get("title", ""),
        "bullets": [], "notes": "",
    })
    if practice_data.get("theory_summary"):
        practice_slides.append({
            "layout": "content",
            "title": "理论要点回顾",
            "bullets": [s.strip() for s in practice_data["theory_summary"].split("。") if s.strip()],
            "notes": "", "subtitle": "",
        })
    for p in practice_data.get("practices", []):
        practice_slides.append({
            "layout": "content",
            "title": p.get("title", "练习"),
            "bullets": [
                p.get("description", ""),
                f"时长：{p.get('duration', '')}" if p.get("duration") else "",
                f"所需材料：{p.get('materials', '')}" if p.get("materials") else "",
                f"预期成果：{p.get('expected_outcome', '')}" if p.get("expected_outcome") else "",
            ],
            "notes": "", "subtitle": "",
        })

    from app.services.ppt_service import append_slides_to_pptx
    ppt_params = ppt_item.params or {}
    merged_bytes = append_slides_to_pptx(
        ppt_item.file_path,
        practice_slides,
        style=ppt_params.get("style", "modern"),
        palette=ppt_params.get("palette"),
    )
    fname = f"merged_{uuid.uuid4().hex[:8]}.pptx"
    fpath = os.path.join(settings.FILES_DIR, fname)
    with open(fpath, "wb") as f:
        f.write(merged_bytes)

    merged_result = {**(ppt_item.result or {}), "merged_practice": True}
    item = await _save(db, owner.id, ppt_item.lesson_id, "ppt",
                       {**ppt_params, "merged_practice_id": practice_id},
                       merged_result, fpath)
    return {"id": item.id, "message": "合并成功"}


@router.get("/practice/{result_id}/download")
async def download_practice(
    result_id: str,
    for_user_id: Optional[str] = Query(None, description="管理员：下载指定用户的记录"),
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_active_user),
):
    owner = await resolve_documents_owner(db, current_user, for_user_id)
    item = await _get_result(result_id, owner.id, db)
    data = item.result or {}

    from docx import Document
    from docx.shared import Pt
    doc = Document()
    style = doc.styles["Normal"]
    style.font.size = Pt(11)
    style.font.name = "SimSun"

    doc.add_heading(data.get("title", "课上练习"), level=0)

    if data.get("theory_summary"):
        doc.add_heading("理论要点", level=1)
        doc.add_paragraph(data["theory_summary"])

    doc.add_heading("练习与实操", level=1)
    for p in data.get("practices", []):
        doc.add_heading(f"{p.get('id', '')}. {p.get('title', '')}", level=2)
        tp = {"individual": "个人", "group": "小组", "hands_on": "实操", "discussion": "讨论"}
        doc.add_paragraph(f"类型：{tp.get(p.get('type', ''), p.get('type', ''))}")
        doc.add_paragraph(p.get("description", ""))
        if p.get("duration"):
            doc.add_paragraph(f"建议时长：{p['duration']}")
        if p.get("materials"):
            doc.add_paragraph(f"所需材料：{p['materials']}")
        if p.get("expected_outcome"):
            doc.add_paragraph(f"预期成果：{p['expected_outcome']}")
        doc.add_paragraph("")

    if data.get("assessment_criteria"):
        doc.add_heading("评价标准", level=1)
        doc.add_paragraph(data["assessment_criteria"])

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return Response(content=buf.getvalue(),
                    media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    headers={"Content-Disposition": _cd(data.get("title", "课上练习"), "docx")})


# ── History ──────────────────────────────────────────────────────────

def _serialize_result(i: CourseToolResult, *, include_result: bool = False) -> dict:
    safe_params = dict(i.params or {})
    # do not leak the bulky resolved source text back on every history call
    safe_params.pop("_resolved_source", None)
    row = {
        "id": i.id,
        "tool_type": i.tool_type,
        "params": safe_params,
        "title": (i.result or {}).get("title", ""),
        "status": getattr(i, "status", None) or "completed",
        "error_message": getattr(i, "error_message", None),
        "has_file": bool(i.file_path),
        "created_at": str(i.created_at),
    }
    if include_result:
        row["result"] = i.result or {}
    return row


@router.get("/history")
async def list_history(
    lesson_id: Optional[str] = None,
    tool_type: Optional[str] = None,
    status: Optional[str] = None,
    for_user_id: Optional[str] = Query(None, description="管理员：查看指定用户的课程工具记录"),
    limit: int = 200,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_active_user),
):
    safe_limit = max(1, min(limit, 500))
    owner = await resolve_documents_owner(db, current_user, for_user_id)
    q = select(CourseToolResult).where(CourseToolResult.user_id == owner.id)
    if lesson_id:
        q = q.where(CourseToolResult.lesson_id == lesson_id)
    if tool_type:
        q = q.where(CourseToolResult.tool_type == tool_type)
    if status:
        q = q.where(CourseToolResult.status == status)
    q = q.order_by(CourseToolResult.created_at.desc()).limit(safe_limit)
    res = await db.execute(q)
    items = res.scalars().all()
    return [_serialize_result(i) for i in items]


@router.get("/results/{result_id}")
async def get_result(
    result_id: str,
    for_user_id: Optional[str] = Query(None, description="管理员：读取指定用户的记录"),
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_active_user),
):
    owner = await resolve_documents_owner(db, current_user, for_user_id)
    item = await _get_result(result_id, owner.id, db)
    return _serialize_result(item, include_result=True)


@router.delete("/results/{result_id}")
async def delete_result(
    result_id: str,
    for_user_id: Optional[str] = Query(None, description="管理员：删除指定用户的记录"),
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_active_user),
):
    owner = await resolve_documents_owner(db, current_user, for_user_id)
    item = await _get_result(result_id, owner.id, db)
    # try to remove the backing file (best-effort)
    fp = item.file_path
    if fp and os.path.isfile(fp):
        try:
            os.remove(fp)
        except Exception:
            pass
    await db.delete(item)
    await db.commit()
    return {"ok": True, "id": result_id}
