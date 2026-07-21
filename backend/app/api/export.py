from fastapi import APIRouter, Depends, HTTPException, Form, File, UploadFile, Query
from fastapi.responses import Response, StreamingResponse, FileResponse
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select
from urllib.parse import quote
from datetime import datetime, date, timedelta, timezone
from io import BytesIO
from typing import Optional
from pydantic import BaseModel
import asyncio
import json
import os
import uuid
import tempfile
import traceback
from concurrent.futures import ThreadPoolExecutor
from loguru import logger

from app.core.config import settings
from app.core.deps import get_db, get_current_active_user, require_capability, resolve_documents_owner, require_export_payment, user_access_level, ACCESS_ADMIN
from app.core.database import async_session_maker
from app.models.user import User
from app.models.lesson import LessonPlan, LessonSeries, DocumentVersion, ExportRecord

router = APIRouter(
    prefix="/export",
    tags=["导出"],
    # can_export 能力门 + 付费闸门（管理员/白名单豁免；普通用户扣 1 次导出额度）
    dependencies=[Depends(require_capability("can_export")), Depends(require_export_payment)],
)

# ───────────────────────────────────────────────────────────────────
# 同步渲染器并发控制：PDF / DOCX 渲染会阻塞事件循环，扔到线程池跑
# 用 Semaphore 限制最大并行渲染数，避免单机 OOM
# ───────────────────────────────────────────────────────────────────
def _env_int(key: str, default: int, lo: int = 1, hi: int = 64) -> int:
    try:
        return max(lo, min(hi, int(os.getenv(key, default))))
    except Exception:
        return default


_PDF_RENDER_LIMIT = _env_int("PDF_RENDER_LIMIT", 3, 1, 16)
_render_semaphore = asyncio.Semaphore(_PDF_RENDER_LIMIT)
_render_executor = ThreadPoolExecutor(max_workers=_PDF_RENDER_LIMIT, thread_name_prefix="pdf-render")


async def run_in_executor(fn, *args, **kwargs):
    """Run blocking renderer in dedicated thread pool with concurrency cap."""
    loop = asyncio.get_running_loop()
    async with _render_semaphore:
        if kwargs:
            from functools import partial
            fn = partial(fn, *args, **kwargs)
            return await loop.run_in_executor(_render_executor, fn)
        return await loop.run_in_executor(_render_executor, fn, *args)


# ───────────────────────────────────────────────────────────────────
# 临时缓存目录（系列/批量异步导出产物）
# 由 supervisor / job_handlers 在启动后周期性清理过期文件
# ───────────────────────────────────────────────────────────────────
_BACKEND_DIR = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
TMP_EXPORTS_DIR = os.path.join(_BACKEND_DIR, "storage", "tmp_exports")
os.makedirs(TMP_EXPORTS_DIR, exist_ok=True)
TMP_EXPORTS_TTL_HOURS = _env_int("TMP_EXPORTS_TTL_HOURS", 168, 1, 24 * 30)  # 默认 7 天

DEFAULT_TEMPLATE_DIR = os.path.join(
    os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))),
    "templates"
)
DEFAULT_TEMPLATE_PATH = os.path.join(DEFAULT_TEMPLATE_DIR, "旋轉對稱圖形教案.pdf")

_cached_default_template: Optional[str] = None


def _json_default(obj):
    """Handle non-serializable objects for json.dumps."""
    if isinstance(obj, (datetime, date)):
        return obj.isoformat()
    if hasattr(obj, "__dict__"):
        return str(obj)
    return str(obj)


def _content_disposition(title: str, ext: str) -> str:
    ascii_fallback = "".join(c for c in (title or "") if ord(c) < 128 and c.isalnum() or c in " _-").strip()
    if not ascii_fallback:
        ascii_fallback = "lesson_plan"
    try:
        utf8_name = quote(f"{title}.{ext}")
        return f"attachment; filename=\"{ascii_fallback}.{ext}\"; filename*=UTF-8''{utf8_name}"
    except Exception:
        return f"attachment; filename=\"{ascii_fallback}.{ext}\""


def _optimized_ready(lesson: LessonPlan) -> bool:
    fc = getattr(lesson, "final_content", None)
    return bool(isinstance(fc, dict) and fc.get("full_optimized"))


async def _get_owner_lesson(
    lesson_id: str, db: AsyncSession, owner: User, requester: Optional[User] = None
) -> LessonPlan:
    result = await db.execute(
        select(LessonPlan).where(LessonPlan.id == lesson_id, LessonPlan.user_id == owner.id)
    )
    lesson = result.scalar_one_or_none()
    if not lesson:
        raise HTTPException(status_code=404, detail="教案不存在")
    if not lesson.final_content:
        raise HTTPException(status_code=400, detail="教案尚未生成完成")
    # 非管理员：优秀教案未生成完成前不可导出
    if requester is not None and user_access_level(requester) != ACCESS_ADMIN and not _optimized_ready(lesson):
        raise HTTPException(status_code=403, detail="优秀教案尚未生成完成，暂不能导出")
    return lesson


async def _get_owner_version(
    version_id: str,
    db: AsyncSession,
    owner: User,
    lesson_id: Optional[str] = None,
) -> DocumentVersion:
    res = await db.execute(
        select(DocumentVersion).where(
            DocumentVersion.id == version_id,
            DocumentVersion.user_id == owner.id,
        )
    )
    v = res.scalar_one_or_none()
    if not v:
        raise HTTPException(status_code=404, detail="文档版本不存在")
    if lesson_id and v.lesson_plan_id and v.lesson_plan_id != lesson_id:
        raise HTTPException(status_code=400, detail="版本与教案不匹配")
    return v


async def _record_export_safely(
    db: AsyncSession,
    user_id: str,
    *,
    lesson_plan_id: Optional[str] = None,
    version_id: Optional[str] = None,
    format: str,
    file_name: str,
    file_size: Optional[int] = None,
    file_path: Optional[str] = None,
    job_id: Optional[str] = None,
    expires_at: Optional[datetime] = None,
    source_kind: str = "lesson",
    status: str = "done",
    params: Optional[dict] = None,
) -> Optional[ExportRecord]:
    """
    插入 ExportRecord；任何异常都被吞掉（导出本身已成功，不能因为 record 失败就 500）。
    返回新建的 record（异步流程会用到 record.id）。
    """
    try:
        rec = ExportRecord(
            id=str(uuid.uuid4()),
            user_id=user_id,
            lesson_plan_id=lesson_plan_id,
            version_id=version_id,
            format=format,
            file_name=file_name,
            file_size=file_size,
            file_path=file_path,
            job_id=job_id,
            expires_at=expires_at,
            source_kind=source_kind,
            status=status,
            params=params,
        )
        db.add(rec)
        await db.commit()
        await db.refresh(rec)
        return rec
    except Exception as e:
        logger.warning(f"_record_export_safely failed (non-fatal): {e}")
        try:
            await db.rollback()
        except Exception:
            pass
        return None


def _markdown_to_doc_sections(version: DocumentVersion, lesson_title: Optional[str] = None) -> dict:
    """Convert a DocumentVersion's markdown content into the legacy 'd' dict shape."""
    title = lesson_title or version.title or "教案"
    return {
        "title": title,
        "meta": {},
        "full_optimized": version.content_markdown or "",
        "full_draft": "",
        "stages": [],
    }


def _parse_fc(fc):
    """Ensure final_content is a dict."""
    if fc is None:
        return {}
    if isinstance(fc, str):
        try:
            return json.loads(fc)
        except (json.JSONDecodeError, ValueError):
            return {"raw": fc}
    if isinstance(fc, dict):
        return fc
    return {"raw": str(fc)}


@router.get("/json/{lesson_id}")
async def export_json(
    lesson_id: str,
    version_id: Optional[str] = Query(None, description="导出指定 DocumentVersion 的 JSON 元数据"),
    for_user_id: Optional[str] = Query(None, description="管理员：代导出指定用户的资源"),
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_active_user),
):
    owner = await resolve_documents_owner(db, current_user, for_user_id)
    try:
        lesson = await _get_owner_lesson(lesson_id, db, owner, current_user)
        ver = None
        if version_id:
            ver = await _get_owner_version(version_id, db, owner, lesson_id)
            export_data = {
                "title": ver.title,
                "lesson_plan_id": lesson.id,
                "lesson_title": lesson.title,
                "subject": lesson.subject,
                "grade_level": lesson.grade_level,
                "version_id": ver.id,
                "version_number": ver.version_number,
                "source_kind": ver.source_kind,
                "change_source": ver.change_source,
                "created_at": str(ver.created_at) if ver.created_at else None,
                "content_markdown": ver.content_markdown,
            }
            file_title = ver.title or lesson.title
        else:
            fc = _parse_fc(lesson.final_content)
            export_data = {
                "title": lesson.title,
                "subject": lesson.subject,
                "grade_level": lesson.grade_level,
                "topic": getattr(lesson, "topic", None) or None,
                "student_type": getattr(lesson, "student_type", None) or None,
                "region": getattr(lesson, "region", None) or None,
                "status": lesson.status,
                "created_at": str(lesson.created_at) if lesson.created_at else None,
                "completed_at": str(lesson.completed_at) if lesson.completed_at else None,
                "final_content": fc,
            }
            file_title = lesson.title

        json_str = json.dumps(export_data, ensure_ascii=False, indent=2, default=_json_default)
        body = json_str.encode("utf-8")

        await _record_export_safely(
            db, owner.id,
            lesson_plan_id=lesson_id,
            version_id=ver.id if ver else None,
            format="json",
            file_name=f"{file_title}.json",
            file_size=len(body),
        )

        return Response(
            content=body,
            media_type="application/json; charset=utf-8",
            headers={"Content-Disposition": _content_disposition(file_title, "json")},
        )
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"JSON export failed for {lesson_id}: {e}\n{traceback.format_exc()}")
        raise HTTPException(status_code=500, detail=f"导出失败: {str(e)}")


@router.get("/txt/{lesson_id}")
async def export_txt(
    lesson_id: str,
    version_id: Optional[str] = Query(None, description="导出指定 DocumentVersion 的 markdown 内容"),
    for_user_id: Optional[str] = Query(None, description="管理员：代导出指定用户的资源"),
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_active_user),
):
    owner = await resolve_documents_owner(db, current_user, for_user_id)
    try:
        lesson = await _get_owner_lesson(lesson_id, db, owner, current_user)
        ver = None
        if version_id:
            ver = await _get_owner_version(version_id, db, owner, lesson_id)
            text_body = ver.content_markdown or ""
            file_title = ver.title or lesson.title
            body = text_body.encode("utf-8")
            await _record_export_safely(
                db, owner.id,
                lesson_plan_id=lesson_id,
                version_id=ver.id,
                format="txt",
                file_name=f"{file_title}.txt",
                file_size=len(body),
            )
            return Response(
                content=body,
                media_type="text/plain; charset=utf-8",
                headers={"Content-Disposition": _content_disposition(file_title, "txt")},
            )

        fc = _parse_fc(lesson.final_content)

        lines = []
        title = lesson.title or "教案"
        lines.append("=" * 50)
        lines.append(f"  {title}")
        lines.append("=" * 50)
        lines.append(f"学科: {lesson.subject}")
        lines.append(f"学段: {lesson.grade_level}")

        topic = getattr(lesson, "topic", None)
        student_type = getattr(lesson, "student_type", None)
        if topic:
            lines.append(f"主题: {topic}")
        if student_type:
            lines.append(f"学生类别: {student_type}")
        lines.append("")

        full_draft = fc.get("full_draft", "") or ""
        full_optimized = fc.get("full_optimized", "") or ""

        if full_draft:
            lines.append("-" * 50)
            lines.append("【初步教案】")
            lines.append("-" * 50)
            lines.append(str(full_draft))
            lines.append("")

        if full_optimized:
            lines.append("-" * 50)
            lines.append("【优化教案】")
            lines.append("-" * 50)
            lines.append(str(full_optimized))
            lines.append("")

        stages = fc.get("stages", {}) or {}
        if stages and isinstance(stages, dict):
            lines.append("-" * 50)
            lines.append("【各环节详情】")
            lines.append("-" * 50)
            lines.append("")

            stage_keys = sorted(
                stages.keys(),
                key=lambda k: int(k.replace("stage_", "")) if k.startswith("stage_") and k.replace("stage_", "").isdigit() else 0,
            )
            for stage_key in stage_keys:
                stage_data = stages[stage_key]
                if isinstance(stage_data, dict):
                    name = stage_data.get("name", stage_key)
                    expert = stage_data.get("expert", "")
                    content = stage_data.get("content", "")
                    draft = stage_data.get("draft", "")

                    lines.append(f">> {name}")
                    if expert:
                        lines.append(f"   采纳专家: {expert}")
                    lines.append("")
                    if content:
                        lines.append(str(content))
                    elif draft:
                        lines.append(f"[初步草稿]\n{str(draft)}")
                    lines.append("")

        text = "\n".join(lines)
        if not text.strip():
            text = json.dumps(fc, ensure_ascii=False, indent=2, default=_json_default)
        body = text.encode("utf-8")

        await _record_export_safely(
            db, owner.id,
            lesson_plan_id=lesson_id,
            version_id=None,
            format="txt",
            file_name=f"{title}.txt",
            file_size=len(body),
        )

        return Response(
            content=body,
            media_type="text/plain; charset=utf-8",
            headers={"Content-Disposition": _content_disposition(title, "txt")},
        )
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"TXT export failed for {lesson_id}: {e}\n{traceback.format_exc()}")
        raise HTTPException(status_code=500, detail=f"导出失败: {str(e)}")


def _build_doc_sections(lesson: LessonPlan) -> dict:
    """Build structured data for document generation."""
    fc = _parse_fc(lesson.final_content)
    title = lesson.title or "教案"
    meta = {
        "学科": lesson.subject,
        "学段": lesson.grade_level,
    }
    topic = getattr(lesson, "topic", None)
    student_type = getattr(lesson, "student_type", None)
    region = getattr(lesson, "region", None)
    if topic:
        meta["主题"] = topic
    if student_type:
        meta["学生类别"] = student_type
    if region:
        meta["地区"] = region

    full_draft = fc.get("full_draft", "") or ""
    full_optimized = fc.get("full_optimized", "") or ""

    stage_list = []
    stages = fc.get("stages", {}) or {}
    if isinstance(stages, dict):
        for stage_key in sorted(stages.keys()):
            sd = stages[stage_key]
            if isinstance(sd, dict):
                stage_list.append({
                    "key": stage_key,
                    "model_name": sd.get("model_name", ""),
                    "stage_name": sd.get("stage_name", sd.get("name", stage_key)),
                    "expert": sd.get("expert", ""),
                    "content": sd.get("content", ""),
                    "draft": sd.get("draft", ""),
                })

    return {
        "title": title,
        "meta": meta,
        "full_draft": str(full_draft),
        "full_optimized": str(full_optimized),
        "stages": stage_list,
    }


@router.get("/markdown/{lesson_id}")
async def export_markdown(
    lesson_id: str,
    version_id: Optional[str] = Query(None, description="导出指定 DocumentVersion 的 markdown 原文"),
    for_user_id: Optional[str] = Query(None, description="管理员：代导出指定用户的资源"),
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_active_user),
):
    owner = await resolve_documents_owner(db, current_user, for_user_id)
    try:
        lesson = await _get_owner_lesson(lesson_id, db, owner, current_user)
        if version_id:
            ver = await _get_owner_version(version_id, db, owner, lesson_id)
            md_body = (ver.content_markdown or "").encode("utf-8")
            file_title = ver.title or lesson.title
            await _record_export_safely(
                db, owner.id,
                lesson_plan_id=lesson_id,
                version_id=ver.id,
                format="markdown",
                file_name=f"{file_title}.md",
                file_size=len(md_body),
            )
            return Response(
                content=md_body,
                media_type="text/markdown; charset=utf-8",
                headers={"Content-Disposition": _content_disposition(file_title, "md")},
            )

        d = _build_doc_sections(lesson)

        lines = [f"# {d['title']}", ""]
        for k, v in d["meta"].items():
            lines.append(f"- **{k}**: {v}")
        lines.append("")

        if d["full_optimized"]:
            lines.append("## 优化教案")
            lines.append("")
            lines.append(d["full_optimized"])
            lines.append("")

        if d["full_draft"]:
            lines.append("## 初步教案")
            lines.append("")
            lines.append(d["full_draft"])
            lines.append("")

        if d["stages"]:
            lines.append("## 各环节详情")
            lines.append("")
            for s in d["stages"]:
                header = f"{s['model_name']} - {s['stage_name']}" if s["model_name"] else s["stage_name"]
                lines.append(f"### {header}")
                if s["expert"]:
                    lines.append(f"*采纳专家: {s['expert']}*")
                lines.append("")
                lines.append(s["content"] or s["draft"] or "(无内容)")
                lines.append("")

        md_text = "\n".join(lines)
        body = md_text.encode("utf-8")
        await _record_export_safely(
            db, owner.id,
            lesson_plan_id=lesson_id,
            version_id=None,
            format="markdown",
            file_name=f"{d['title']}.md",
            file_size=len(body),
        )
        return Response(
            content=body,
            media_type="text/markdown; charset=utf-8",
            headers={"Content-Disposition": _content_disposition(d["title"], "md")},
        )
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Markdown export failed for {lesson_id}: {e}\n{traceback.format_exc()}")
        raise HTTPException(status_code=500, detail=f"导出失败: {str(e)}")


def _build_docx_bytes(d: dict) -> bytes:
    """Build DOCX from structured lesson data; returns raw bytes (blocking — call via run_in_executor)."""
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    style = doc.styles["Normal"]
    style.font.size = Pt(11)
    style.font.name = "SimSun"

    title_para = doc.add_heading(d["title"], level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for k, v in d["meta"].items():
        p = doc.add_paragraph()
        run_key = p.add_run(f"{k}: ")
        run_key.bold = True
        run_key.font.size = Pt(11)
        run_val = p.add_run(str(v))
        run_val.font.size = Pt(11)

    doc.add_paragraph("")

    if d["full_optimized"]:
        doc.add_heading("优化教案", level=1)
        for para_text in d["full_optimized"].split("\n"):
            if para_text.strip():
                p = doc.add_paragraph(para_text.strip())
                p.paragraph_format.space_after = Pt(4)
        doc.add_paragraph("")

    if d["full_draft"]:
        doc.add_heading("初步教案", level=1)
        for para_text in d["full_draft"].split("\n"):
            if para_text.strip():
                p = doc.add_paragraph(para_text.strip())
                p.paragraph_format.space_after = Pt(4)
        doc.add_paragraph("")

    if d["stages"]:
        doc.add_heading("各环节详情", level=1)
        for s in d["stages"]:
            header = f"{s['model_name']} - {s['stage_name']}" if s["model_name"] else s["stage_name"]
            doc.add_heading(header, level=2)
            if s["expert"]:
                p = doc.add_paragraph()
                run = p.add_run(f"采纳专家: {s['expert']}")
                run.italic = True
                run.font.size = Pt(10)
            content = s["content"] or s["draft"] or ""
            for para_text in content.split("\n"):
                if para_text.strip():
                    p = doc.add_paragraph(para_text.strip())
                    p.paragraph_format.space_after = Pt(4)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


@router.get("/docx/{lesson_id}")
async def export_docx(
    lesson_id: str,
    version_id: Optional[str] = Query(None, description="导出指定 DocumentVersion 的 markdown 内容"),
    for_user_id: Optional[str] = Query(None, description="管理员：代导出指定用户的资源"),
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_active_user),
):
    owner = await resolve_documents_owner(db, current_user, for_user_id)
    try:
        lesson = await _get_owner_lesson(lesson_id, db, owner, current_user)
        if version_id:
            ver = await _get_owner_version(version_id, db, owner, lesson_id)
            d = _markdown_to_doc_sections(ver, lesson_title=lesson.title)
        else:
            ver = None
            d = _build_doc_sections(lesson)

        docx_bytes = await run_in_executor(_build_docx_bytes, d)
        file_name = f"{d['title']}.docx"

        await _record_export_safely(
            db, owner.id,
            lesson_plan_id=lesson_id,
            version_id=ver.id if ver else None,
            format="docx",
            file_name=file_name,
            file_size=len(docx_bytes),
        )

        return Response(
            content=docx_bytes,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": _content_disposition(d["title"], "docx")},
        )
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"DOCX export failed for {lesson_id}: {e}\n{traceback.format_exc()}")
        raise HTTPException(status_code=500, detail=f"导出失败: {str(e)}")


_PDF_FONT_READY = False
_PDF_FONT_NAME = "Helvetica"


def _ensure_pdf_font():
    """Register a Chinese font with ReportLab / xhtml2pdf (one-time).

    Docker slim images ship no CJK fonts by default; install e.g. ``fonts-wqy-zenhei``.
    Debian/Ubuntu WenQuanYi is often ``.ttc`` — older code only scanned ``.ttf``, so PDF
    fell back to Helvetica and Chinese appeared as garbled squares.
    """
    global _PDF_FONT_READY, _PDF_FONT_NAME
    if _PDF_FONT_READY:
        return
    import platform
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    candidates = []
    if platform.system() == "Windows":
        fd = r"C:\Windows\Fonts"
        candidates = [
            ("SimHei", os.path.join(fd, "simhei.ttf"), None),
            ("MSYH", os.path.join(fd, "msyh.ttc"), 0),
            ("SimSun", os.path.join(fd, "simsun.ttc"), 0),
        ]
    else:
        env_path = (settings.PDF_CJK_FONT_PATH or "").strip()
        if env_path and os.path.isfile(env_path):
            ext = os.path.splitext(env_path)[1].lower()
            sub = 0 if ext == ".ttc" else None
            base = os.path.splitext(os.path.basename(env_path))[0]
            candidates.append((f"EnvCJK-{base}", env_path, sub))

        # Common Linux/Docker paths (fonts-wqy-zenhei → wqy-zenhei.ttc)
        for name, path, sub in (
            ("WQY-ZenHei", "/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc", 0),
            ("WQY-MicroHei", "/usr/share/fonts/truetype/wqy/wqy-microhei.ttc", 0),
        ):
            candidates.append((name, path, sub))

        for d in ("/usr/share/fonts", "/usr/local/share/fonts", os.path.expanduser("~/.fonts")):
            if os.path.isdir(d):
                for root, _, files in os.walk(d):
                    for f in files:
                        fl = f.lower()
                        if not fl.endswith((".ttf", ".ttc")):
                            continue
                        if not any(k in fl for k in ("noto", "cjk", "wqy", "simhei", "droid", "sourcehansans")):
                            continue
                        sub = 0 if fl.endswith(".ttc") else None
                        candidates.append((os.path.splitext(f)[0], os.path.join(root, f), sub))
    for name, path, sub in candidates:
        if os.path.isfile(path):
            try:
                if sub is not None:
                    pdfmetrics.registerFont(TTFont(name, path, subfontIndex=sub))
                else:
                    pdfmetrics.registerFont(TTFont(name, path))
                _PDF_FONT_NAME = name
                _PDF_FONT_READY = True
                logger.info(f"PDF font registered: {name} ({path})")
                return
            except Exception as e:
                logger.warning(f"Font register failed {name}: {e}")
    _PDF_FONT_READY = True
    logger.warning("No Chinese font found; PDF will use Helvetica")


def _build_pdf_bytes(d: dict) -> bytes:
    """Build PDF from structured lesson data using ReportLab Platypus."""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib.colors import HexColor
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable

    _ensure_pdf_font()
    fn = _PDF_FONT_NAME

    S = {
        "title": ParagraphStyle("T", fontName=fn, fontSize=18, alignment=1, spaceAfter=10, leading=26),
        "h2": ParagraphStyle("H2", fontName=fn, fontSize=14, spaceBefore=16, spaceAfter=8, leading=20),
        "h3": ParagraphStyle("H3", fontName=fn, fontSize=12, spaceBefore=12, spaceAfter=6, leading=18),
        "body": ParagraphStyle("B", fontName=fn, fontSize=10.5, leading=18, spaceAfter=3),
        "meta": ParagraphStyle("M", fontName=fn, fontSize=10.5, leading=16, spaceAfter=2),
    }

    def esc(t: str) -> str:
        return t.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

    buf = BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        leftMargin=2*cm, rightMargin=2*cm, topMargin=2*cm, bottomMargin=2*cm,
    )
    elems = []

    elems.append(Paragraph(esc(d["title"]), S["title"]))
    elems.append(Spacer(1, 6))
    for k, v in d["meta"].items():
        elems.append(Paragraph(f"<b>{esc(k)}:</b> {esc(str(v))}", S["meta"]))
    elems.append(Spacer(1, 8))
    elems.append(HRFlowable(width="100%", thickness=0.5, color=HexColor("#cccccc")))

    for label, text in [("优化教案", d["full_optimized"]), ("初步教案", d["full_draft"])]:
        if text:
            elems.append(Paragraph(esc(label), S["h2"]))
            for line in text.split("\n"):
                line = line.strip()
                if line:
                    elems.append(Paragraph(esc(line), S["body"]))

    if d["stages"]:
        elems.append(Paragraph(esc("各环节详情"), S["h2"]))
        for s in d["stages"]:
            header = f"{s['model_name']} - {s['stage_name']}" if s["model_name"] else s["stage_name"]
            elems.append(Paragraph(esc(header), S["h3"]))
            if s["expert"]:
                elems.append(Paragraph(f"<i>采纳专家: {esc(s['expert'])}</i>", S["meta"]))
            content = s["content"] or s["draft"] or ""
            for line in content.split("\n"):
                line = line.strip()
                if line:
                    elems.append(Paragraph(esc(line), S["body"]))

    doc.build(elems)
    return buf.getvalue()


@router.get("/pdf/{lesson_id}")
async def export_pdf(
    lesson_id: str,
    version_id: Optional[str] = Query(None, description="导出指定 DocumentVersion 的 markdown 内容"),
    for_user_id: Optional[str] = Query(None, description="管理员：代导出指定用户的资源"),
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_active_user),
):
    owner = await resolve_documents_owner(db, current_user, for_user_id)
    try:
        lesson = await _get_owner_lesson(lesson_id, db, owner, current_user)
        if version_id:
            ver = await _get_owner_version(version_id, db, owner, lesson_id)
            d = _markdown_to_doc_sections(ver, lesson_title=lesson.title)
        else:
            ver = None
            d = _build_doc_sections(lesson)

        pdf_bytes = await run_in_executor(_build_pdf_bytes, d)
        file_name = f"{d['title']}.pdf"

        await _record_export_safely(
            db, owner.id,
            lesson_plan_id=lesson_id,
            version_id=ver.id if ver else None,
            format="pdf",
            file_name=file_name,
            file_size=len(pdf_bytes),
        )

        return Response(
            content=pdf_bytes,
            media_type="application/pdf",
            headers={"Content-Disposition": _content_disposition(d["title"], "pdf")},
        )
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"PDF export failed for {lesson_id}: {e}\n{traceback.format_exc()}")
        raise HTTPException(status_code=500, detail=f"导出失败: {str(e)}")


# ---------------------------------------------------------------------------
# Styled PDF: AI-generated HTML matching a template format
# ---------------------------------------------------------------------------

_STYLED_PDF_SYSTEM = """你是教案排版专家。根据范本格式将教案排版为可打印PDF的HTML页面。
要求：
1. 严格按范本的表格结构、标题层级、段落排列来排版
2. 样式按范本的字号、粗细、颜色、间距、边框设计
3. 内容使用提供的教案，不得修改或删减
4. 输出完整HTML，CSS内联在<style>中，@page设A4纸、@media print优化
5. 字体: "Microsoft YaHei","SimHei","SimSun",sans-serif
6. 只输出HTML代码（从<!DOCTYPE html>开始），无其他文字"""


async def _parse_template_file(file_path: str, ext: str) -> str:
    from app.services.document_parser import DocumentParserService
    parser = DocumentParserService()
    if ext in ("txt", "md", "json"):
        return await parser._parse_txt(file_path)
    elif ext == "pdf":
        return await parser._parse_pdf(file_path)
    elif ext in ("docx", "doc"):
        return await parser._parse_docx(file_path)
    elif ext == "html":
        return await parser._parse_txt(file_path)
    return await parser._parse_txt(file_path)


def _clean_ai_html(raw: str) -> str:
    """Strip markdown code fences from AI response if present."""
    text = raw.strip()
    if text.startswith("```html"):
        text = text[7:]
    elif text.startswith("```HTML"):
        text = text[7:]
    elif text.startswith("```"):
        text = text[3:]
    if text.endswith("```"):
        text = text[:-3]
    return text.strip()


_BUILTIN_DEFAULT_TEMPLATE = """【教案范本格式】
课题名称：___
班级：___    人数：___    教材来源：___    设计者：___    时间：___

一、学生背景分析
（描述学生已有知识基础、学习特点、认知水平）

二、教学目标
（一）认知目标  （二）情意目标  （三）技能目标

三、具体目标
1. ___  2. ___  3. ___

四、学生能力分析
（一）认知能力水平  （二）学习基础与能力  （三）学习困难与挑战

五、教学内容分析
（一）教学重点  （二）教学难点

六、教学流程
| 目标代号 | 活动流程 | 时间 | 教学资源 | 评量 |
|---------|---------|------|---------|------|

七、评价方式

八、教学反思"""


async def _get_default_template_text() -> str:
    """Return cached default template text, parsing the PDF only once. Falls back to built-in template."""
    global _cached_default_template
    if _cached_default_template is not None:
        return _cached_default_template
    if os.path.exists(DEFAULT_TEMPLATE_PATH):
        try:
            _cached_default_template = await _parse_template_file(DEFAULT_TEMPLATE_PATH, "pdf")
            return _cached_default_template
        except Exception as e:
            logger.warning(f"Failed to parse default template PDF: {e}, using built-in fallback")
    else:
        logger.warning(f"Default template not found at {DEFAULT_TEMPLATE_PATH}, using built-in fallback")
    _cached_default_template = _BUILTIN_DEFAULT_TEMPLATE
    return _cached_default_template


@router.post("/styled-pdf/generate/{lesson_id}")
async def generate_styled_pdf_html(
    lesson_id: str,
    template_type: str = Form("default"),
    template_file: Optional[UploadFile] = File(None),
    content_version: str = Form("draft"),
    for_user_id: Optional[str] = Query(None, description="管理员：代指定用户入队任务"),
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_active_user),
):
    """Start background task to generate styled PDF HTML, result saved to final_content."""
    owner = await resolve_documents_owner(db, current_user, for_user_id)
    lesson = await _get_owner_lesson(lesson_id, db, owner, current_user)
    fc = _parse_fc(lesson.final_content)

    if content_version == "optimized":
        content = fc.get("full_optimized", "") or ""
    else:
        content = fc.get("full_draft", "") or ""

    if not content:
        raise HTTPException(status_code=400, detail="所选版本的教案内容为空")

    if template_type == "upload" and template_file:
        ext = template_file.filename.rsplit(".", 1)[-1].lower() if "." in template_file.filename else "txt"
        if ext not in ("txt", "md", "json", "pdf", "docx", "doc", "html"):
            raise HTTPException(status_code=400, detail="不支持的文件格式，请上传 txt、md、json、pdf、docx、doc 或 html 文件")
        with tempfile.NamedTemporaryFile(delete=False, suffix=f".{ext}") as tmp:
            tmp_path = tmp.name
            tmp.write(await template_file.read())
        try:
            template_text = await _parse_template_file(tmp_path, ext)
        finally:
            os.unlink(tmp_path)
    else:
        template_text = await _get_default_template_text()

    prompt = (
        "根据以下范本格式，将教案排版为完整HTML页面。\n\n"
        "## 范本格式：\n"
        f"{template_text[:6000]}\n\n"
        "## 教案内容：\n"
        f"{content}\n\n"
        "严格按范本格式结构排版，生成完整HTML代码。"
    )

    fc = {
        **fc,
        "styled_pdf_status": "generating",
        "styled_pdf_content_version": content_version,
        "styled_pdf_prompt": prompt,
    }
    fc.pop("styled_pdf_error", None)
    lesson.final_content = fc
    await db.commit()

    from app.tasks.queue_manager import enqueue
    await enqueue(
        target_id=lesson_id,
        kind="styled_pdf",
        user_id=owner.id,
        max_attempts=2,
    )
    await _record_export_safely(
        db, owner.id,
        lesson_plan_id=lesson_id,
        format="html",
        file_name=f"{(lesson.title or 'lesson')}_styled.html",
        source_kind="styled_pdf",
        status="queued",
        params={"content_version": content_version, "template_type": template_type},
    )
    return {"status": "queued", "message": "排版PDF生成任务已入队"}


_MATERIAL_GEN_SYSTEM = """你是一位专业的课程材料制作老师，擅长将教学知识点转换为互动式、美观的HTML演示页面。你会生成完整的、可直接运行的HTML代码，包含现代化的CSS样式和交互式JavaScript功能。"""

MATERIAL_DOUBAO_PROVIDER = "doubao"

_MATERIAL_STAGE_A_SYSTEM = """你是一位专业的课程材料策划老师。请从完整教案中提炼结构化知识点，输出严格 JSON（不要 markdown 代码块、不要解释文字）。

JSON 格式：
{
  "title": "课程标题",
  "summary": "100-200字课程概述",
  "sections": [
    {
      "id": "s1",
      "title": "知识点标题",
      "icon": "fa-book",
      "content": "200-400字详细说明，可含 <ul><li>、<strong> 等简单 HTML",
      "diagram_hint": "示意图/实验/现象的文字描述（供页面展示）",
      "quiz": [{"question": "思考题", "answer": "参考答案"}]
    }
  ]
}

硬性要求：
1. sections 至少 6 个，覆盖教案主要教学环节与核心概念
2. 每个 section 的 content 纯文本不少于 150 字（不含 HTML 标签）
3. 每个 section 至少 1 道 quiz（question + answer 均非空）
4. icon 从 Font Awesome 6 类名选取（如 fa-book、fa-flask、fa-lightbulb）
5. 只输出 JSON 本体"""


def _parse_material_json(raw: str) -> dict:
    """Parse AI JSON response for material Stage A."""
    t = raw.strip()
    if t.startswith("```"):
        first_nl = t.index("\n") if "\n" in t else 3
        t = t[first_nl + 1:]
    if t.endswith("```"):
        t = t[:-3]
    return json.loads(t.strip())


def _material_baseline(meta: dict) -> str:
    """Teacher standard + K12/special-ed hints from lesson metadata."""
    from app.services import teacher_standard as _teacher_standard
    from app.services import k12_skills as _k12_skills
    from app.services import special_ed_skills as _special_ed

    locale = meta.get("locale") or "zh-CN"
    text = "\n\n" + _teacher_standard.standard(locale)
    if _k12_skills.is_k12(meta.get("education_level")):
        text += "\n\n" + _k12_skills.lesson_skills(locale)
    if _special_ed.is_special_ed(
        subject=meta.get("subject"),
        grade_level=meta.get("grade_level"),
        topic=meta.get("topic"),
        title=meta.get("title"),
        student_type=meta.get("student_type"),
    ):
        text += "\n\n" + _special_ed.lesson_skills(locale)
    return text


def _build_material_stage_a_prompt(meta: dict) -> str:
    title = meta.get("title") or "课程演示"
    content = meta.get("content") or ""
    version = meta.get("content_version") or "draft"
    version_label = "优化教案" if version == "optimized" else "初步教案"
    return (
        f'课程标题：{title}\n'
        f'内容版本：{version_label}\n\n'
        f'完整教案内容：\n{content[:12000]}\n\n'
        "请提炼至少 6 个知识点 section，每个 content 200-400 字，每节至少 1 道思考题。"
    )


async def _generate_material_outline(ai, meta: dict) -> dict:
    """Stage A: 豆包从教案抽取结构化 JSON。"""
    from app.services.material_html_service import normalize_material_data

    prompt = _build_material_stage_a_prompt(meta)
    sys = _MATERIAL_STAGE_A_SYSTEM + _material_baseline(meta)
    raw = await ai.generate(
        prompt,
        provider_name=MATERIAL_DOUBAO_PROVIDER,
        max_tokens=8000,
        temperature=0.4,
        system_message=sys,
    )
    return normalize_material_data(_parse_material_json(raw))


async def _expand_material_sections(ai, meta: dict, data: dict) -> dict:
    """补全/扩写 sections（校验不合格时重试一次）。"""
    from app.services.material_html_service import normalize_material_data

    prompt = (
        f'课程：{meta.get("title") or "课程"}\n'
        f'当前 JSON（sections 不足或过短）：\n{json.dumps(data, ensure_ascii=False)[:6000]}\n\n'
        f'教案节选：\n{(meta.get("content") or "")[:8000]}\n\n'
        "请扩写并补全 sections：至少 6 个，每个 content ≥150 字，每节至少 1 道 quiz。"
        "只输出完整 JSON 本体。"
    )
    sys = _MATERIAL_STAGE_A_SYSTEM + _material_baseline(meta)
    raw = await ai.generate(
        prompt,
        provider_name=MATERIAL_DOUBAO_PROVIDER,
        max_tokens=8000,
        temperature=0.35,
        system_message=sys,
    )
    return normalize_material_data(_parse_material_json(raw))


def _material_fallback_prompt(meta: dict) -> str:
    """Last-resort 单轮 HTML prompt（豆包 fallback）。"""
    title = meta.get("title") or "课程演示"
    content = meta.get("content") or ""
    return f"""你是一位专业的课程材料制作老师，请为课程"{title}"生成一个丰富、交互式的HTML演示页面。

完整教案内容：
{content[:12000]}

要求：
- 完整 HTML（<!DOCTYPE html> 起），内联 CSS/JS
- 顶部控制栏：全屏、主题切换、章节导航
- 至少 6 个知识点卡片，每卡 200+ 字 + 思考题（点击显隐答案）
- 所有按钮必须绑定可运行的 JavaScript（addEventListener）
- 使用 Font Awesome CDN

请直接输出完整 HTML 代码。"""


@router.post("/material/generate/{lesson_id}")
async def generate_course_material_html(
    lesson_id: str,
    content_version: str = Form("draft"),
    for_user_id: Optional[str] = Query(None, description="管理员：代指定用户入队任务"),
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_active_user),
):
    """Start background task to generate interactive HTML course material, result saved to final_content."""
    owner = await resolve_documents_owner(db, current_user, for_user_id)
    lesson = await _get_owner_lesson(lesson_id, db, owner, current_user)
    fc = _parse_fc(lesson.final_content)

    if content_version == "optimized":
        content = fc.get("full_optimized", "") or ""
    else:
        content = fc.get("full_draft", "") or ""

    if not content:
        raise HTTPException(status_code=400, detail="所选版本的教案内容为空")

    lesson_title = lesson.title or "课程演示"

    material_meta = {
        "title": lesson_title,
        "content": content,
        "content_version": content_version,
        "locale": getattr(lesson, "locale", None) or "zh-CN",
        "subject": getattr(lesson, "subject", None) or "",
        "grade_level": getattr(lesson, "grade_level", None) or "",
        "topic": getattr(lesson, "topic", None) or "",
        "education_level": getattr(lesson, "education_level", None) or "",
        "student_type": getattr(lesson, "student_type", None) or "",
    }

    status_key = f"material_{content_version}_status"
    error_key = f"material_{content_version}_error"
    meta_key = f"material_{content_version}_meta"
    fc = {**fc, status_key: "generating", meta_key: material_meta}
    fc.pop(error_key, None)
    fc.pop(f"material_{content_version}_engine", None)
    lesson.final_content = fc
    await db.commit()

    await _record_export_safely(
        db, owner.id,
        lesson_plan_id=lesson_id,
        format="html",
        file_name=f"{(lesson.title or 'lesson')}_material_{content_version}.html",
        source_kind="material",
        status="queued",
        params={"content_version": content_version},
    )

    from app.tasks.queue_manager import enqueue
    await enqueue(
        target_id=lesson_id,
        kind=f"material_{content_version}",
        user_id=owner.id,
        max_attempts=2,
    )
    return {"status": "queued", "message": "教学材料生成任务已入队"}


# ---------------------------------------------------------------------------
# 立体几何「图片入口」：上传题目图片 → qwen-vl 识别 spec →（前端确认）→ 精确出交互3D页
# ---------------------------------------------------------------------------

_ALLOWED_IMAGE_MIME = {"image/png", "image/jpeg", "image/jpg", "image/webp", "image/gif"}
_MAX_IMAGE_BYTES = 8 * 1024 * 1024  # 8MB


@router.post("/geometry/recognize-image")
async def recognize_geometry_image(
    file: UploadFile = File(...),
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_active_user),
):
    """读题目图片 → 结构化 problem spec（回显给用户确认，不落库、图片不落盘）。"""
    mime = (file.content_type or "").lower()
    if mime not in _ALLOWED_IMAGE_MIME:
        raise HTTPException(status_code=400, detail="仅支持 PNG/JPG/WEBP/GIF 图片")
    data = await file.read()
    if not data:
        raise HTTPException(status_code=400, detail="图片为空")
    if len(data) > _MAX_IMAGE_BYTES:
        raise HTTPException(status_code=400, detail="图片过大（上限 8MB）")

    from app.services.ai_service import AIService
    from app.services.geometry_skill.vision import extract_spec_from_image
    from app.services.geometry_skill.driver import build_lesson_data

    spec = await extract_spec_from_image(data, mime, AIService())
    if not spec:
        return {"ok": False, "reason": "未能从图片识别出受支持的立体几何题，请换一张更清晰的图或改用文字。"}
    # 二次确认可解：用确定性 driver 试算，不通过则视为不支持
    try:
        build_lesson_data(spec)
    except Exception as e:
        logger.info(f"[geometry] recognized spec not solvable: {e}")
        return {"ok": False, "reason": "识别到的题目暂不在可精确求解范围内。"}

    return {"ok": True, "spec": spec, "title": spec.get("title") or ""}


@router.post("/material/generate-from-spec/{lesson_id}")
async def generate_material_from_spec(
    lesson_id: str,
    spec: str = Form(...),
    content_version: str = Form("draft"),
    for_user_id: Optional[str] = Query(None, description="管理员：代指定用户写入"),
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_active_user),
):
    """用（确认后的）立体几何 spec 精确生成交互 3D 教学页，写入 material_*_html，复用预览/下载。"""
    owner = await resolve_documents_owner(db, current_user, for_user_id)
    lesson = await _get_owner_lesson(lesson_id, db, owner, current_user)  # 校验归属

    try:
        spec_obj = json.loads(spec)
    except Exception:
        raise HTTPException(status_code=400, detail="spec 不是合法 JSON")

    from app.services.geometry_skill import generate_solid_geometry_html
    from app.services.geometry_skill.driver import UnsupportedSpec
    try:
        html = generate_solid_geometry_html(spec_obj)
    except UnsupportedSpec as e:
        raise HTTPException(status_code=400, detail=f"该题目不在可精确求解范围内：{e}")
    except Exception as e:
        logger.warning(f"[geometry] generate-from-spec failed: {e}")
        raise HTTPException(status_code=500, detail="生成失败，请重试或改用文字入口")

    version = "optimized" if content_version == "optimized" else "draft"
    await _bg_update_fc(lesson_id, {
        f"material_{version}_status": "done",
        f"material_{version}_html": html,
        f"material_{version}_engine": "edu-solid-geometry",
    })
    await _record_export_safely(
        db, owner.id, lesson_plan_id=lesson_id,
        format="html", file_name=f"{(lesson.title or 'lesson')}_material_{version}.html",
        source_kind="material", status="done",
        file_size=len(html.encode("utf-8")),
        params={"content_version": version, "variant": "solid_geometry_image"},
    )
    sio = _get_sio()
    if sio:
        try:
            await sio.emit("bg_task_complete", {
                "lesson_id": lesson_id, "task": "material",
                "content_version": version, "status": "done",
                "engine": "edu-solid-geometry",
            }, room=f"lesson_{lesson_id}")
        except Exception:
            pass
    return {"ok": True, "content_version": version}


# ---------------------------------------------------------------------------
# Background task helpers
# ---------------------------------------------------------------------------

def _get_sio():
    try:
        from app.main import sio
        return sio
    except Exception:
        return None


async def _bg_update_fc(lesson_id: str, updates: dict):
    """Merge updates into lesson.final_content in a fresh DB session."""
    async with async_session_maker() as session:
        result = await session.execute(select(LessonPlan).where(LessonPlan.id == lesson_id))
        lesson = result.scalar_one_or_none()
        if not lesson:
            return
        fc = {**_parse_fc(lesson.final_content), **updates}
        lesson.final_content = fc
        await session.commit()


async def _finalize_bg_export_record(
    lesson_id: str,
    *,
    source_kind: str,
    status: str,
    error: Optional[str] = None,
    content_version: Optional[str] = None,
    file_size: Optional[int] = None,
) -> None:
    """Mark the most recent queued ExportRecord for (lesson_id, source_kind[, content_version])
    as completed / failed. Non-fatal on any error (the bg task already succeeded/failed
    independently of the audit record)."""
    try:
        async with async_session_maker() as session:
            q = (
                select(ExportRecord)
                .where(
                    ExportRecord.lesson_plan_id == lesson_id,
                    ExportRecord.source_kind == source_kind,
                    ExportRecord.status == "queued",
                )
                .order_by(ExportRecord.created_at.desc())
                .limit(8)
            )
            rows = (await session.execute(q)).scalars().all()
            target: Optional[ExportRecord] = None
            for r in rows:
                if content_version is None:
                    target = r
                    break
                p = r.params or {}
                if (p.get("content_version") or "draft") == content_version:
                    target = r
                    break
            if not target:
                return
            target.status = status
            if status == "failed" and error:
                target.error_message = error[:2000]
            if file_size is not None:
                target.file_size = file_size
            await session.commit()
    except Exception as e:
        logger.warning(f"_finalize_bg_export_record failed (non-fatal): {e}")


async def run_styled_pdf_job(lesson_id: str):
    """`kind=styled_pdf` queue handler：从 lesson.final_content.styled_pdf_prompt 还原任务参数。"""
    async with async_session_maker() as session:
        res = await session.execute(select(LessonPlan).where(LessonPlan.id == lesson_id))
        lesson = res.scalar_one_or_none()
    if not lesson:
        logger.warning(f"[styled_pdf] lesson not found: {lesson_id}")
        return
    fc = _parse_fc(lesson.final_content)
    prompt = fc.get("styled_pdf_prompt") or ""
    if not prompt:
        await _bg_update_fc(lesson_id, {
            "styled_pdf_status": "error",
            "styled_pdf_error": "任务参数缺失：styled_pdf_prompt 为空",
        })
        return
    await _bg_generate_styled_pdf(lesson_id, prompt)


async def run_material_job(lesson_id: str):
    """`kind=material_*` queue handler。Lesson 中保存了 material_<version>_prompt。"""
    async with async_session_maker() as session:
        res = await session.execute(select(LessonPlan).where(LessonPlan.id == lesson_id))
        lesson = res.scalar_one_or_none()
    if not lesson:
        logger.warning(f"[material] lesson not found: {lesson_id}")
        return
    fc = _parse_fc(lesson.final_content)
    # 找还在 generating 的 version
    for v in ("draft", "optimized"):
        if fc.get(f"material_{v}_status") == "generating":
            meta = fc.get(f"material_{v}_meta")
            # 兼容旧任务：仍读 prompt 字段
            if not meta and fc.get(f"material_{v}_prompt"):
                content = fc.get("full_optimized", "") if v == "optimized" else fc.get("full_draft", "")
                meta = {
                    "title": lesson.title or "课程演示",
                    "content": content or "",
                    "content_version": v,
                    "locale": getattr(lesson, "locale", None) or "zh-CN",
                    "subject": getattr(lesson, "subject", None) or "",
                    "grade_level": getattr(lesson, "grade_level", None) or "",
                    "topic": getattr(lesson, "topic", None) or "",
                    "education_level": getattr(lesson, "education_level", None) or "",
                    "student_type": getattr(lesson, "student_type", None) or "",
                    "legacy_prompt": fc.get(f"material_{v}_prompt"),
                }
            if not meta:
                await _bg_update_fc(lesson_id, {
                    f"material_{v}_status": "error",
                    f"material_{v}_error": "任务参数缺失：material_*_meta 为空",
                })
                return
            await _bg_generate_material(lesson_id, v, meta)
            return
    logger.info(f"[material] no generating version for lesson {lesson_id}; skipping")


async def _bg_generate_styled_pdf(lesson_id: str, prompt: str):
    """Background task: generate styled PDF HTML via AI, save to final_content."""
    try:
        from app.services.ai_service import AIService
        ai = AIService()
        collected = []
        async for chunk in ai.generate_stream(
            prompt=prompt,
            provider_name="qwen",
            temperature=0.15,
            max_tokens=10000,
            system_message=_STYLED_PDF_SYSTEM,
        ):
            collected.append(chunk)

        html_raw = "".join(collected)
        html_clean = _clean_ai_html(html_raw)

        await _bg_update_fc(lesson_id, {
            "styled_pdf_status": "done",
            "styled_pdf_html": html_clean,
        })
        await _finalize_bg_export_record(
            lesson_id, source_kind="styled_pdf", status="done",
            file_size=len(html_clean.encode("utf-8")),
        )

        sio = _get_sio()
        if sio:
            try:
                await sio.emit("bg_task_complete", {
                    "lesson_id": lesson_id, "task": "styled_pdf", "status": "done",
                }, room=f"lesson_{lesson_id}")
            except Exception:
                pass

        logger.info(f"Styled PDF background generation complete for {lesson_id}")
    except Exception as e:
        logger.error(f"Styled PDF bg task failed for {lesson_id}: {e}\n{traceback.format_exc()}")
        await _bg_update_fc(lesson_id, {
            "styled_pdf_status": "error",
            "styled_pdf_error": str(e),
        })
        await _finalize_bg_export_record(
            lesson_id, source_kind="styled_pdf", status="failed", error=str(e),
        )
        sio = _get_sio()
        if sio:
            try:
                await sio.emit("bg_task_complete", {
                    "lesson_id": lesson_id, "task": "styled_pdf", "status": "error",
                    "error": str(e),
                }, room=f"lesson_{lesson_id}")
            except Exception:
                pass


async def _try_solid_geometry_material(lesson_id: str, content_version: str) -> Optional[str]:
    """数学立体几何题 → 用移植的 edu-solid-geometry skill 生成精确可交互 3D 教学网页。

    命中返回 HTML 字符串；非数学/非立体几何/抽取失败/题型不支持一律返回 None，
    由调用方静默回退到现有 AI 生成，保证不劣化既有能力。
    """
    try:
        async with async_session_maker() as session:
            res = await session.execute(select(LessonPlan).where(LessonPlan.id == lesson_id))
            lesson = res.scalar_one_or_none()
        if not lesson:
            return None
        fc = _parse_fc(lesson.final_content)
        content = (fc.get("full_optimized" if content_version == "optimized" else "full_draft") or "")
        if not content:
            return None

        from app.services.geometry_skill.extract import looks_like_solid_geometry, extract_spec
        if not looks_like_solid_geometry(lesson.subject, content):
            return None

        from app.services.ai_service import AIService
        from app.services.geometry_skill import generate_solid_geometry_html
        spec = await extract_spec(content, AIService(), title=(lesson.title or ""))
        if not spec:
            return None
        html = generate_solid_geometry_html(spec)
        logger.info(f"[material] solid-geometry skill hit for lesson {lesson_id} "
                    f"(body={spec.get('body')}, query={(spec.get('query') or {}).get('type')})")
        return html
    except Exception as e:
        logger.warning(f"[material] solid-geometry skill path failed (fallback to AI): {e}")
        return None


async def _try_chem_reaction_material(lesson_id: str, content_version: str) -> Optional[str]:
    """化学反应 → 用移植的 edu-chem-reaction skill 生成微观 3D 交互反应演示网页。

    命中返回 HTML；非化学/未命中内置反应预设/任何异常一律返回 None，静默回退现有 AI 生成。
    """
    try:
        async with async_session_maker() as session:
            res = await session.execute(select(LessonPlan).where(LessonPlan.id == lesson_id))
            lesson = res.scalar_one_or_none()
        if not lesson:
            return None
        fc = _parse_fc(lesson.final_content)
        content = (fc.get("full_optimized" if content_version == "optimized" else "full_draft") or "")
        if not content:
            return None

        from app.services.chem_skill import looks_like_chemistry, classify_reaction, generate_reaction_html
        if not looks_like_chemistry(lesson.subject, content):
            return None

        from app.services.ai_service import AIService
        key = await classify_reaction(content, AIService())
        if not key:
            return None
        html = generate_reaction_html(key)
        logger.info(f"[material] chem-reaction skill hit for lesson {lesson_id} (reaction={key})")
        return html
    except Exception as e:
        logger.warning(f"[material] chem-reaction skill path failed (fallback to AI): {e}")
        return None


async def _bg_generate_material(lesson_id: str, content_version: str, meta: dict):
    """Background task: doubao two-stage material (JSON → template HTML), fallback to single-shot."""
    from app.services.material_html_service import (
        build_material_html, validate_material_data,
    )

    status_key = f"material_{content_version}_status"
    html_key = f"material_{content_version}_html"
    error_key = f"material_{content_version}_error"
    engine_key = f"material_{content_version}_engine"

    async def _finish(html: str, engine: str):
        await _bg_update_fc(lesson_id, {
            status_key: "done",
            html_key: html,
            engine_key: engine,
        })
        await _finalize_bg_export_record(
            lesson_id, source_kind="material", status="done",
            content_version=content_version,
            file_size=len(html.encode("utf-8")),
        )
        sio = _get_sio()
        if sio:
            try:
                await sio.emit("bg_task_complete", {
                    "lesson_id": lesson_id, "task": "material",
                    "content_version": content_version, "status": "done",
                    "engine": engine,
                }, room=f"lesson_{lesson_id}")
            except Exception:
                pass

    try:
        # 数学立体几何题：优先走确定性精确求解 + Three.js 交互页
        skill_html = await _try_solid_geometry_material(lesson_id, content_version)
        if skill_html:
            await _finish(skill_html, "edu-solid-geometry")
            logger.info(f"Material (solid-geometry skill) complete for {lesson_id}/{content_version}")
            return

        # 化学反应 skill
        chem_html = await _try_chem_reaction_material(lesson_id, content_version)
        if chem_html:
            await _finish(chem_html, "edu-chem-reaction")
            logger.info(f"Material (chem-reaction skill) complete for {lesson_id}/{content_version}")
            return

        from app.services.ai_service import AIService
        ai = AIService()
        locale = meta.get("locale") or "zh-CN"

        # 兼容旧 prompt 任务：直接走 fallback 单轮
        if meta.get("legacy_prompt"):
            collected = []
            async for chunk in ai.generate_stream(
                prompt=meta["legacy_prompt"],
                provider_name=MATERIAL_DOUBAO_PROVIDER,
                temperature=0.7,
                max_tokens=16000,
                system_message=_MATERIAL_GEN_SYSTEM,
            ):
                collected.append(chunk)
            html_clean = _clean_ai_html("".join(collected))
            await _finish(html_clean, "doubao_single_shot")
            logger.info(f"Material (legacy single-shot) complete for {lesson_id}/{content_version}")
            return

        # Stage A: 豆包抽 JSON 结构
        data = await _generate_material_outline(ai, meta)
        ok, reason = validate_material_data(data)
        if not ok:
            logger.warning(f"[material] outline validation failed ({reason}), expanding once")
            data = await _expand_material_sections(ai, meta, data)
            ok, reason = validate_material_data(data)

        if ok:
            html_out = build_material_html(data, lang=locale)
            await _finish(html_out, "doubao_two_stage")
            logger.info(
                f"Material (doubao_two_stage) complete for {lesson_id}/{content_version} "
                f"sections={len(data.get('sections') or [])} bytes={len(html_out)}"
            )
            return

        # Fallback: 豆包单轮 HTML
        logger.warning(f"[material] two-stage validation still failed ({reason}), fallback single-shot")
        fb_prompt = _material_fallback_prompt(meta)
        collected = []
        async for chunk in ai.generate_stream(
            prompt=fb_prompt,
            provider_name=MATERIAL_DOUBAO_PROVIDER,
            temperature=0.7,
            max_tokens=16000,
            system_message=_MATERIAL_GEN_SYSTEM + _material_baseline(meta),
        ):
            collected.append(chunk)
        html_clean = _clean_ai_html("".join(collected))
        await _finish(html_clean, "doubao_single_shot")
        logger.info(f"Material (doubao_single_shot fallback) complete for {lesson_id}/{content_version}")

    except Exception as e:
        logger.error(f"Material bg task failed for {lesson_id}: {e}\n{traceback.format_exc()}")
        await _bg_update_fc(lesson_id, {
            status_key: "error",
            error_key: str(e),
        })
        await _finalize_bg_export_record(
            lesson_id, source_kind="material", status="failed", error=str(e),
            content_version=content_version,
        )
        sio = _get_sio()
        if sio:
            try:
                await sio.emit("bg_task_complete", {
                    "lesson_id": lesson_id, "task": "material",
                    "content_version": content_version, "status": "error",
                    "error": str(e),
                }, room=f"lesson_{lesson_id}")
            except Exception:
                pass


# ---------------------------------------------------------------------------


class HtmlToPdfRequest(BaseModel):
    html: str
    title: str = "lesson_plan"
    lesson_plan_id: Optional[str] = None


@router.post("/styled-pdf/html-to-pdf")
async def convert_html_to_pdf(
    body: HtmlToPdfRequest,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_active_user),
):
    """Convert an HTML string to a downloadable PDF using xhtml2pdf."""
    if not body.html.strip():
        raise HTTPException(status_code=400, detail="HTML内容为空")
    try:
        from xhtml2pdf import pisa

        _ensure_pdf_font()

        html_with_font = body.html
        if "<style" in html_with_font and _PDF_FONT_NAME != "Helvetica":
            font_css = (
                f'@page {{ size: A4; margin: 1.5cm; }}\n'
                f'body {{ font-family: "{_PDF_FONT_NAME}", "Microsoft YaHei", "SimHei", "SimSun", sans-serif; }}\n'
            )
            html_with_font = html_with_font.replace("<style>", f"<style>\n{font_css}", 1)
            html_with_font = html_with_font.replace("<style ", f"<style>{font_css}</style><style ", 1) if "<style>" not in body.html else html_with_font

        buf = BytesIO()
        pisa_status = pisa.CreatePDF(
            html_with_font.encode("utf-8"),
            dest=buf,
            encoding="utf-8",
        )
        if pisa_status.err:
            logger.warning(f"xhtml2pdf reported {pisa_status.err} errors during conversion")

        buf.seek(0)
        pdf_bytes = buf.getvalue()
        if len(pdf_bytes) < 100:
            raise Exception("生成的PDF文件过小，可能转换失败")

        await _record_export_safely(
            db, current_user.id,
            lesson_plan_id=body.lesson_plan_id,
            format="pdf",
            file_name=f"{body.title}.pdf",
            file_size=len(pdf_bytes),
            source_kind="styled_pdf",
            params={"title": body.title},
        )
        return Response(
            content=pdf_bytes,
            media_type="application/pdf",
            headers={"Content-Disposition": _content_disposition(body.title, "pdf")},
        )
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"HTML-to-PDF conversion failed: {e}\n{traceback.format_exc()}")
        raise HTTPException(status_code=500, detail=f"PDF转换失败: {str(e)}")


# ---------------------------------------------------------------------------
# Series-level exports: merged single file / zip per-lesson
# ---------------------------------------------------------------------------

ALLOWED_SERIES_FORMATS = {"docx", "pdf", "md", "txt", "json"}


async def _load_series_lessons(series_id: str, db: AsyncSession, user: User) -> tuple[LessonSeries, list[LessonPlan]]:
    r = await db.execute(
        select(LessonSeries).where(LessonSeries.id == series_id, LessonSeries.user_id == user.id)
    )
    series = r.scalar_one_or_none()
    if not series:
        raise HTTPException(status_code=404, detail="系列不存在")
    r2 = await db.execute(
        select(LessonPlan)
        .where(LessonPlan.sequence_id == series_id, LessonPlan.user_id == user.id)
        .order_by(LessonPlan.sequence_order)
    )
    lessons = list(r2.scalars().all())
    if not lessons:
        raise HTTPException(status_code=400, detail="该系列尚未生成任何教案")
    return series, lessons


def _lesson_plain_text(d: dict) -> str:
    lines = [d["title"], "=" * 40]
    for k, v in d["meta"].items():
        lines.append(f"{k}: {v}")
    lines.append("")
    if d["full_optimized"]:
        lines.append("【优化教案】")
        lines.append(d["full_optimized"])
        lines.append("")
    if d["full_draft"]:
        lines.append("【初步教案】")
        lines.append(d["full_draft"])
        lines.append("")
    if d["stages"]:
        lines.append("【各环节详情】")
        for s in d["stages"]:
            header = f"{s['model_name']} - {s['stage_name']}" if s["model_name"] else s["stage_name"]
            lines.append(f">> {header}")
            if s["expert"]:
                lines.append(f"   采纳专家: {s['expert']}")
            lines.append(s["content"] or s["draft"] or "")
            lines.append("")
    return "\n".join(lines)


def _lesson_markdown(d: dict, level_offset: int = 0) -> str:
    H = lambda n: "#" * max(1, n + level_offset)
    lines = [f"{H(1)} {d['title']}", ""]
    for k, v in d["meta"].items():
        lines.append(f"- **{k}**: {v}")
    lines.append("")
    if d["full_optimized"]:
        lines.append(f"{H(2)} 优化教案")
        lines.append("")
        lines.append(d["full_optimized"])
        lines.append("")
    if d["full_draft"]:
        lines.append(f"{H(2)} 初步教案")
        lines.append("")
        lines.append(d["full_draft"])
        lines.append("")
    if d["stages"]:
        lines.append(f"{H(2)} 各环节详情")
        lines.append("")
        for s in d["stages"]:
            header = f"{s['model_name']} - {s['stage_name']}" if s["model_name"] else s["stage_name"]
            lines.append(f"{H(3)} {header}")
            if s["expert"]:
                lines.append(f"*采纳专家: {s['expert']}*")
            lines.append("")
            lines.append(s["content"] or s["draft"] or "(无内容)")
            lines.append("")
    return "\n".join(lines)


def _fetch_series_exercises_map(series_id: str, user_id: str) -> dict:
    """Pre-fetch all exercises/practice CourseToolResult keyed by lesson_id for a series.

    Returns empty dict if the table isn't available.
    """
    return {}


async def _fetch_exercises_for_lessons(
    db: AsyncSession, user_id: str, lesson_ids: list[str]
) -> dict[str, list[dict]]:
    """Return {lesson_id: [ {tool_type, title, result} ... ]}."""
    from app.models.course_tool import CourseToolResult
    if not lesson_ids:
        return {}
    r = await db.execute(
        select(CourseToolResult)
        .where(
            CourseToolResult.user_id == user_id,
            CourseToolResult.lesson_id.in_(lesson_ids),
            CourseToolResult.tool_type.in_(["exercises", "practice"]),
        )
        .order_by(CourseToolResult.created_at.desc())
    )
    items = r.scalars().all()
    out: dict[str, list[dict]] = {}
    for it in items:
        out.setdefault(it.lesson_id, []).append({
            "tool_type": it.tool_type,
            "title": (it.result or {}).get("title", ""),
            "result": it.result or {},
        })
    return out


def _render_exercises_text(block: dict) -> str:
    """Turn a CourseToolResult.result (exercises or practice) into plain text."""
    tt = block["tool_type"]
    data = block["result"] or {}
    out: list[str] = []
    if tt == "exercises":
        out.append(f"【习题：{data.get('title', '')}】")
        for ex in data.get("exercises", []):
            out.append(f"{ex.get('id', '')}. {ex.get('question', '')}")
            for opt in ex.get("options", []) or []:
                out.append(f"    {opt}")
            if ex.get("answer"):
                out.append(f"  答案：{ex['answer']}")
            if ex.get("explanation"):
                out.append(f"  解析：{ex['explanation']}")
            out.append("")
    elif tt == "practice":
        out.append(f"【课上练习：{data.get('title', '')}】")
        if data.get("theory_summary"):
            out.append("理论要点：")
            out.append(data["theory_summary"])
            out.append("")
        for p in data.get("practices", []):
            out.append(f"- {p.get('title', '')}: {p.get('description', '')}")
        if data.get("assessment_criteria"):
            out.append("评价标准：")
            out.append(data["assessment_criteria"])
    return "\n".join(out)


def _render_exercises_md(block: dict, level_offset: int = 0) -> str:
    H = lambda n: "#" * max(1, n + level_offset)
    tt = block["tool_type"]
    data = block["result"] or {}
    out: list[str] = []
    if tt == "exercises":
        out.append(f"{H(3)} 习题：{data.get('title', '')}")
        out.append("")
        for ex in data.get("exercises", []):
            out.append(f"**{ex.get('id', '')}. {ex.get('question', '')}**")
            for opt in ex.get("options", []) or []:
                out.append(f"- {opt}")
            if ex.get("answer"):
                out.append(f"> 答案：{ex['answer']}")
            if ex.get("explanation"):
                out.append(f"> 解析：{ex['explanation']}")
            out.append("")
    elif tt == "practice":
        out.append(f"{H(3)} 课上练习：{data.get('title', '')}")
        out.append("")
        if data.get("theory_summary"):
            out.append(f"{H(4)} 理论要点")
            out.append("")
            out.append(data["theory_summary"])
            out.append("")
        for p in data.get("practices", []):
            out.append(f"- **{p.get('title', '')}**: {p.get('description', '')}")
        if data.get("assessment_criteria"):
            out.append("")
            out.append(f"{H(4)} 评价标准")
            out.append("")
            out.append(data["assessment_criteria"])
    return "\n".join(out)


def _render_exercises_docx(doc, block: dict):
    from docx.shared import Pt
    tt = block["tool_type"]
    data = block["result"] or {}
    if tt == "exercises":
        doc.add_heading(f"习题：{data.get('title', '')}", level=3)
        for ex in data.get("exercises", []):
            p = doc.add_paragraph()
            run = p.add_run(f"{ex.get('id', '')}. {ex.get('question', '')}")
            run.bold = True
            for opt in ex.get("options", []) or []:
                doc.add_paragraph(f"    {opt}")
            if ex.get("answer"):
                doc.add_paragraph(f"  答案：{ex['answer']}")
            if ex.get("explanation"):
                doc.add_paragraph(f"  解析：{ex['explanation']}")
    elif tt == "practice":
        doc.add_heading(f"课上练习：{data.get('title', '')}", level=3)
        if data.get("theory_summary"):
            doc.add_heading("理论要点", level=4)
            doc.add_paragraph(data["theory_summary"])
        for p_item in data.get("practices", []):
            doc.add_paragraph(f"- {p_item.get('title', '')}: {p_item.get('description', '')}")
        if data.get("assessment_criteria"):
            doc.add_heading("评价标准", level=4)
            doc.add_paragraph(data["assessment_criteria"])


def _render_exercises_pdf_elems(block: dict, S):
    from reportlab.platypus import Paragraph
    def esc(t: str) -> str:
        return (t or "").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    elems = []
    tt = block["tool_type"]
    data = block["result"] or {}
    if tt == "exercises":
        elems.append(Paragraph(esc(f"习题：{data.get('title', '')}"), S["h3"]))
        for ex in data.get("exercises", []):
            elems.append(Paragraph(f"<b>{esc(str(ex.get('id', '')))}. {esc(ex.get('question', ''))}</b>", S["body"]))
            for opt in ex.get("options", []) or []:
                elems.append(Paragraph(esc(opt), S["body"]))
            if ex.get("answer"):
                elems.append(Paragraph(f"<i>答案：{esc(ex['answer'])}</i>", S["body"]))
            if ex.get("explanation"):
                elems.append(Paragraph(f"<i>解析：{esc(ex['explanation'])}</i>", S["body"]))
    elif tt == "practice":
        elems.append(Paragraph(esc(f"课上练习：{data.get('title', '')}"), S["h3"]))
        if data.get("theory_summary"):
            elems.append(Paragraph(esc(data["theory_summary"]), S["body"]))
        for p_item in data.get("practices", []):
            elems.append(Paragraph(esc(f"- {p_item.get('title', '')}: {p_item.get('description', '')}"), S["body"]))
        if data.get("assessment_criteria"):
            elems.append(Paragraph(f"<i>评价标准：{esc(data['assessment_criteria'])}</i>", S["body"]))
    return elems


def _series_prefix(series: LessonSeries, lesson: LessonPlan, idx: int) -> str:
    """Compute W{week}-L{lesson_num}-{order} style prefix for file names & section headings."""
    lpw = max(1, series.lessons_per_week or 1)
    order = lesson.sequence_order or (idx + 1)
    week = ((order - 1) // lpw) + 1
    sub = ((order - 1) % lpw) + 1
    return f"W{week:02d}-L{sub}"


@router.get("/series/{series_id}/export-merged")
async def export_series_merged(
    series_id: str,
    format: str = "docx",
    include_exercises: bool = False,
    for_user_id: Optional[str] = Query(None, description="管理员：代导出指定用户的系列"),
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_active_user),
):
    """Merge all lessons in a series into a single file (docx/pdf/md/txt/json)."""
    owner = await resolve_documents_owner(db, current_user, for_user_id)
    fmt = (format or "docx").lower()
    if fmt not in ALLOWED_SERIES_FORMATS:
        raise HTTPException(status_code=400, detail=f"不支持的格式：{fmt}")

    try:
        series, lessons = await _load_series_lessons(series_id, db, owner)
    except HTTPException:
        raise

    lesson_ids = [l.id for l in lessons]
    exercises_map: dict[str, list[dict]] = {}
    if include_exercises:
        exercises_map = await _fetch_exercises_for_lessons(db, owner.id, lesson_ids)

    built = []
    for idx, lesson in enumerate(lessons):
        if not lesson.final_content:
            continue
        d = _build_doc_sections(lesson)
        built.append({"idx": idx, "lesson": lesson, "data": d,
                      "prefix": _series_prefix(series, lesson, idx)})

    if not built:
        raise HTTPException(status_code=400, detail="系列中没有任何已完成的教案")

    series_title = series.title or "系列教案"

    bundle_params = {"series_id": series_id, "format": fmt, "include_exercises": include_exercises, "type": "merged"}

    if fmt == "json":
        payload = {
            "series": {
                "id": series.id,
                "title": series.title,
                "subject": series.subject,
                "grade_level": series.grade_level,
                "total_weeks": series.total_weeks,
                "lessons_per_week": series.lessons_per_week,
                "education_level": getattr(series, "education_level", "k12"),
                "major": getattr(series, "major", None),
            },
            "lessons": [],
        }
        for b in built:
            item = {
                "prefix": b["prefix"],
                "sequence_order": b["lesson"].sequence_order,
                "title": b["lesson"].title,
                "final_content": _parse_fc(b["lesson"].final_content),
            }
            if include_exercises:
                item["exercises"] = exercises_map.get(b["lesson"].id, [])
            payload["lessons"].append(item)
        text = json.dumps(payload, ensure_ascii=False, indent=2, default=_json_default)
        body = text.encode("utf-8")
        await _record_export_safely(
            db, owner.id, lesson_plan_id=None, version_id=None,
            format="json", file_name=f"{series_title}.json",
            file_size=len(body), source_kind="bundle", params=bundle_params,
        )
        return Response(
            content=body,
            media_type="application/json; charset=utf-8",
            headers={"Content-Disposition": _content_disposition(series_title, "json")},
        )

    if fmt == "txt":
        parts = [series_title, "=" * 40, ""]
        if series.subject:
            parts.append(f"学科：{series.subject}")
        if series.grade_level:
            parts.append(f"学段：{series.grade_level}")
        if getattr(series, "major", None):
            parts.append(f"专业：{series.major}")
        parts.append(f"总课时：{len(built)}")
        parts.append("")
        for b in built:
            parts.append(f"====== {b['prefix']}  {b['data']['title']} ======")
            parts.append(_lesson_plain_text(b["data"]))
            if include_exercises:
                for blk in exercises_map.get(b["lesson"].id, []):
                    parts.append("")
                    parts.append(_render_exercises_text(blk))
            parts.append("")
        body = "\n".join(parts).encode("utf-8")
        await _record_export_safely(
            db, owner.id, lesson_plan_id=None, version_id=None,
            format="txt", file_name=f"{series_title}.txt",
            file_size=len(body), source_kind="bundle", params=bundle_params,
        )
        return Response(
            content=body,
            media_type="text/plain; charset=utf-8",
            headers={"Content-Disposition": _content_disposition(series_title, "txt")},
        )

    if fmt == "md":
        parts = [f"# {series_title}", ""]
        if series.subject:
            parts.append(f"- **学科**: {series.subject}")
        if series.grade_level:
            parts.append(f"- **学段**: {series.grade_level}")
        if getattr(series, "major", None):
            parts.append(f"- **专业**: {series.major}")
        parts.append(f"- **总课时**: {len(built)}")
        parts.append("")
        for b in built:
            parts.append(f"## {b['prefix']}  {b['data']['title']}")
            parts.append("")
            # shift lesson's own H1 to H3 within the merged doc
            parts.append(_lesson_markdown(b["data"], level_offset=2))
            if include_exercises:
                for blk in exercises_map.get(b["lesson"].id, []):
                    parts.append("")
                    parts.append(_render_exercises_md(blk, level_offset=2))
            parts.append("")
        body = "\n".join(parts).encode("utf-8")
        await _record_export_safely(
            db, owner.id, lesson_plan_id=None, version_id=None,
            format="markdown", file_name=f"{series_title}.md",
            file_size=len(body), source_kind="bundle", params=bundle_params,
        )
        return Response(
            content=body,
            media_type="text/markdown; charset=utf-8",
            headers={"Content-Disposition": _content_disposition(series_title, "md")},
        )

    if fmt == "docx":
        docx_bytes = await run_in_executor(
            _build_series_merged_docx, series_title, built, series, include_exercises, exercises_map
        )
        await _record_export_safely(
            db, owner.id, lesson_plan_id=None, version_id=None,
            format="docx", file_name=f"{series_title}.docx",
            file_size=len(docx_bytes), source_kind="bundle", params=bundle_params,
        )
        return Response(
            content=docx_bytes,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": _content_disposition(series_title, "docx")},
        )

    if fmt == "pdf":
        pdf_bytes = await run_in_executor(
            _build_series_merged_pdf, series_title, built, series, include_exercises, exercises_map
        )
        await _record_export_safely(
            db, owner.id, lesson_plan_id=None, version_id=None,
            format="pdf", file_name=f"{series_title}.pdf",
            file_size=len(pdf_bytes), source_kind="bundle", params=bundle_params,
        )
        return Response(
            content=pdf_bytes,
            media_type="application/pdf",
            headers={"Content-Disposition": _content_disposition(series_title, "pdf")},
        )

    raise HTTPException(status_code=400, detail=f"暂不支持的格式：{fmt}")


def _build_series_merged_docx(
    series_title: str,
    built: list,
    series: LessonSeries,
    include_exercises: bool,
    exercises_map: dict,
) -> bytes:
    """Synchronous DOCX builder for series merged export. Call via run_in_executor."""
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    style = doc.styles["Normal"]
    style.font.size = Pt(11)
    style.font.name = "SimSun"

    title_para = doc.add_heading(series_title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_lines: list[tuple[str, str]] = []
    if series.subject:
        meta_lines.append(("学科", series.subject))
    if series.grade_level:
        meta_lines.append(("学段", series.grade_level))
    if getattr(series, "major", None):
        meta_lines.append(("专业", series.major))
    meta_lines.append(("总课时", str(len(built))))
    for k, v in meta_lines:
        p = doc.add_paragraph()
        rk = p.add_run(f"{k}: ")
        rk.bold = True
        p.add_run(str(v))
    doc.add_paragraph("")

    for b in built:
        doc.add_page_break()
        doc.add_heading(f"{b['prefix']}  {b['data']['title']}", level=1)
        d = b["data"]
        for k, v in d["meta"].items():
            p = doc.add_paragraph()
            rk = p.add_run(f"{k}: ")
            rk.bold = True
            p.add_run(str(v))
        doc.add_paragraph("")
        if d["full_optimized"]:
            doc.add_heading("优化教案", level=2)
            for line in d["full_optimized"].split("\n"):
                if line.strip():
                    doc.add_paragraph(line.strip())
        if d["full_draft"]:
            doc.add_heading("初步教案", level=2)
            for line in d["full_draft"].split("\n"):
                if line.strip():
                    doc.add_paragraph(line.strip())
        if d["stages"]:
            doc.add_heading("各环节详情", level=2)
            for s in d["stages"]:
                header = f"{s['model_name']} - {s['stage_name']}" if s["model_name"] else s["stage_name"]
                doc.add_heading(header, level=3)
                if s["expert"]:
                    pp = doc.add_paragraph()
                    run = pp.add_run(f"采纳专家: {s['expert']}")
                    run.italic = True
                for line in (s["content"] or s["draft"] or "").split("\n"):
                    if line.strip():
                        doc.add_paragraph(line.strip())
        if include_exercises:
            for blk in exercises_map.get(b["lesson"].id, []):
                _render_exercises_docx(doc, blk)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _build_series_merged_pdf(
    series_title: str,
    built: list,
    series: LessonSeries,
    include_exercises: bool,
    exercises_map: dict,
) -> bytes:
    """Synchronous PDF builder for series merged export. Call via run_in_executor."""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib.colors import HexColor
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable, PageBreak

    _ensure_pdf_font()
    fn = _PDF_FONT_NAME
    S = {
        "title": ParagraphStyle("T", fontName=fn, fontSize=18, alignment=1, spaceAfter=10, leading=26),
        "h1": ParagraphStyle("H1", fontName=fn, fontSize=15, spaceBefore=18, spaceAfter=10, leading=22),
        "h2": ParagraphStyle("H2", fontName=fn, fontSize=13, spaceBefore=14, spaceAfter=6, leading=20),
        "h3": ParagraphStyle("H3", fontName=fn, fontSize=12, spaceBefore=10, spaceAfter=6, leading=18),
        "body": ParagraphStyle("B", fontName=fn, fontSize=10.5, leading=18, spaceAfter=3),
        "meta": ParagraphStyle("M", fontName=fn, fontSize=10.5, leading=16, spaceAfter=2),
    }

    def esc(t: str) -> str:
        return (t or "").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

    buf = BytesIO()
    pdoc = SimpleDocTemplate(
        buf, pagesize=A4,
        leftMargin=2 * cm, rightMargin=2 * cm, topMargin=2 * cm, bottomMargin=2 * cm,
    )
    elems = [Paragraph(esc(series_title), S["title"]), Spacer(1, 6)]
    if series.subject:
        elems.append(Paragraph(f"<b>学科:</b> {esc(series.subject)}", S["meta"]))
    if series.grade_level:
        elems.append(Paragraph(f"<b>学段:</b> {esc(series.grade_level)}", S["meta"]))
    if getattr(series, "major", None):
        elems.append(Paragraph(f"<b>专业:</b> {esc(series.major)}", S["meta"]))
    elems.append(Paragraph(f"<b>总课时:</b> {len(built)}", S["meta"]))
    elems.append(HRFlowable(width="100%", thickness=0.5, color=HexColor("#cccccc")))

    for b in built:
        elems.append(PageBreak())
        elems.append(Paragraph(esc(f"{b['prefix']}  {b['data']['title']}"), S["h1"]))
        d = b["data"]
        for k, v in d["meta"].items():
            elems.append(Paragraph(f"<b>{esc(k)}:</b> {esc(str(v))}", S["meta"]))
        elems.append(Spacer(1, 6))
        for label, text in [("优化教案", d["full_optimized"]), ("初步教案", d["full_draft"])]:
            if text:
                elems.append(Paragraph(esc(label), S["h2"]))
                for line in text.split("\n"):
                    line = line.strip()
                    if line:
                        elems.append(Paragraph(esc(line), S["body"]))
        if d["stages"]:
            elems.append(Paragraph(esc("各环节详情"), S["h2"]))
            for s in d["stages"]:
                header = f"{s['model_name']} - {s['stage_name']}" if s["model_name"] else s["stage_name"]
                elems.append(Paragraph(esc(header), S["h3"]))
                if s["expert"]:
                    elems.append(Paragraph(f"<i>采纳专家: {esc(s['expert'])}</i>", S["meta"]))
                for line in (s["content"] or s["draft"] or "").split("\n"):
                    line = line.strip()
                    if line:
                        elems.append(Paragraph(esc(line), S["body"]))
        if include_exercises:
            for blk in exercises_map.get(b["lesson"].id, []):
                elems.extend(_render_exercises_pdf_elems(blk, S))

    pdoc.build(elems)
    return buf.getvalue()


def _build_series_zip_bytes(
    fmt: str,
    series: LessonSeries,
    lessons_built: list,
    include_exercises: bool,
    exercises_map: dict,
) -> bytes:
    """Synchronous ZIP builder for series. Each `lessons_built` item: {idx, lesson, data, prefix, base}."""
    import zipfile

    buf = BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        for entry in lessons_built:
            lesson = entry["lesson"]
            d = entry["data"]
            base = entry["base"]
            prefix = entry["prefix"]
            ext_blocks = exercises_map.get(lesson.id, []) if include_exercises else []

            if fmt == "txt":
                text = _lesson_plain_text(d)
                for blk in ext_blocks:
                    text += "\n\n" + _render_exercises_text(blk)
                zf.writestr(f"{base}.txt", text.encode("utf-8"))
            elif fmt == "md":
                text = _lesson_markdown(d)
                for blk in ext_blocks:
                    text += "\n\n" + _render_exercises_md(blk)
                zf.writestr(f"{base}.md", text.encode("utf-8"))
            elif fmt == "json":
                payload = {
                    "prefix": prefix,
                    "sequence_order": lesson.sequence_order,
                    "title": lesson.title,
                    "final_content": _parse_fc(lesson.final_content),
                }
                if include_exercises:
                    payload["exercises"] = ext_blocks
                zf.writestr(
                    f"{base}.json",
                    json.dumps(payload, ensure_ascii=False, indent=2, default=_json_default).encode("utf-8"),
                )
            elif fmt == "docx":
                from docx import Document
                from docx.shared import Pt
                doc = Document()
                st = doc.styles["Normal"]
                st.font.size = Pt(11)
                st.font.name = "SimSun"
                doc.add_heading(d["title"], level=0)
                for k, v in d["meta"].items():
                    p = doc.add_paragraph()
                    rk = p.add_run(f"{k}: ")
                    rk.bold = True
                    p.add_run(str(v))
                if d["full_optimized"]:
                    doc.add_heading("优化教案", level=1)
                    for line in d["full_optimized"].split("\n"):
                        if line.strip():
                            doc.add_paragraph(line.strip())
                if d["full_draft"]:
                    doc.add_heading("初步教案", level=1)
                    for line in d["full_draft"].split("\n"):
                        if line.strip():
                            doc.add_paragraph(line.strip())
                if d["stages"]:
                    doc.add_heading("各环节详情", level=1)
                    for s in d["stages"]:
                        header = f"{s['model_name']} - {s['stage_name']}" if s["model_name"] else s["stage_name"]
                        doc.add_heading(header, level=2)
                        if s["expert"]:
                            pp = doc.add_paragraph()
                            run = pp.add_run(f"采纳专家: {s['expert']}")
                            run.italic = True
                        for line in (s["content"] or s["draft"] or "").split("\n"):
                            if line.strip():
                                doc.add_paragraph(line.strip())
                for blk in ext_blocks:
                    _render_exercises_docx(doc, blk)
                sub = BytesIO()
                doc.save(sub)
                zf.writestr(f"{base}.docx", sub.getvalue())
            elif fmt == "pdf":
                pdf_bytes = _build_pdf_bytes(d)
                zf.writestr(f"{base}.pdf", pdf_bytes)
                if ext_blocks:
                    from reportlab.lib.pagesizes import A4
                    from reportlab.lib.styles import ParagraphStyle
                    from reportlab.lib.units import cm
                    from reportlab.platypus import SimpleDocTemplate
                    _ensure_pdf_font()
                    fn = _PDF_FONT_NAME
                    S = {
                        "h3": ParagraphStyle("H3", fontName=fn, fontSize=12, spaceBefore=10, spaceAfter=6, leading=18),
                        "body": ParagraphStyle("B", fontName=fn, fontSize=10.5, leading=18, spaceAfter=3),
                    }
                    sub = BytesIO()
                    pdoc = SimpleDocTemplate(
                        sub, pagesize=A4,
                        leftMargin=2 * cm, rightMargin=2 * cm, topMargin=2 * cm, bottomMargin=2 * cm,
                    )
                    elems: list = []
                    for blk in ext_blocks:
                        elems.extend(_render_exercises_pdf_elems(blk, S))
                    if elems:
                        pdoc.build(elems)
                        zf.writestr(f"{base}-exercises.pdf", sub.getvalue())

    return buf.getvalue()


def _safe_filename(name: str) -> str:
    import re
    name = re.sub(r"[\\/:*?\"<>|]", "_", name or "")
    return name.strip() or "lesson"


@router.get("/series/{series_id}/export-zip")
async def export_series_zip(
    series_id: str,
    format: str = "docx",
    include_exercises: bool = False,
    for_user_id: Optional[str] = Query(None, description="管理员：代导出指定用户的系列"),
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_active_user),
):
    """Export each lesson as its own file and package into a zip archive."""
    owner = await resolve_documents_owner(db, current_user, for_user_id)
    fmt = (format or "docx").lower()
    if fmt not in ALLOWED_SERIES_FORMATS:
        raise HTTPException(status_code=400, detail=f"不支持的格式：{fmt}")

    series, lessons = await _load_series_lessons(series_id, db, owner)

    lesson_ids = [l.id for l in lessons]
    exercises_map: dict[str, list[dict]] = {}
    if include_exercises:
        exercises_map = await _fetch_exercises_for_lessons(db, owner.id, lesson_ids)

    lessons_built: list[dict] = []
    for idx, lesson in enumerate(lessons):
        if not lesson.final_content:
            continue
        d = _build_doc_sections(lesson)
        prefix = _series_prefix(series, lesson, idx)
        base = f"{prefix}-{_safe_filename(d['title'])}"
        lessons_built.append({"idx": idx, "lesson": lesson, "data": d, "prefix": prefix, "base": base})

    if not lessons_built:
        raise HTTPException(status_code=400, detail="系列中没有任何已完成的教案")

    zip_bytes = await run_in_executor(
        _build_series_zip_bytes, fmt, series, lessons_built, include_exercises, exercises_map
    )

    zip_title = f"{series.title or 'series'}_{fmt}"
    await _record_export_safely(
        db, owner.id,
        lesson_plan_id=None,
        version_id=None,
        format="zip",
        file_name=f"{zip_title}.zip",
        file_size=len(zip_bytes),
        source_kind="bundle",
        params={"series_id": series_id, "format": fmt, "include_exercises": include_exercises, "type": "zip"},
    )

    return Response(
        content=zip_bytes,
        media_type="application/zip",
        headers={"Content-Disposition": _content_disposition(zip_title, "zip")},
    )


# ───────────────────────────────────────────────────────────────────
# 异步系列导出：QueueJob kind=export_bundle，结果落到 tmp_exports/，7 天 TTL
# ───────────────────────────────────────────────────────────────────

class _AsyncExportResponse(BaseModel):
    record_id: str
    status: str
    job_enqueued: bool
    expires_at: Optional[datetime] = None


def _bundle_ext_for(format_: str, bundle_type: str) -> str:
    if bundle_type == "zip":
        return "zip"
    return {"markdown": "md"}.get(format_, format_)


@router.post("/series/{series_id}/export-merged-async", response_model=_AsyncExportResponse)
async def export_series_merged_async(
    series_id: str,
    format: str = Query("docx"),
    include_exercises: bool = Query(False),
    for_user_id: Optional[str] = Query(None, description="管理员：代导出指定用户的系列"),
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_active_user),
):
    """异步合并导出：写入 ExportRecord(status=queued)，QueueJob 接力执行，7 天 TTL。"""
    owner = await resolve_documents_owner(db, current_user, for_user_id)
    fmt = (format or "docx").lower()
    if fmt not in ALLOWED_SERIES_FORMATS:
        raise HTTPException(status_code=400, detail=f"不支持的格式：{fmt}")

    series, _lessons = await _load_series_lessons(series_id, db, owner)
    series_title = series.title or "系列教案"
    file_name = f"{series_title}.{_bundle_ext_for(fmt, 'merged')}"

    record = ExportRecord(
        id=str(uuid.uuid4()),
        user_id=owner.id,
        lesson_plan_id=None,
        version_id=None,
        format=fmt,
        file_name=file_name,
        source_kind="bundle",
        status="queued",
        params={
            "series_id": series_id,
            "format": fmt,
            "include_exercises": include_exercises,
            "type": "merged",
        },
        expires_at=datetime.now(timezone.utc) + timedelta(hours=TMP_EXPORTS_TTL_HOURS),
    )
    db.add(record)
    await db.commit()
    await db.refresh(record)

    from app.tasks.queue_manager import enqueue
    ok = await enqueue(
        target_id=record.id,
        kind="export_bundle",
        user_id=owner.id,
        max_attempts=2,
    )
    if not ok:
        record.status = "failed"
        record.error_message = "任务入队失败或已有相同任务在队列中"
        await db.commit()
        await db.refresh(record)

    return _AsyncExportResponse(
        record_id=record.id,
        status=record.status,
        job_enqueued=bool(ok),
        expires_at=record.expires_at,
    )


@router.post("/series/{series_id}/export-zip-async", response_model=_AsyncExportResponse)
async def export_series_zip_async(
    series_id: str,
    format: str = Query("docx"),
    include_exercises: bool = Query(False),
    for_user_id: Optional[str] = Query(None, description="管理员：代导出指定用户的系列"),
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_active_user),
):
    """异步逐课打包导出（zip）。"""
    owner = await resolve_documents_owner(db, current_user, for_user_id)
    fmt = (format or "docx").lower()
    if fmt not in ALLOWED_SERIES_FORMATS:
        raise HTTPException(status_code=400, detail=f"不支持的格式：{fmt}")

    series, _lessons = await _load_series_lessons(series_id, db, owner)
    series_title = series.title or "系列教案"
    zip_name = f"{series_title}_{fmt}.zip"

    record = ExportRecord(
        id=str(uuid.uuid4()),
        user_id=owner.id,
        lesson_plan_id=None,
        version_id=None,
        format="zip",
        file_name=zip_name,
        source_kind="bundle",
        status="queued",
        params={
            "series_id": series_id,
            "format": fmt,
            "include_exercises": include_exercises,
            "type": "zip",
        },
        expires_at=datetime.now(timezone.utc) + timedelta(hours=TMP_EXPORTS_TTL_HOURS),
    )
    db.add(record)
    await db.commit()
    await db.refresh(record)

    from app.tasks.queue_manager import enqueue
    ok = await enqueue(
        target_id=record.id,
        kind="export_bundle",
        user_id=owner.id,
        max_attempts=2,
    )
    if not ok:
        record.status = "failed"
        record.error_message = "任务入队失败或已有相同任务在队列中"
        await db.commit()
        await db.refresh(record)

    return _AsyncExportResponse(
        record_id=record.id,
        status=record.status,
        job_enqueued=bool(ok),
        expires_at=record.expires_at,
    )


# ───────────────────────────────────────────────────────────────────
# Queue handler：执行 ExportRecord(status=queued) 的实际构建
# 由 app/tasks/job_handlers.py 在启动时注册到 kind=export_bundle
# ───────────────────────────────────────────────────────────────────

async def _bundle_emit_socket(user_id: str, record_id: str, status: str, **extra):
    sio = _get_sio()
    if not sio:
        return
    try:
        payload = {"record_id": record_id, "status": status, **extra}
        await sio.emit("export_record_update", payload, room=f"user_{user_id}")
    except Exception:
        pass


async def _bundle_mark_status(record_id: str, **fields):
    """更新 ExportRecord 状态字段，使用独立 session。"""
    async with async_session_maker() as session:
        res = await session.execute(select(ExportRecord).where(ExportRecord.id == record_id))
        rec = res.scalar_one_or_none()
        if not rec:
            return None
        for k, v in fields.items():
            setattr(rec, k, v)
        await session.commit()
        await session.refresh(rec)
        return rec


async def run_bundle_export_job(record_id: str):
    """`kind=export_bundle` handler：把 queued 的 ExportRecord 实际渲染并落盘。"""
    async with async_session_maker() as session:
        res = await session.execute(select(ExportRecord).where(ExportRecord.id == record_id))
        record = res.scalar_one_or_none()
        if not record:
            logger.warning(f"[bundle] record not found: {record_id}")
            return
        user_id = record.user_id
        params = record.params or {}
        series_id = params.get("series_id")
        fmt = params.get("format") or record.format
        include_exercises = bool(params.get("include_exercises"))
        bundle_type = params.get("type") or "merged"

    if not series_id or fmt not in ALLOWED_SERIES_FORMATS:
        await _bundle_mark_status(
            record_id, status="failed",
            error_message=f"invalid params: series_id={series_id} fmt={fmt}",
        )
        await _bundle_emit_socket(user_id, record_id, "failed", error="invalid params")
        return

    await _bundle_mark_status(record_id, status="running")
    await _bundle_emit_socket(user_id, record_id, "running")

    try:
        async with async_session_maker() as session:
            r = await session.execute(
                select(LessonSeries).where(LessonSeries.id == series_id, LessonSeries.user_id == user_id)
            )
            series = r.scalar_one_or_none()
            if not series:
                raise RuntimeError(f"series not found: {series_id}")

            r2 = await session.execute(
                select(LessonPlan)
                .where(LessonPlan.sequence_id == series_id, LessonPlan.user_id == user_id)
                .order_by(LessonPlan.sequence_order)
            )
            lessons = list(r2.scalars().all())

            exercises_map: dict[str, list[dict]] = {}
            if include_exercises:
                exercises_map = await _fetch_exercises_for_lessons(
                    session, user_id, [l.id for l in lessons]
                )

        built: list[dict] = []
        for idx, lesson in enumerate(lessons):
            if not lesson.final_content:
                continue
            d = _build_doc_sections(lesson)
            prefix = _series_prefix(series, lesson, idx)
            base = f"{prefix}-{_safe_filename(d['title'])}"
            built.append({"idx": idx, "lesson": lesson, "data": d, "prefix": prefix, "base": base})

        if not built:
            raise RuntimeError("系列中没有任何已完成的教案")

        series_title = series.title or "系列教案"

        if bundle_type == "zip":
            data_bytes = await run_in_executor(
                _build_series_zip_bytes, fmt, series, built, include_exercises, exercises_map
            )
            ext = "zip"
        elif fmt == "json":
            payload = {
                "series": {
                    "id": series.id, "title": series.title, "subject": series.subject,
                    "grade_level": series.grade_level,
                },
                "lessons": [
                    {"prefix": b["prefix"], "title": b["lesson"].title,
                     "final_content": _parse_fc(b["lesson"].final_content),
                     **({"exercises": exercises_map.get(b["lesson"].id, [])} if include_exercises else {})}
                    for b in built
                ],
            }
            data_bytes = json.dumps(payload, ensure_ascii=False, indent=2, default=_json_default).encode("utf-8")
            ext = "json"
        elif fmt == "txt":
            parts = [series_title, "=" * 40, ""]
            for b in built:
                parts.append(f"====== {b['prefix']}  {b['data']['title']} ======")
                parts.append(_lesson_plain_text(b["data"]))
                parts.append("")
            data_bytes = "\n".join(parts).encode("utf-8")
            ext = "txt"
        elif fmt == "md":
            parts = [f"# {series_title}", ""]
            for b in built:
                parts.append(f"## {b['prefix']}  {b['data']['title']}")
                parts.append(_lesson_markdown(b["data"], level_offset=2))
                parts.append("")
            data_bytes = "\n".join(parts).encode("utf-8")
            ext = "md"
        elif fmt == "docx":
            data_bytes = await run_in_executor(
                _build_series_merged_docx, series_title, built, series, include_exercises, exercises_map
            )
            ext = "docx"
        elif fmt == "pdf":
            data_bytes = await run_in_executor(
                _build_series_merged_pdf, series_title, built, series, include_exercises, exercises_map
            )
            ext = "pdf"
        else:
            raise RuntimeError(f"未知格式：{fmt}")

        file_path = os.path.join(TMP_EXPORTS_DIR, f"{record_id}.{ext}")
        with open(file_path, "wb") as fp:
            fp.write(data_bytes)

        await _bundle_mark_status(
            record_id,
            status="done",
            file_path=file_path,
            file_size=len(data_bytes),
        )
        await _bundle_emit_socket(
            user_id, record_id, "done",
            file_size=len(data_bytes),
            download_url=f"/api/v1/documents/exports/{record_id}/download",
        )
        logger.info(f"[bundle] done record={record_id} ext={ext} size={len(data_bytes)}")
    except Exception as e:
        logger.exception(f"[bundle] failed record={record_id}: {e}")
        await _bundle_mark_status(
            record_id, status="failed",
            error_message=f"{type(e).__name__}: {e}",
        )
        await _bundle_emit_socket(user_id, record_id, "failed", error=str(e))
        raise


async def cleanup_expired_exports() -> int:
    """删除 tmp_exports/ 下过期文件，把对应 ExportRecord.status 标为 expired。"""
    cleaned = 0
    now = datetime.now(timezone.utc)
    async with async_session_maker() as session:
        res = await session.execute(
            select(ExportRecord)
            .where(
                ExportRecord.expires_at.isnot(None),
                ExportRecord.expires_at < now,
                ExportRecord.status.in_(["done", "running", "queued"]),
            )
        )
        rows = list(res.scalars().all())
        for r in rows:
            if r.file_path and os.path.exists(r.file_path):
                try:
                    os.remove(r.file_path)
                except Exception as e:
                    logger.warning(f"[bundle-gc] unlink failed {r.file_path}: {e}")
            r.status = "expired"
            r.file_path = None
            cleaned += 1
        await session.commit()
    if cleaned:
        logger.info(f"[bundle-gc] expired {cleaned} export records")
    return cleaned
