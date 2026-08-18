"""Fill syllabus / calendar / lesson templates from DeepSeek JSON."""
from __future__ import annotations

import os
import shutil
from typing import Any, Dict, List, Optional

from docx import Document
from openpyxl import load_workbook
from openpyxl.worksheet.worksheet import Worksheet

_TEMPLATES_DIR = os.path.join(
    os.path.dirname(os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))),
    "templates",
    "zhuke_materials",
)


def templates_dir() -> str:
    return _TEMPLATES_DIR


def _tpl(name: str) -> str:
    return os.path.join(_TEMPLATES_DIR, name)


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(str(text or ""), level=level)


def _add_para(doc: Document, label: str, value: Any) -> None:
    p = doc.add_paragraph()
    run = p.add_run(f"{label}：")
    run.bold = True
    p.add_run("" if value is None else str(value))


def fill_syllabus_docx(data: Dict[str, Any], out_path: str, *, use_lab: bool = False) -> str:
    """Write a structured syllabus docx from JSON.

    Official template is kept as a companion copy; field mapping across complex
    附件3/4 tables is lossy — structured output guarantees usable delivery.
    """
    os.makedirs(os.path.dirname(out_path) or ".", exist_ok=True)
    src = _tpl("syllabus_lab.docx" if use_lab else "syllabus_theory.docx")
    companion = out_path.replace(".docx", "_template_blank.docx")
    if os.path.isfile(src):
        shutil.copy2(src, companion)

    doc = Document()
    doc.add_heading(f"珠海科技学院教学大纲 — {data.get('course_name') or ''}", 0)
    for key, label in (
        ("course_name", "课程名称"),
        ("course_code", "课程代码"),
        ("credits", "学分"),
        ("total_hours", "总学时"),
        ("theory_hours", "理论学时"),
        ("lab_hours", "实验学时"),
        ("course_nature", "课程性质"),
        ("applicable_major", "适用专业"),
        ("prerequisites", "先修课程"),
        ("course_intro", "课程简介"),
        ("teaching_methods", "教学方法"),
    ):
        _add_para(doc, label, data.get(key))

    objs = data.get("course_objectives") or []
    if isinstance(objs, list):
        _add_heading(doc, "课程目标", 1)
        for i, o in enumerate(objs, 1):
            doc.add_paragraph(f"{i}. {o}")

    assessment = data.get("assessment") or {}
    if isinstance(assessment, dict):
        _add_heading(doc, "考核方式", 1)
        _add_para(doc, "平时", assessment.get("usual_percent"))
        _add_para(doc, "实验", assessment.get("lab_percent"))
        _add_para(doc, "期末", assessment.get("final_percent"))
        _add_para(doc, "说明", assessment.get("notes"))

    chapters = data.get("chapters") or []
    if isinstance(chapters, list):
        _add_heading(doc, "教学内容与学时分配", 1)
        for ch in chapters:
            if not isinstance(ch, dict):
                continue
            _add_heading(doc, f"第{ch.get('no', '')}章 {ch.get('title', '')}", 2)
            _add_para(doc, "学时", ch.get("hours"))
            _add_para(doc, "类型", ch.get("theory_or_lab"))
            _add_para(doc, "目标", ch.get("objectives"))
            _add_para(doc, "内容", ch.get("content"))
            _add_para(doc, "重点", ch.get("key_points"))
            _add_para(doc, "难点", ch.get("difficulties"))
            _add_para(doc, "方法", ch.get("methods_note"))

    for key, label in (("textbooks", "教材"), ("references", "参考书"), ("other_notes", "其他说明")):
        val = data.get(key)
        _add_heading(doc, label, 1)
        if isinstance(val, list):
            for item in val:
                doc.add_paragraph(str(item))
        else:
            doc.add_paragraph("" if val is None else str(val))

    doc.save(out_path)
    return out_path


def _clear_week_rows(ws: Worksheet, start_row: int = 7) -> None:
    # Keep footer rows that start with 注/本课程 — find first footer-like row
    max_r = ws.max_row or start_row
    footer_start = None
    for r in range(start_row, max_r + 1):
        v = ws.cell(r, 1).value
        if isinstance(v, str) and (v.startswith("注") or v.startswith("本课程") or "院系" in v):
            footer_start = r
            break
    end = (footer_start - 1) if footer_start else max_r
    for r in range(start_row, end + 1):
        for c in range(1, 12):
            ws.cell(r, c).value = None


def _write_weeks(
    ws: Worksheet,
    weeks: List[Dict[str, Any]],
    *,
    meta: Dict[str, Any],
    schedule: str,
    teacher: str = "",
) -> None:
    # Meta band (best-effort on sample layout)
    if meta.get("course_code"):
        ws.cell(3, 3).value = meta.get("course_code")
    if meta.get("course_name"):
        ws.cell(3, 7).value = meta.get("course_name")
    if meta.get("class_name"):
        ws.cell(3, 11).value = meta.get("class_name")
    # weekday text into 星期 column sample uses 一/二…
    weekday = ""
    if schedule:
        # e.g. 周三 第3–4节 → 三
        for ch in ("一", "二", "三", "四", "五", "六", "日"):
            if f"周{ch}" in schedule or f"星期{ch}" in schedule:
                weekday = ch
                break

    _clear_week_rows(ws, 7)
    for i, w in enumerate(weeks):
        r = 7 + i
        ws.cell(r, 1).value = w.get("week") or (i + 1)
        ws.cell(r, 4).value = weekday or ""
        if teacher:
            ws.cell(r, 5).value = teacher
        ws.cell(r, 8).value = w.get("hours") or 2
        content = w.get("teaching_content") or w.get("experiment_name") or ""
        ws.cell(r, 9).value = content
        # image columns left empty (J/K…)


def fill_calendar_xlsx(
    weeks_data: Dict[str, Any],
    out_theory: str,
    out_lab: Optional[str] = None,
    *,
    schedule: str = "",
    teacher: str = "",
) -> Dict[str, str]:
    os.makedirs(os.path.dirname(out_theory) or ".", exist_ok=True)
    meta = weeks_data.get("meta") or {}
    if not isinstance(meta, dict):
        meta = {}
    if schedule and not meta.get("schedule"):
        meta = {**meta, "schedule": schedule}

    theory_src = _tpl("calendar_theory.xlsx")
    shutil.copy2(theory_src, out_theory)
    wb = load_workbook(out_theory)
    ws = wb.active
    theory_weeks = weeks_data.get("theory_weeks") or []
    if isinstance(theory_weeks, list):
        _write_weeks(ws, theory_weeks, meta=meta, schedule=schedule or meta.get("schedule") or "", teacher=teacher)
    wb.save(out_theory)

    paths = {"theory": out_theory}
    lab_weeks = weeks_data.get("lab_weeks") or []
    if isinstance(lab_weeks, list) and lab_weeks and out_lab:
        lab_src = _tpl("calendar_lab.xlsx")
        if os.path.isfile(lab_src):
            shutil.copy2(lab_src, out_lab)
            wb2 = load_workbook(out_lab)
            _write_weeks(
                wb2.active,
                lab_weeks,
                meta=meta,
                schedule=schedule or meta.get("schedule") or "",
                teacher=teacher,
            )
            wb2.save(out_lab)
            paths["lab"] = out_lab
    return paths


def fill_lessons_docx(lessons_data: Dict[str, Any], out_path: str) -> str:
    os.makedirs(os.path.dirname(out_path) or ".", exist_ok=True)
    # Keep official template blank companion
    src = _tpl("lesson_plan.docx")
    if os.path.isfile(src):
        shutil.copy2(src, out_path.replace(".docx", "_template_blank.docx"))

    doc = Document()
    doc.add_heading("珠海科技学院教学设计（教案）", 0)
    lessons = lessons_data.get("lessons") or []
    if not isinstance(lessons, list):
        lessons = []
    for les in lessons:
        if not isinstance(les, dict):
            continue
        title = les.get("title") or f"第{les.get('unit_index', '')}次课"
        _add_heading(doc, f"课次 {les.get('unit_index', '')} · 周{les.get('week', '')} · {title}", 1)
        for key, label in (
            ("class_hours", "学时"),
            ("schedule_text", "上课时间"),
            ("learning_situation", "学情分析"),
            ("objectives", "教学目标"),
            ("key_points", "教学重点"),
            ("difficulties", "教学难点"),
            ("methods_and_means", "教学方法与手段"),
            ("homework", "作业"),
            ("reflection", "教学反思"),
            ("materials", "教学材料"),
        ):
            _add_para(doc, label, les.get(key))
        process = les.get("process") or []
        if isinstance(process, list) and process:
            _add_heading(doc, "教学过程", 2)
            table = doc.add_table(rows=1, cols=5)
            table.style = "Table Grid"
            hdr = table.rows[0].cells
            hdr[0].text = "环节"
            hdr[1].text = "分钟"
            hdr[2].text = "教师活动"
            hdr[3].text = "学生活动"
            hdr[4].text = "设计意图"
            for step in process:
                if not isinstance(step, dict):
                    continue
                row = table.add_row().cells
                row[0].text = str(step.get("phase") or "")
                row[1].text = str(step.get("minutes") or "")
                row[2].text = str(step.get("teacher_activity") or "")
                row[3].text = str(step.get("student_activity") or "")
                row[4].text = str(step.get("intent") or "")
        doc.add_page_break()

    # remove trailing page break if any empty
    doc.save(out_path)
    return out_path
