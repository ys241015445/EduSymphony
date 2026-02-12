"""
导出服务
支持Word、PDF、TXT、JSON格式导出
"""
import json
from typing import Dict, Optional
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from weasyprint import HTML, CSS
from minio import Minio
from loguru import logger

from app.core.config import settings

class ExportService:
    """导出服务类"""
    
    def __init__(self):
        # 初始化MinIO客户端
        try:
            self.minio_client = Minio(
                settings.MINIO_ENDPOINT,
                access_key=settings.MINIO_ACCESS_KEY,
                secret_key=settings.MINIO_SECRET_KEY,
                secure=settings.MINIO_SECURE
            )
            
            # 确保bucket存在
            if not self.minio_client.bucket_exists(settings.MINIO_BUCKET):
                self.minio_client.make_bucket(settings.MINIO_BUCKET)
            
            logger.info("✅ MinIO客户端初始化成功")
        except Exception as e:
            logger.error(f"❌ MinIO客户端初始化失败: {str(e)}")
            self.minio_client = None
    
    async def export_word(self, lesson_data: Dict) -> str:
        """
        导出Word文档
        
        Args:
            lesson_data: 教案数据
        
        Returns:
            文件URL
        """
        try:
            doc = Document()
            
            # 添加标题
            title = doc.add_heading(lesson_data['title'], 0)
            title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            # 添加基本信息
            info_table = doc.add_table(rows=4, cols=2)
            info_table.style = 'Light Grid Accent 1'
            
            info_data = [
                ('学科', lesson_data['subject']),
                ('年级', lesson_data['grade_level']),
                ('地区', lesson_data['region']),
                ('教学模型', lesson_data.get('teaching_model', ''))
            ]
            
            for i, (label, value) in enumerate(info_data):
                info_table.rows[i].cells[0].text = label
                info_table.rows[i].cells[1].text = value
            
            doc.add_paragraph()  # 空行
            
            # 添加各阶段内容
            final_content = lesson_data.get('final_content', {})
            stages = final_content.get('stages', {})
            
            for stage_id, stage_data in stages.items():
                # 阶段标题
                stage_heading = doc.add_heading(stage_data['name'], level=1)
                
                # 专家信息
                expert_para = doc.add_paragraph()
                expert_para.add_run(f"设计专家：{stage_data.get('expert', '未知')}").italic = True
                
                # 内容
                content_para = doc.add_paragraph(stage_data['content'])
                content_para.paragraph_format.line_spacing = 1.5
                
                doc.add_paragraph()  # 空行
            
            # 保存到临时文件
            temp_dir = Path("/tmp/exports")
            temp_dir.mkdir(exist_ok=True)
            
            filename = f"lesson_{lesson_data['id']}_{datetime.now().strftime('%Y%m%d%H%M%S')}.docx"
            file_path = temp_dir / filename
            
            doc.save(str(file_path))
            
            # 上传到MinIO
            url = await self._upload_to_minio(file_path, filename)
            
            # 清理临时文件
            file_path.unlink(missing_ok=True)
            
            return url
            
        except Exception as e:
            logger.error(f"Word导出失败: {str(e)}")
            raise Exception(f"Word导出失败: {str(e)}")
    
    async def export_pdf(self, lesson_data: Dict) -> str:
        """
        导出PDF文档
        
        Args:
            lesson_data: 教案数据
        
        Returns:
            文件URL
        """
        try:
            # 生成HTML
            html_content = self._generate_html(lesson_data)
            
            # CSS样式
            css_content = """
            @page {
                size: A4;
                margin: 2cm;
            }
            body {
                font-family: "SimSun", serif;
                font-size: 12pt;
                line-height: 1.6;
            }
            h1 {
                color: #2c3e50;
                text-align: center;
                font-size: 24pt;
                margin-bottom: 20pt;
            }
            h2 {
                color: #34495e;
                font-size: 18pt;
                margin-top: 15pt;
                margin-bottom: 10pt;
                border-bottom: 2px solid #3498db;
            }
            table {
                width: 100%;
                border-collapse: collapse;
                margin: 10pt 0;
            }
            table td {
                padding: 8pt;
                border: 1px solid #bdc3c7;
            }
            .expert-info {
                color: #7f8c8d;
                font-style: italic;
                margin-bottom: 10pt;
            }
            .content {
                text-align: justify;
                margin-bottom: 15pt;
            }
            """
            
            # 生成PDF
            temp_dir = Path("/tmp/exports")
            temp_dir.mkdir(exist_ok=True)
            
            filename = f"lesson_{lesson_data['id']}_{datetime.now().strftime('%Y%m%d%H%M%S')}.pdf"
            file_path = temp_dir / filename
            
            HTML(string=html_content).write_pdf(
                str(file_path),
                stylesheets=[CSS(string=css_content)]
            )
            
            # 上传到MinIO
            url = await self._upload_to_minio(file_path, filename)
            
            # 清理临时文件
            file_path.unlink(missing_ok=True)
            
            return url
            
        except Exception as e:
            logger.error(f"PDF导出失败: {str(e)}")
            raise Exception(f"PDF导出失败: {str(e)}")
    
    async def export_txt(self, lesson_data: Dict, clean: bool = True) -> str:
        """
        导出TXT文本
        
        Args:
            lesson_data: 教案数据
            clean: 是否为纯净版（无标注）
        
        Returns:
            文件URL
        """
        try:
            content_lines = []
            
            # 标题
            content_lines.append("=" * 60)
            content_lines.append(lesson_data['title'].center(60))
            content_lines.append("=" * 60)
            content_lines.append("")
            
            # 基本信息
            content_lines.append(f"学科：{lesson_data['subject']}")
            content_lines.append(f"年级：{lesson_data['grade_level']}")
            content_lines.append(f"地区：{lesson_data['region']}")
            content_lines.append(f"教学模型：{lesson_data.get('teaching_model', '')}")
            content_lines.append("")
            content_lines.append("-" * 60)
            content_lines.append("")
            
            # 各阶段内容
            final_content = lesson_data.get('final_content', {})
            stages = final_content.get('stages', {})
            
            for stage_id, stage_data in stages.items():
                content_lines.append(f"\n【{stage_data['name']}】")
                
                if not clean:
                    content_lines.append(f"设计专家：{stage_data.get('expert', '未知')}")
                
                content_lines.append("")
                content_lines.append(stage_data['content'])
                content_lines.append("")
                content_lines.append("-" * 60)
            
            # 合并所有行
            full_content = "\n".join(content_lines)
            
            # 保存到临时文件
            temp_dir = Path("/tmp/exports")
            temp_dir.mkdir(exist_ok=True)
            
            suffix = "clean" if clean else "annotated"
            filename = f"lesson_{lesson_data['id']}_{suffix}_{datetime.now().strftime('%Y%m%d%H%M%S')}.txt"
            file_path = temp_dir / filename
            
            file_path.write_text(full_content, encoding='utf-8')
            
            # 上传到MinIO
            url = await self._upload_to_minio(file_path, filename)
            
            # 清理临时文件
            file_path.unlink(missing_ok=True)
            
            return url
            
        except Exception as e:
            logger.error(f"TXT导出失败: {str(e)}")
            raise Exception(f"TXT导出失败: {str(e)}")
    
    async def export_json(self, lesson_data: Dict) -> str:
        """
        导出JSON格式
        
        Args:
            lesson_data: 教案数据
        
        Returns:
            文件URL
        """
        try:
            # 保存到临时文件
            temp_dir = Path("/tmp/exports")
            temp_dir.mkdir(exist_ok=True)
            
            filename = f"lesson_{lesson_data['id']}_{datetime.now().strftime('%Y%m%d%H%M%S')}.json"
            file_path = temp_dir / filename
            
            file_path.write_text(
                json.dumps(lesson_data, ensure_ascii=False, indent=2),
                encoding='utf-8'
            )
            
            # 上传到MinIO
            url = await self._upload_to_minio(file_path, filename)
            
            # 清理临时文件
            file_path.unlink(missing_ok=True)
            
            return url
            
        except Exception as e:
            logger.error(f"JSON导出失败: {str(e)}")
            raise Exception(f"JSON导出失败: {str(e)}")
    
    def _generate_html(self, lesson_data: Dict) -> str:
        """生成HTML内容"""
        html = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <title>{lesson_data['title']}</title>
        </head>
        <body>
            <h1>{lesson_data['title']}</h1>
            
            <table>
                <tr>
                    <td width="20%">学科</td>
                    <td>{lesson_data['subject']}</td>
                </tr>
                <tr>
                    <td>年级</td>
                    <td>{lesson_data['grade_level']}</td>
                </tr>
                <tr>
                    <td>地区</td>
                    <td>{lesson_data['region']}</td>
                </tr>
                <tr>
                    <td>教学模型</td>
                    <td>{lesson_data.get('teaching_model', '')}</td>
                </tr>
            </table>
        """
        
        final_content = lesson_data.get('final_content', {})
        stages = final_content.get('stages', {})
        
        for stage_id, stage_data in stages.items():
            content_html = stage_data['content'].replace('\n', '<br>')
            html += f"""
            <h2>{stage_data['name']}</h2>
            <p class="expert-info">设计专家：{stage_data.get('expert', '未知')}</p>
            <div class="content">{content_html}</div>
            """
        
        html += """
        </body>
        </html>
        """
        
        return html
    
    async def _upload_to_minio(self, file_path: Path, filename: str) -> str:
        """上传文件到MinIO"""
        if not self.minio_client:
            raise Exception("MinIO客户端未初始化")
        
        try:
            object_name = f"exports/{filename}"
            
            self.minio_client.fput_object(
                settings.MINIO_BUCKET,
                object_name,
                str(file_path)
            )
            
            # 生成访问URL（7天有效期）
            url = self.minio_client.presigned_get_object(
                settings.MINIO_BUCKET,
                object_name,
                expires=7*24*3600  # 7天
            )
            
            logger.info(f"✅ 文件上传成功: {object_name}")
            return url
            
        except Exception as e:
            logger.error(f"❌ 文件上传失败: {str(e)}")
            raise Exception(f"文件上传失败: {str(e)}")

