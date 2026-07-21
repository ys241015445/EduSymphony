"""珠科教案助手 — service layer.

Parses an uploaded 教学日历 (xlsx/docx/pdf), drives Kimi K2 to produce per-lesson
content blocks, and assembles a docx using `templates/zhuke_lesson_template.docx`
as the visual baseline. Designed to be called from
`backend/app/api/semester_helper.py`.

Public surface:
- :func:`parse_schedule` — best-effort header detection + row extraction.
- :class:`KimiAgent` — thin Moonshot wrapper that returns JSON per lesson.
- :func:`build_docx` — clones the template's first lesson block N times.

Parsing design (2026-05 rewrite):
- All three formats (xlsx / docx / pdf) are reduced to a uniform list of
  ``logical_tables`` where each table is a ``List[List[str]]`` of logical cells
  (merged cells collapsed to a single value).
- We then SCAN every row across every table and pick the row with the largest
  number of header-alias hits as the canonical header row.
- The metadata band ABOVE the header row is mined for ``LABEL：VALUE`` pairs
  (works whether the label and value sit in the same cell or in two adjacent
  cells, which is how 珠科课程教学进度表 actually formats its cover band).
- Body rows below the header AND any other tables whose logical column count
  matches the header's are mapped through the same ``field_map`` — this is what
  recovers the second-page continuation table that has no header of its own.
"""
from __future__ import annotations

import io
import json
import os
import re
from copy import deepcopy
from dataclasses import dataclass, field
from typing import Any, Dict, List, Optional, Tuple

from loguru import logger

from app.core.config import settings


# ─────────────────────────────────────────────────────────────────
# Template location
# ─────────────────────────────────────────────────────────────────

# `backend/templates/zhuke_lesson_template.docx` lives alongside `backend/app`.
_TEMPLATE_PATH = os.path.join(
    os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))),
    "templates",
    "zhuke_lesson_template.docx",
)


def template_path() -> str:
    return _TEMPLATE_PATH


def template_exists() -> bool:
    return os.path.isfile(_TEMPLATE_PATH)


# ─────────────────────────────────────────────────────────────────
# docx → pdf via LibreOffice (headless soffice)
# ─────────────────────────────────────────────────────────────────


def _find_soffice() -> Optional[str]:
    """Locate LibreOffice's headless converter binary across the common
    Windows / macOS / Linux install locations. Returns None when the user
    hasn't installed LibreOffice — the caller should fall back to the lossy
    text-only PDF renderer (or surface a friendly error)."""
    import shutil
    env_path = os.environ.get("SOFFICE_PATH", "").strip()
    if env_path and os.path.isfile(env_path):
        return env_path
    cand = shutil.which("soffice") or shutil.which("soffice.exe")
    if cand:
        return cand
    for p in (
        r"C:\Program Files\LibreOffice\program\soffice.exe",
        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
        "/Applications/LibreOffice.app/Contents/MacOS/soffice",
        "/usr/bin/soffice",
        "/usr/lib/libreoffice/program/soffice",
    ):
        if os.path.isfile(p):
            return p
    return None


def convert_docx_to_pdf_via_soffice(docx_bytes: bytes) -> Optional[bytes]:
    """Convert a docx blob to PDF using LibreOffice headless mode.

    Returns the PDF bytes on success, None on any failure (soffice not
    installed, conversion crashed, output missing, timed out). The caller
    decides whether to fall back to the text-only renderer or 502 the user.

    Why LibreOffice: the old `_render_pdf_from_text` extracted plain text
    then re-laid via reportlab — that destroys all docx formatting (tables,
    fonts, page breaks, underlines, the entire 珠科 layout). LibreOffice
    renders the docx faithfully because it parses Word's own layout engine.
    """
    soffice = _find_soffice()
    if not soffice:
        return None
    import subprocess
    import tempfile
    with tempfile.TemporaryDirectory(prefix="zhuke_pdf_") as td:
        in_path = os.path.join(td, "in.docx")
        with open(in_path, "wb") as f:
            f.write(docx_bytes)
        try:
            # `--headless` avoids spawning the desktop UI; the per-user
            # profile dir keeps concurrent invocations from clobbering each
            # other's UNO listener. Timeout=180s covers 16-page payloads.
            proc = subprocess.run(
                [
                    soffice,
                    "--headless",
                    "-env:UserInstallation=file:///" + os.path.join(td, "lo_profile").replace("\\", "/"),
                    "--convert-to",
                    "pdf",
                    "--outdir",
                    td,
                    in_path,
                ],
                capture_output=True,
                timeout=180,
            )
        except (subprocess.TimeoutExpired, OSError) as e:
            from loguru import logger as _logger
            _logger.warning(f"[zhuke] LibreOffice convert failed: {e}")
            return None
        if proc.returncode != 0:
            from loguru import logger as _logger
            _logger.warning(
                f"[zhuke] LibreOffice exited rc={proc.returncode} "
                f"stderr={proc.stderr[:300].decode(errors='replace')}"
            )
            return None
        out_path = os.path.join(td, "in.pdf")
        if not os.path.isfile(out_path):
            return None
        with open(out_path, "rb") as f:
            return f.read()


# ─────────────────────────────────────────────────────────────────
# Schedule parsing
# ─────────────────────────────────────────────────────────────────

# Maps canonical field → list of accepted header substrings (Chinese,
# case-insensitive, whitespace-normalised). The classifier matches by
# normalized-equality / strict-prefix to avoid metadata cells like
# `上课班级总人数` being misread as `class_name`.
_HEADER_ALIASES: Dict[str, List[str]] = {
    # Body columns that feed `field_map` for lesson rows.
    "lesson_no": ["课次", "序号", "节次序号", "课时序号"],
    "week": ["周次", "教学周", "第几周"],
    "weekday": ["星期", "周次星期", "上课星期", "节次星期"],
    "periods": ["节次", "节数", "上课节次", "时段"],
    "date": ["日期", "上课日期", "授课日期"],
    "hours": ["周学时", "课时", "学时", "课时数"],
    "title": ["授课题目", "项目", "模块", "章节"],
    "content": [
        "授课内容", "教学内容", "本周内容", "课程内容",
        "授课主题", "教学主题", "教学要点", "授课进度", "教学进度",
        "主要内容", "本节内容", "章节内容", "内容",
    ],
    # Cover columns / metadata-band labels.
    "college": ["开课学院", "学院", "院系", "开设学院", "所属学院"],
    "course_name": ["课程名称", "课程名", "课程"],
    "course_type": ["课程性质", "授课类别", "课程类别", "类型"],
    "teacher": ["教师姓名", "主讲教师", "授课教师", "教师", "讲师"],
    "class_name": ["上课班级(含合班情况)", "上课班级", "教学班级", "教学班", "学生班级", "班级"],
}

# Cover-only labels mined from the metadata band (NEVER fed into the column
# field_map for lesson rows — they live in returned `cover` only).
_COVER_ONLY_LABEL_ALIASES: Dict[str, List[str]] = {
    "course_code": ["课程代码", "课程编号"],
    "total_weeks": ["周数", "总周数"],
    "total_hours": ["总学时"],
    "theory_hours": ["理论学时"],
    "practice_hours": ["实验学时", "实践学时"],
    "exam_type": ["考核方式"],
    "student_count": ["上课班级总人数", "学生人数", "总人数", "班级总人数"],
}

# Used by `_classify_header` (header-row scoring) — only the BODY column aliases
# count, so we do not let "课程名称" in a header-band row hijack scoring for a
# real body row that has 课程内容/授课内容. Cover-only labels are mined separately.
_BODY_ALIASES: Dict[str, List[str]] = {
    k: v for k, v in _HEADER_ALIASES.items()
    if k in ("lesson_no", "week", "weekday", "periods", "date", "hours", "title", "content", "teacher")
}


def _norm(s: Any) -> str:
    if s is None:
        return ""
    return re.sub(r"\s+", "", str(s)).lower()


def _strip_label_punct(s: str) -> str:
    """Strip ONLY trailing colon-style punctuation + whitespace used in label
    cells (e.g. `课程性质：` → `课程性质`). We deliberately do NOT strip
    parentheses because aliases like `上课班级(含合班情况)` legitimately end in
    `)` and the closing paren must survive the strip to match."""
    return re.sub(r"[:：、\.。,，\s]+$", "", s)


def _classify_header(cell: Any) -> Optional[str]:
    """Return canonical BODY field key if the cell looks like a header column.

    Header cells are short (~≤12 chars) and equal to (or strictly prefix-equal
    to) one of the body aliases after normalization. We deliberately reject
    long metadata cells like `课程性质：选修课` to keep header-row scoring clean.
    """
    if cell is None:
        return None
    raw = str(cell).strip()
    if not raw:
        return None
    n = _norm(_strip_label_punct(raw))
    if not n:
        return None
    # Short cells only — real table headers are rarely longer than ~12 chars.
    short = len(n) <= 12
    for canonical, aliases in _BODY_ALIASES.items():
        for alias in aliases:
            a = _norm(alias)
            if not a:
                continue
            if n == a:
                return canonical
            if short and a in n and len(n) - len(a) <= 4:
                # e.g. "周学时数" → matches hours alias "周学时"; tight tolerance
                return canonical
    return None


def _classify_cover_label(cell: Any) -> Optional[str]:
    """Return canonical cover key for a `LABEL：` cell (or LABEL portion of a
    `LABEL：VALUE` cell). Uses STRICT equality / prefix-equality after stripping
    trailing punctuation — so `上课班级总人数` will NOT match `班级`.
    """
    if cell is None:
        return None
    raw = str(cell).strip()
    if not raw:
        return None
    label = _strip_label_punct(raw)
    n = _norm(label)
    if not n:
        return None
    # Try cover-only labels first (more specific).
    for canonical, aliases in _COVER_ONLY_LABEL_ALIASES.items():
        for alias in aliases:
            a = _norm(alias)
            if a and n == a:
                return canonical
    # Then HEADER_ALIASES (so `课程名称` / `开课学院` / `教师姓名` etc. catch).
    for canonical, aliases in _HEADER_ALIASES.items():
        if canonical in ("lesson_no", "week", "weekday", "periods", "date", "hours", "title", "content"):
            # These are body columns; don't promote them as cover labels.
            continue
        for alias in aliases:
            a = _norm(alias)
            if a and n == a:
                return canonical
    return None


# ── Row collapse helpers ─────────────────────────────────────────


def _collapse_docx_row(row) -> List[str]:
    """Dedupe a python-docx row by `id(cell._element)` so a horizontally/vertically
    merged cell appears only once. Returns the cleaned cell text list."""
    out: List[str] = []
    seen_ids = set()
    for c in row.cells:
        cid = id(c._element)
        if cid in seen_ids:
            continue
        seen_ids.add(cid)
        out.append((c.text or "").replace("\n", " ").strip())
    return out


def _collapse_pdf_row(row: List[Optional[str]]) -> List[str]:
    """Collapse runs of empties + adjacent duplicates that pdfplumber emits for
    horizontally merged cells.

    Examples:
        ['a', '', '', 'b'] -> ['a', 'b']         # empty = merge with previous
        ['a', 'a', 'a', 'b'] -> ['a', 'b']       # repeats coalesce
        ['', '', 'a'] -> ['a']                   # leading empties dropped
    """
    cleaned: List[str] = []
    for raw in row:
        v = (raw or "").replace("\n", " ").strip()
        if not v:
            # Empty cells continue the previous logical cell (merged); skip.
            if cleaned:
                continue
            # Leading empty -> skip entirely (no prior logical cell).
            continue
        if cleaned and cleaned[-1] == v:
            # Adjacent duplicate from a merged cell — already counted.
            continue
        cleaned.append(v)
    return cleaned


def _collapse_xlsx_row(row: List[str]) -> List[str]:
    """xlsx via openpyxl doesn't duplicate merged values per-cell — empty cells
    are real empties. We keep all cells verbatim so column indices line up."""
    return [(v or "").replace("\n", " ").strip() if isinstance(v, str) else ("" if v is None else str(v).strip())
            for v in row]


# ── Header detection across all tables ───────────────────────────


@dataclass
class _HeaderHit:
    table_idx: int
    row_idx: int
    field_map: Dict[int, str]  # logical col idx -> canonical body field key
    width: int                 # logical column count of that row's table

    @property
    def score(self) -> int:
        return len(self.field_map)


def _pick_header_across_tables(logical_tables: List[List[List[str]]]) -> Optional[_HeaderHit]:
    """Scan every row of every table and return the highest-scoring header hit.

    A "score" = number of distinct canonical body fields the row matches.
    Ties broken by larger logical width (more columns = better).
    """
    best: Optional[_HeaderHit] = None
    for ti, tbl in enumerate(logical_tables):
        for ri, row in enumerate(tbl):
            fm: Dict[int, str] = {}
            for ci, cell in enumerate(row):
                key = _classify_header(cell)
                if key and key not in fm.values():
                    fm[ci] = key
            if not fm:
                continue
            cand = _HeaderHit(ti, ri, fm, width=len(row))
            if best is None or cand.score > best.score or (
                cand.score == best.score and cand.width > best.width
            ):
                best = cand
    return best


# ── Cover extraction from the metadata band ──────────────────────


def _extract_cover_from_band(
    logical_tables: List[List[List[str]]],
    header: _HeaderHit,
) -> Dict[str, str]:
    """Mine all rows BEFORE `header.row_idx` in `header.table_idx` for
    LABEL→VALUE pairs in either same-cell ("LABEL：VALUE") or split-cell
    ("LABEL：" + next non-empty cell) layouts."""
    cover: Dict[str, str] = {}
    if header.table_idx >= len(logical_tables):
        return cover
    band = logical_tables[header.table_idx][: header.row_idx]

    for row in band:
        cells = list(row)
        i = 0
        while i < len(cells):
            cell = cells[i].strip()
            if not cell:
                i += 1
                continue
            # Case A: LABEL：VALUE in one cell.
            m = re.match(r"^(.{1,20}?)[:：]\s*(.+)$", cell)
            if m:
                label_raw, value_raw = m.group(1), m.group(2).strip()
                key = _classify_cover_label(label_raw)
                if key and value_raw:
                    cover.setdefault(key, value_raw)
                    i += 1
                    continue
            # Case B: LABEL：(or LABEL: at end of cell, value in next cell).
            if cell.endswith(":") or cell.endswith("："):
                key = _classify_cover_label(cell)
                # advance to next non-empty cell
                j = i + 1
                while j < len(cells) and not cells[j].strip():
                    j += 1
                if key and j < len(cells):
                    cover.setdefault(key, cells[j].strip())
                    i = j + 1
                    continue
            # Case C: cell is exactly a known label without colon (rare).
            key = _classify_cover_label(cell)
            if key:
                j = i + 1
                while j < len(cells) and not cells[j].strip():
                    j += 1
                if j < len(cells):
                    cover.setdefault(key, cells[j].strip())
                    i = j + 1
                    continue
            i += 1
    return cover


# ── Lesson rows extraction (header + matching continuation tables) ──


def _extract_lessons(
    logical_tables: List[List[List[str]]],
    header: _HeaderHit,
) -> List[Dict[str, str]]:
    """Apply `header.field_map` to:
      (1) the body rows below `header.row_idx` of the header table, AND
      (2) every other table whose logical column count == header.width
          (these are no-header continuation pages typical of 珠科 教学日历).
    """
    out: List[Dict[str, str]] = []

    def _emit(rows: List[List[str]]) -> None:
        for row in rows:
            item: Dict[str, str] = {}
            for col_idx, field_key in header.field_map.items():
                if col_idx < len(row):
                    v = row[col_idx].strip()
                    if v:
                        item[field_key] = v
            if not item:
                continue
            # Require at least content (授课内容) OR title (授课题目) OR a numeric
            # lesson_no, so signature/notes rows like "注：…" / "本课程教师:" are
            # filtered out — those rows have at most 1-2 non-empty cells and none
            # of them are `content` / `title` / `lesson_no`.
            if not (item.get("content") or item.get("title") or _is_lesson_no(item.get("lesson_no", ""))):
                continue
            out.append(item)

    # (1) header table body
    body = logical_tables[header.table_idx][header.row_idx + 1 :]
    _emit(body)

    # (2) continuation tables — same logical width, no header detected.
    for ti, tbl in enumerate(logical_tables):
        if ti == header.table_idx:
            continue
        if not tbl:
            continue
        # We treat the WHOLE table as body if width matches the header.
        same_width = any(len(r) == header.width for r in tbl)
        if not same_width:
            # Sometimes the continuation table is slightly narrower because of
            # different merge layouts (e.g. 9 vs 16 physical cells but BOTH
            # collapse to 7 logical). The collapser already normalizes that, so
            # we additionally accept tables whose MAX width is ≥ field_map's
            # rightmost col index + 1 — i.e. enough cols to address every field.
            max_col_needed = max(header.field_map.keys()) + 1
            if not any(len(r) >= max_col_needed for r in tbl):
                continue
        _emit(tbl)

    return out


def _is_lesson_no(v: str) -> bool:
    if not v:
        return False
    n = v.strip()
    return bool(re.fullmatch(r"\d{1,3}", n))


# ── Preview builder (so frontend can show "what backend saw") ────


def _build_preview(
    logical_tables: List[List[List[str]]],
    header: Optional[_HeaderHit],
) -> List[List[str]]:
    """Return up to 12 rows from the best table (header + 11 following) so the
    frontend can render exactly what we parsed even if 0 lessons matched."""
    if header is None:
        # Pick the widest table's first 12 rows.
        if not logical_tables:
            return []
        best_ti = max(range(len(logical_tables)),
                      key=lambda i: max((len(r) for r in logical_tables[i]), default=0))
        return [r[:14] for r in logical_tables[best_ti][:12]]
    tbl = logical_tables[header.table_idx]
    start = max(0, header.row_idx)
    return [r[:14] for r in tbl[start : start + 12]]


# ── Format-specific entry points ─────────────────────────────────


def _parse_xlsx(file_bytes: bytes) -> Tuple[Dict[str, str], List[Dict[str, str]], List[List[str]]]:
    import openpyxl  # local import to avoid hard dep at module import

    wb = openpyxl.load_workbook(io.BytesIO(file_bytes), data_only=True)
    logical_tables: List[List[List[str]]] = []
    for sn in wb.sheetnames:
        ws = wb[sn]
        rows: List[List[str]] = []
        for r in range(1, min(ws.max_row, 400) + 1):
            raw_row = [ws.cell(r, c).value for c in range(1, min(ws.max_column, 40) + 1)]
            rows.append(_collapse_xlsx_row(raw_row))
        if rows:
            logical_tables.append(rows)
    return _finalize(logical_tables)


def _parse_docx(file_bytes: bytes) -> Tuple[Dict[str, str], List[Dict[str, str]], List[List[str]]]:
    import docx as _docx

    doc = _docx.Document(io.BytesIO(file_bytes))
    logical_tables: List[List[List[str]]] = []
    for tbl in doc.tables:
        rows = [_collapse_docx_row(r) for r in tbl.rows]
        if rows:
            logical_tables.append(rows)
    return _finalize(logical_tables)


def _parse_pdf(file_bytes: bytes) -> Tuple[Dict[str, str], List[Dict[str, str]], List[List[str]]]:
    """Use pdfplumber.extract_tables() then collapse merged cells."""
    import pdfplumber  # type: ignore

    logical_tables: List[List[List[str]]] = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            try:
                tables = page.extract_tables() or []
            except Exception as e:  # pdfplumber occasionally chokes on bad rules
                logger.warning(f"[zhuke] pdfplumber.extract_tables failed on page: {e}")
                tables = []
            for t in tables:
                rows = [_collapse_pdf_row(r) for r in t]
                rows = [r for r in rows if r]  # drop fully-empty rows
                if rows:
                    logical_tables.append(rows)
    return _finalize(logical_tables)


def _finalize(
    logical_tables: List[List[List[str]]],
) -> Tuple[Dict[str, str], List[Dict[str, str]], List[List[str]]]:
    """Common end-of-pipeline: pick header, mine cover, extract lessons, build
    preview. Returns ALWAYS-NON-EMPTY preview so the frontend can show what we
    actually saw even when 0 lessons matched."""
    header = _pick_header_across_tables(logical_tables)
    cover: Dict[str, str] = {}
    lessons: List[Dict[str, str]] = []
    if header is not None:
        cover = _extract_cover_from_band(logical_tables, header)
        lessons = _extract_lessons(logical_tables, header)
        # Backfill cover fields that live in the BODY (e.g. teacher column,
        # which 珠科 templates always put per-lesson rather than in the cover
        # band). Use the most-common non-empty value across lessons.
        _backfill_cover_from_lessons(cover, lessons)
    preview = _build_preview(logical_tables, header)
    return cover, lessons, preview


def _backfill_cover_from_lessons(cover: Dict[str, str], lessons: List[Dict[str, str]]) -> None:
    """Fill cover-eligible single-value body fields (teacher / class_name when
    rare) using the most-common value across lesson rows."""
    if not lessons:
        return
    from collections import Counter

    for cover_key, body_key in (("teacher", "teacher"),):
        if cover.get(cover_key):
            continue
        vals = [l.get(body_key, "").strip() for l in lessons if l.get(body_key)]
        if not vals:
            continue
        most_common, _ = Counter(vals).most_common(1)[0]
        if most_common:
            cover[cover_key] = most_common


def parse_schedule(file_bytes: bytes, ext: str) -> Dict[str, Any]:
    """Return ``{cover, lessons, raw_preview, ext}`` parsed from the upload.

    cover keys: college / course_name / course_type / teacher / class_name
                (+ optional course_code / total_weeks / total_hours /
                 theory_hours / practice_hours / exam_type / student_count)
    lesson keys: lesson_no / title / week / weekday / periods / date / content / hours
    raw_preview: a 2D list with up to 12 logical rows for the UI to display
                 (NON-EMPTY even when lessons==[]).
    """
    e = ext.lower().lstrip(".")
    if e in ("xlsx", "xlsm", "xls"):
        cover, lessons, preview = _parse_xlsx(file_bytes)
    elif e in ("docx", "doc"):
        cover, lessons, preview = _parse_docx(file_bytes)
    elif e == "pdf":
        cover, lessons, preview = _parse_pdf(file_bytes)
    else:
        raise ValueError(f"unsupported schedule format: .{e}")
    return {
        "cover": cover,
        "lessons": lessons,
        "raw_preview": preview,
        "ext": e,
    }


# ─────────────────────────────────────────────────────────────────
# Kimi K2 agent
# ─────────────────────────────────────────────────────────────────

_KIMI_SYSTEM_PROMPT = (
    "你是珠海科技学院的资深课程教研专家，正在为大学课程逐节填写《教学设计（教案）》。"
    "请严格按 JSON 返回结果（不要 Markdown 代码围栏，不要解释），字段及要求如下：\n"
    "{\n"
    '  "学情分析": "200-400字，结合本节课的内容，分三段：一、知识基础；二、能力特点；三、思想特点。",\n'
    '  "教学目标": "200-400字，按 1. 思政目标 2. 知识目标 3. 能力目标 4. 素质目标 四个编号段落。",\n'
    '  "主要教学内容": "100-200字，描述本节课的主要内容，最后一行写 教学重点：xxx  教学难点：xxx。",\n'
    '  "教学方法": "勾选并简述：项目教学法☑ 讨论法☑ PBL☑ 练习法 案例教学法☑ 讲授法☑ 其他 中适用的方法。",\n'
    '  "教学媒体": "勾选并简述：教材☑ 多媒体☑ AI 工具☑ 视频☑ 云平台空间 板书☑ 其他 中适用的媒体。",\n'
    '  "教学过程设计": "600-1000字，必须含 1. 场景导入（15分钟） 2. 讲解具体实现方式（70分钟，含 2.1/2.2/2.3 子节） 3. 课堂小结（5分钟） 三大段。",\n'
    '  "作业布置": "120-250字，思考题与实践任务编号列出。",\n'
    '  "参考资料": "100-200字，严格按 GB/T 7714-2015 国标格式列出 3-5 条文献。'
    '格式规则：\\n'
    '  - 专著: 作者. 书名[M]. 版本. 出版地: 出版者, 出版年: 起止页.\\n'
    '  - 期刊: 作者. 题名[J]. 刊名, 年, 卷(期): 起止页.\\n'
    '  - 会议: 作者. 题名[C]//会议名. 出版地: 出版者, 年: 起止页.\\n'
    '  - 报纸: 作者. 题名[N]. 报纸名, 年-月-日(版次).\\n'
    '  - 电子资源: 作者. 题名[EB/OL]. (发表日期)[引用日期]. URL.\\n'
    '必须分 (一)教材 与 (二)参考资料 两块，每块标题独占一行；'
    '每条文献独占一行，禁止在逗号/连字符处断行；'
    '必须标注文献类型 [M]/[J]/[C]/[N]/[EB/OL]；标点用半角；'
    '3 位以上作者写 et al. 或 等；优先选近 5 年文献。",\n'
    '  "评估与反馈": "150-300字，包含 1. 课堂评估 2. 反馈机制 两个编号段落。"\n'
    "}\n"
    "格式要求：每个编号条目单独占一行（用换行符 \\n 分隔）；教学过程设计的 1./2./3. 及 2.1/2.2/2.3 必须分行。"
    "禁止使用 Markdown（无 **、无 - 列表、无代码围栏）；编号条目之间必须用 \\n 分隔。"
    "全部中文，避免空泛口号，紧扣本节课主题。"
)

# 通用教学标准（全局基线）：珠科属大学教学内容创作，注入精简版；仍须只输出 JSON。
from app.services import teacher_standard as _teacher_standard  # noqa: E402
_KIMI_SYSTEM_PROMPT = (
    _KIMI_SYSTEM_PROMPT + "\n\n" + _teacher_standard.standard_brief()
    + "（注意：以上为教学标准，本次仍必须严格只输出上述 JSON 结构。）"
)

_CHECKBOX_ITEMS = (
    "项目教学法",
    "讨论法",
    "PBL",
    "练习法",
    "案例教学法",
    "讲授法",
    "教材",
    "多媒体",
    "AI 工具",
    "视频",
    "云平台空间",
    "板书",
)

from app.core.kimi_zhuke_config import KIMI_K2_CONCURRENCY, KIMI_K2_MODEL, KIMI_K2_TIMEOUT_SEC

_kimi_executor: Optional["ThreadPoolExecutor"] = None


def get_kimi_executor():
    """Dedicated thread pool for parallel Kimi SubAgent calls."""
    global _kimi_executor
    from concurrent.futures import ThreadPoolExecutor

    if _kimi_executor is None:
        _kimi_executor = ThreadPoolExecutor(
            max_workers=KIMI_K2_CONCURRENCY,
            thread_name_prefix="kimi-zhuke",
        )
    return _kimi_executor


_RE_GB7714_TYPE = re.compile(r"\[(M|J|N|C|D|R|S|P|EB/OL|DB/OL|J/OL|M/OL)\]")
_RE_REF_SECTION_HEAD = re.compile(r"^\([一二]\)")
_RE_REF_NUMBERED = re.compile(r"^\[\s*\d+\s*\]")
_RE_GB7714_INLINE_TYPE = re.compile(r"\[(?:M|J|N|C|D|R|S|P|EB/OL|DB/OL|J/OL|M/OL)\]")


def format_section_for_docx(text: str, *, reference: bool = False) -> str:
    """Single entry point: clean AI output for docx paragraph layout.

    Rules (additive — order matters):
      - Strip ```code fences / ** bold / __underline / leading ``- `` bullets
      - Normalise weird numbering variants (``1、``, ``(1)``) to ``1. ``
      - Force every standalone enumeration onto its own line
      - Collapse runs of ASCII spaces / tabs to 1 (Word renders them at full
        width which looks broken in CJK paragraphs)
      - Collapse triple+ blank lines to 2
    """
    if not text:
        return ""
    t = str(text).strip()
    # 1) Markdown residue & basic normalization
    t = re.sub(r"^```(?:json)?\s*|\s*```$", "", t, flags=re.IGNORECASE)
    t = re.sub(r"\r\n?", "\n", t)
    t = re.sub(r"\*\*(.+?)\*\*", r"\1", t)
    t = re.sub(r"__(.+?)__", r"\1", t)
    t = re.sub(r"(?m)^-\s+", "", t)
    t = re.sub(r"(?m)^(\d+)、\s*", r"\1. ", t)
    t = re.sub(r"(?m)^\((\d+)\)\s*", r"\1. ", t)
    t = re.sub(r"(?<!\n)\((\d+)\)\s*", r"\n\1. ", t)
    if reference:
        t = re.sub(r"[ \t]{2,}", " ", t)
        t = re.sub(r" +\n", "\n", t)
        t = re.sub(r"\n{3,}", "\n\n", t)
        return _format_reference_section(t)
    # 2) 重点/难点 on their own line
    t = re.sub(r"(?<!\n)(教学重点：)", r"\n\1", t)
    t = re.sub(r"(?<!\n)(教学难点：)", r"\n\1", t)
    # 3) Arabic numbering after sentence-ending punctuation
    t = re.sub(r"(?<=[。；!！?？])\s*(\d+\.)", r"\n\1", t)
    t = re.sub(r"(?<=[。；!！?？])\s*(\d+\.\d+\s)", r"\n\1", t)
    t = re.sub(r"(?<=[。；;])\s*(一、|二、|三、|四、|五、|六、|七、|八、|九、|十、)", r"\n\1", t)
    # 4) Standalone Arabic enumeration
    t = re.sub(r"(?<![\n\d])(\d+\.)", r"\n\1", t)
    t = re.sub(r"(?<!\n)(\d+\.\d+\s)", r"\n\1", t)
    # 5) Chinese enumeration after CJK char (not after another enumeration)
    t = re.sub(
        r"(?<=[^\n])(?<![\d.])(一、|二、|三、|四、|五、|六、|七、|八、|九、|十、)",
        r"\n\1",
        t,
    )
    # 6) Parenthesised enumeration — both Chinese and Arabic, both bracket styles.
    t = re.sub(
        r"(?<!\n)（\s*([一二三四五六七八九十]|\d+)\s*）",
        r"\n（\1）",
        t,
    )
    t = re.sub(
        r"(?<!\n)\(\s*([一二三四五六七八九十]|\d+)\s*\)",
        r"\n(\1)",
        t,
    )
    # 7) Checkbox method names — collapse residual spaces (the lint check
    # depends on this to stay clean and skip Kimi layout review).
    for item in _CHECKBOX_ITEMS:
        t = re.sub(rf"{re.escape(item)}\s*☑", f"{item}☑", t)
    # 8) Collapse runs of ASCII spaces / tabs (Word CJK paragraphs render
    # double-space as a visible gap — Kimi sometimes adds 2-3 between Chinese
    # phrases which looks broken in 仿宋).
    t = re.sub(r"[ \t]{2,}", " ", t)
    # 9) Tidy trailing spaces & collapse blank lines — Word renders each "\n\n"
    # as an empty paragraph (visible gap). Keep one newline between items only.
    t = re.sub(r" +\n", "\n", t)
    t = re.sub(r"\n{2,}", "\n", t)
    return t.strip()


def _format_reference_section(text: str) -> str:
    """GB/T 7714-2015 reference block: join broken lines, halfwidth punct."""
    if not text:
        return ""
    t = str(text).strip()
    # Normalise numbered ref markers: "[ 1 ]" -> "[1]".
    t = re.sub(r"\[\s*(\d+)\s*\]", r"[\1]", t)
    # Section headers on their own line, content follows on next line.
    t = re.sub(r"\((一)\)教材\s*", r"(一)教材\n", t)
    t = re.sub(r"\((二)\)参考资料\s*", r"(二)参考资料\n", t)
    t = re.sub(r"(?<!\n)\(([一二])\)", r"\n(\1)", t)
    # Numbered refs [1] [2] each on own line.
    t = re.sub(r"(?<!\n)\[(\d+)\]", r"\n[\1]", t)
    # Each GB/T type marker starts a new entry when glued to prior text.
    t = re.sub(
        r"(?<!\n)([^\n]{8,}?)(\[(?:M|J|N|C|D|R|S|P|EB/OL|DB/OL|J/OL|M/OL)\])",
        r"\n\1\2",
        t,
    )
    # Join lines broken mid-citation (comma/colon/hyphen continuation).
    raw_lines = [ln.strip() for ln in t.split("\n")]
    merged: List[str] = []
    for line in raw_lines:
        if not line:
            continue
        if (
            line.startswith("(一)教材")
            or line.startswith("(二)参考资料")
            or _RE_REF_NUMBERED.match(line)
            or not merged
        ):
            merged.append(line)
            continue
        prev = merged[-1]
        prev_stripped = prev.rstrip()
        if prev_stripped.startswith("(一)教材") or prev_stripped.startswith("(二)参考资料"):
            merged.append(line)
            continue
        ends_with_terminal = prev_stripped.endswith((".", "。"))
        ends_with_continuation = prev_stripped.endswith(
            (",", "，", ":", "：", "-", "—")
        )
        if ends_with_continuation or (
            not ends_with_terminal and not _RE_GB7714_INLINE_TYPE.search(line[:20])
        ):
            joiner = "" if prev_stripped.endswith("-") else " "
            merged[-1] = prev_stripped + joiner + line.lstrip("-")
        else:
            merged.append(line)
    t = "\n".join(merged)
    # GB/T 7714 uses halfwidth punctuation in bibliographic entries.
    t = t.replace("，", ",").replace("。", ".").replace("：", ":")
    t = t.replace("；", ";").replace("（", "(").replace("）", ")")
    t = re.sub(r"\n{2,}", "\n", t)
    return t.strip()


def normalize_section_text(text: str) -> str:
    """Alias for :func:`format_section_for_docx` (backward compatible)."""
    return format_section_for_docx(text)


def normalize_sections(sections: Dict[str, Any]) -> Dict[str, str]:
    return {
        str(k): normalize_section_text(str(v))
        for k, v in (sections or {}).items()
        if v is not None and str(v).strip()
    }


def _zhuke_tmp_dir_for_preflight() -> str:
    d = os.path.join(settings.FILES_DIR, "tmp_zhuke")
    os.makedirs(d, exist_ok=True)
    return d


def validate_zhuke_preflight(*, skip_ai: bool = False) -> None:
    """Raise RuntimeError when template / disk / Kimi config is unusable."""
    if not template_exists():
        raise RuntimeError("后端缺少珠科教案模板文件，请联系管理员")
    tmp = _zhuke_tmp_dir_for_preflight()
    probe = os.path.join(tmp, ".write_probe")
    try:
        with open(probe, "w", encoding="utf-8") as f:
            f.write("ok")
        os.remove(probe)
    except OSError as e:
        raise RuntimeError(f"珠科临时目录不可写：{tmp}（{e}）") from e
    if not skip_ai:
        KimiAgent()._client()


_RE_NUMBERED_INLINE = re.compile(r"(?<=[\u4e00-\u9fff。；;!！?？])\s*\d+\.")
_RE_SUBSECTION_INLINE = re.compile(r"(?<=[\u4e00-\u9fff。；;!！?？])\s*\d+\.\d+\s")
_RE_CHINESE_ENUM_INLINE = re.compile(r"(?<=[\u4e00-\u9fff。；;])\s*(二、|三、|四、|五、|六、)")
_RE_FOCUS_NOT_LINE = re.compile(r"[^\n]教学重点：")
_RE_DIFF_NOT_LINE = re.compile(r"[^\n]教学难点：")
_RE_CHECKBOX_SPACE = re.compile(
    r"(?:项目教学法|讨论法|PBL|练习法|案例教学法|讲授法|教材|多媒体|AI 工具|视频|云平台空间|板书)\s+☑"
)
_RE_MARKDOWN_RESIDUE = re.compile(r"\*\*|```|^-\s+", re.MULTILINE)


def lint_sections_format(sections: Dict[str, Any]) -> List[str]:
    """Return human-readable issue strings; empty list means rule layer is clean.

    Scope intentionally narrowed (2026-05): only flag issues that the local
    `format_section_for_docx` normaliser CANNOT fix on its own. Anything that
    normalize already handles (阿拉伯/子节/中文序号未换行 / 教学重点 / 教学
    难点 inline) is skipped here so we do NOT trigger a second 60-180s Kimi
    layout-review call per lesson for issues that are already fully repaired
    by post-processing. Saves roughly half the wall-clock per lesson.
    """
    issues: List[str] = []
    for key, raw in (sections or {}).items():
        text = str(raw or "")
        if not text.strip():
            continue
        label = str(key)
        if _RE_MARKDOWN_RESIDUE.search(text):
            issues.append(f"{label}: 含 Markdown 残留")
        if _RE_CHECKBOX_SPACE.search(text):
            issues.append(f"{label}: 勾选项与方法名之间有多余空格")
        if label == "参考资料" and text.strip() and not _RE_GB7714_TYPE.search(text):
            issues.append("参考资料: 缺少文献类型标注 [M]/[J]/[C]/[N]/[EB/OL] 等")
    return issues


def format_sections_for_docx(sections: Dict[str, str]) -> Dict[str, str]:
    """Apply local docx layout rules to every section before lint / optional Kimi review."""
    return {
        key: format_section_for_docx(value or "", reference=(key == "参考资料"))
        for key, value in sections.items()
    }


def layout_review_on_lint_enabled() -> bool:
    return os.getenv("ZHUKE_LAYOUT_REVIEW_ON_LINT", "").strip().lower() in (
        "1",
        "true",
        "yes",
    )


def layout_review_always_enabled() -> bool:
    return os.getenv("ZHUKE_LAYOUT_REVIEW_ALWAYS", "").strip().lower() in (
        "1",
        "true",
        "yes",
    )


_LAYOUT_REVIEW_SYSTEM = (
    "你是珠科教案 Word 排版质检专家。你会收到一课 9 个栏目的 JSON 文本。"
    "只修正排版格式，不改变语义与字数；禁止删改专业术语和实质内容。"
    "检查并修正："
    "1) 编号条目(1./2./2.1/一、/（一）)各自独占一行；"
    "2) 教学重点：与 教学难点：各占一行；"
    "3) 勾选项格式为「方法名☑」不含空格；"
    "4) 删除 Markdown(**、- 列表、```)；"
    "5) 禁止把本应分段的条目挤在同一行。"
    "返回相同 key 的 JSON object，不要 Markdown 围栏，不要解释。"
)


@dataclass
class KimiAgent:
    """Thin OpenAI-SDK wrapper that calls Moonshot."""

    api_key: str = field(default_factory=lambda: settings.KIMI_API_KEY or os.getenv("MOONSHOT_API_KEY", ""))
    base_url: str = field(default_factory=lambda: settings.KIMI_BASE_URL or "https://api.moonshot.cn/v1")
    model: str = field(default_factory=lambda: KIMI_K2_MODEL)
    timeout_sec: float = field(default_factory=lambda: KIMI_K2_TIMEOUT_SEC)

    def _client(self):
        from openai import OpenAI

        if not self.api_key:
            raise RuntimeError("KIMI_API_KEY 未配置；请在 backend/.env 填入 KIMI_API_KEY=sk-...")
        return OpenAI(api_key=self.api_key, base_url=self.base_url, timeout=self.timeout_sec)

    def generate_lesson(
        self,
        *,
        course_name: str,
        lesson_title: str,
        time_label: str,
        hours: str,
        outline: str,
        major: str,
    ) -> Dict[str, str]:
        """Synchronous call (this runs from FastAPI but blocks the event loop;
        the API endpoint should wrap with `run_in_executor` if many lessons)."""
        user_msg = (
            f"课程名称：{course_name}\n"
            f"专业方向：{major}\n"
            f"本节课题目：{lesson_title}\n"
            f"授课时间：{time_label}\n"
            f"授课学时数：{hours}\n"
            f"教学日历给出的授课内容大纲：{outline or '（无）'}\n\n"
            "请基于上述信息按系统提示给出 JSON。"
        )
        client = self._client()
        # NOTE: Moonshot's `kimi-k2.6` rejects any temperature other than 1 with
        # `invalid temperature: only 1 is allowed for this model`. We omit the
        # field entirely so the server uses its native default — compatible with
        # both kimi-k2.6 (strict temp=1) and older models like kimi-k2-0905-preview
        # where the default is also 1.0. If you really need a custom temperature
        # for a different Kimi model, set env `KIMI_K2_TEMPERATURE`.
        create_kwargs: Dict[str, Any] = {
            "model": self.model,
            "messages": [
                {"role": "system", "content": _KIMI_SYSTEM_PROMPT},
                {"role": "user", "content": user_msg},
            ],
            "response_format": {"type": "json_object"},
        }
        temp_env = os.getenv("KIMI_K2_TEMPERATURE", "").strip()
        if temp_env:
            try:
                create_kwargs["temperature"] = float(temp_env)
            except ValueError:
                pass
        completion = client.chat.completions.create(**create_kwargs)
        raw = (completion.choices[0].message.content or "").strip()
        try:
            data = json.loads(raw)
        except Exception:
            # Best-effort recovery: extract first {...} block.
            m = re.search(r"\{[\s\S]*\}", raw)
            data = json.loads(m.group(0)) if m else {}
        return normalize_sections({k: str(v).strip() for k, v in (data or {}).items()})


@dataclass
class LessonSubAgent(KimiAgent):
    """Per-lesson SubAgent — isolated client, no shared conversation context."""

    def generate_lesson(
        self,
        *,
        course_name: str,
        lesson_title: str,
        time_label: str,
        hours: str,
        outline: str,
        major: str,
    ) -> Dict[str, str]:
        return super().generate_lesson(
            course_name=course_name,
            lesson_title=lesson_title,
            time_label=time_label,
            hours=hours,
            outline=outline,
            major=major,
        )


@dataclass
class LayoutReviewAgent(KimiAgent):
    """Post-generation layout QA — fixes line breaks / numbering before docx write."""

    def review_sections(self, sections: Dict[str, str]) -> Dict[str, str]:
        if not sections:
            return {}
        client = self._client()
        user_msg = (
            "请检查并修正以下 JSON 各栏目的排版（只改格式）：\n"
            + json.dumps(sections, ensure_ascii=False)
        )
        create_kwargs: Dict[str, Any] = {
            "model": self.model,
            "messages": [
                {"role": "system", "content": _LAYOUT_REVIEW_SYSTEM},
                {"role": "user", "content": user_msg},
            ],
            "response_format": {"type": "json_object"},
        }
        temp_env = os.getenv("KIMI_K2_TEMPERATURE", "").strip()
        if temp_env:
            try:
                create_kwargs["temperature"] = float(temp_env)
            except ValueError:
                pass
        completion = client.chat.completions.create(**create_kwargs)
        raw = (completion.choices[0].message.content or "").strip()
        try:
            data = json.loads(raw)
        except Exception:
            m = re.search(r"\{[\s\S]*\}", raw)
            data = json.loads(m.group(0)) if m else sections
        merged = {**sections, **{k: str(v).strip() for k, v in (data or {}).items() if v is not None}}
        return normalize_sections(merged)


# ─────────────────────────────────────────────────────────────────
# DOCX assembly
# ─────────────────────────────────────────────────────────────────

# Labels used to locate cover paragraph lines in the template.
_COVER_PARA_LABELS: List[Tuple[str, str]] = [
    ("学", "college"),       # "学       院"
    ("专", "major"),         # "专       业"
    ("班", "class_name"),    # "班       级"
    ("授 课 类 别", "course_type"),
    ("课 程 名 称", "course_name"),
    ("主 讲 教 师", "teacher"),
]


def _para_starts_with_cover_label(text: str, label_seed: str) -> bool:
    norm = re.sub(r"\s+", "", text)
    if label_seed == "学":
        return norm.startswith("学院")
    if label_seed == "专":
        return norm.startswith("专业")
    if label_seed == "班":
        return norm.startswith("班级")
    if label_seed == "授 课 类 别":
        return norm.startswith("授课类别")
    if label_seed == "课 程 名 称":
        return norm.startswith("课程名称")
    if label_seed == "主 讲 教 师":
        return norm.startswith("主讲教师")
    return False


def _replace_cover_runs(paragraph, label_seed: str, value: str) -> None:
    """Write the user-provided value into the underlined fill-in area of a
    cover paragraph WITHOUT touching the label runs or the underline formatting.

    The official 珠科 template uses two kinds of runs in each cover line:
      1. Non-underlined label runs (e.g. `学` / `       ` / `院 `) that
         visually compose `学       院`.
      2. Trailing underlined runs of blank spaces (`[U=True]`) that visually
         render as the fill-in line `_________________`.

    The OLD code wrote `f"{label}：{value}"` into the first run and cleared
    the rest, destroying the underline runs entirely. Users lost the
    fill-in line and the value rendered as plain text without underline.

    The NEW algorithm:
      - Find the contiguous trailing block of underlined runs.
      - Replace the first of them with `{value}` (preserving its U=True,
        font, size — runs inherit theme defaults if unset).
      - Clear the remaining underlined runs so the line ends right after
        the value.
      - Leave label runs (and any separator like `：`) completely alone.

    Falls back to the legacy overwrite path only if no underlined run is
    found (defensive against future template variants).
    """
    runs = paragraph.runs
    if not runs:
        return

    # Locate the trailing block of underlined runs (the fill-in area).
    underline_idx_start = None
    for i, r in enumerate(runs):
        if _run_is_underlined(r):
            if underline_idx_start is None:
                underline_idx_start = i
        else:
            # Reset if a non-underlined run breaks the trailing block; the
            # template puts ALL underlined runs at the tail so this works.
            if underline_idx_start is not None and i < len(runs):
                # If we hit a non-underline run after starting one, we are no
                # longer in the trailing block; reset.
                underline_idx_start = None

    safe_value = (value or "").strip()
    if underline_idx_start is None:
        # Legacy fallback: no underlined fill-in area, write `{label}：{value}`
        # into first run and clear the rest. Same as the old behavior.
        label_text = {
            "学": "学       院",
            "专": "专       业",
            "班": "班       级",
            "授 课 类 别": "授 课 类 别",
            "课 程 名 称": "课 程 名 称",
            "主 讲 教 师": "主 讲 教 师",
        }.get(label_seed, label_seed)
        runs[0].text = f"{label_text}：{safe_value}"
        _set_run_font_cjk(runs[0])
        for r in runs[1:]:
            r.text = ""
        return

    # Measure total visible "underline width" the template originally drew —
    # sum of all underline-run text lengths. Treats CJK glyphs as width 2 and
    # ASCII as width 1 so the right edge lands roughly at the same column
    # regardless of what value the user supplies. Without this preservation
    # the underline would shrink visually whenever we wrote a short value
    # (e.g. "数学" replaces 17 spaces → fill-in line collapses to 4 cells).
    def _width(s: str) -> int:
        return sum(2 if ord(c) > 127 else 1 for c in s)

    total_underline_width = sum(_width(r.text or "") for r in runs[underline_idx_start:] if _run_is_underlined(r))
    value_width = _width(safe_value)
    pad_width = max(0, total_underline_width - value_width)
    # Pad with full-width ideographic spaces (U+3000) so the underline stays
    # visually flush with the template (half-width ASCII spaces would render
    # noticeably narrower under CJK fonts).
    pad_spaces = "\u3000" * (pad_width // 2)
    # Write value into the first underlined run, padding into the same run so
    # the entire fill-in area keeps a single contiguous underline format.
    runs[underline_idx_start].text = safe_value + pad_spaces
    _set_run_font_cjk(runs[underline_idx_start])
    # Underline is preserved by the original run's rPr — we only replaced
    # text. Clear any subsequent underlined runs (their width is already
    # absorbed by the pad_spaces above).
    for r in runs[underline_idx_start + 1:]:
        if _run_is_underlined(r):
            r.text = ""


def _run_is_underlined(run) -> bool:
    """python-docx returns None for "unset" and True/enum for set. We treat
    any truthy value (True, 'single', etc.) as underlined."""
    u = run.font.underline if hasattr(run.font, "underline") else None
    if u is None:
        # Fall back to legacy attr that some run wrappers expose.
        u = getattr(run, "underline", None)
    return bool(u) and u is not False


# 珠科教案模板正文规定使用「小四号仿宋」(see template P039 "正文采用小四号
# 仿宋、1.25 倍行间距"). Centralized here so worker output matches the spec.
_ZHUKE_CJK_FONT = "仿宋"
_ZHUKE_CJK_FONT_SIZE_PT = 12  # 小四 = 12pt
_ZHUKE_LINE_SPACING = 1.25
_ZHUKE_FIRST_LINE_INDENT_CHARS = 2  # 中文段落首行缩进 2 字

# Tier-2 section labels used for row lookup / orphan-table detection.
_TIER2_SECTION_LABELS = (
    "学情分析",
    "教学目标",
    "主要教学内容",
    "教学方法",
    "教学媒体",
    "教学过程设计",
    "复习思考及作业布置",
    "参考资料",
    "评估与反馈",
    "教学重点",
    "教学难点",
)

_RE_INDENT_SKIP_PATTERNS = (
    re.compile(r"^\d+\.\s"),
    re.compile(r"^\d+\.\d+\s"),
    re.compile(r"^[一二三四五六七八九十]+、"),
    re.compile(r"^（[一二三四五六七八九十\d]+）"),
    re.compile(r"^\([一二三四五六七八九十\d]+\)"),
    re.compile(r"^教学重点："),
    re.compile(r"^教学难点："),
    re.compile(r"^\(一\)教材"),
    re.compile(r"^\(二\)参考资料"),
)


def _should_first_line_indent(line: str) -> bool:
    """True when a line is prose body text that should get 2-char first-line indent."""
    s = (line or "").strip()
    if not s:
        return False
    # Short metadata lines (学时、节次) — no indent.
    if len(s) <= 24 and "。" not in s and "；" not in s and "!" not in s and "？" not in s:
        return False
    for pat in _RE_INDENT_SKIP_PATTERNS:
        if pat.match(s):
            return False
    return True


def _apply_body_line_format(paragraph, line: str, *, reference: bool = False) -> None:
    """Line spacing + conditional first-line indent for one body line."""
    pf = paragraph.paragraph_format
    pf.line_spacing = _ZHUKE_LINE_SPACING
    if reference:
        _apply_zhuke_reference_paragraph_format(paragraph)
        return
    from docx.oxml.ns import qn

    pPr = paragraph._p.get_or_add_pPr()
    ind = pPr.find(qn("w:ind"))
    if _should_first_line_indent(line):
        _set_first_line_indent_chars(paragraph)
    else:
        pf.first_line_indent = None
        if ind is not None:
            for k in ("w:firstLine", "w:firstLineChars", "w:hanging", "w:hangingChars"):
                if ind.get(qn(k)) is not None:
                    del ind.attrib[qn(k)]


def _set_run_font_cjk(run, name: str = _ZHUKE_CJK_FONT, size_pt: Optional[int] = None) -> None:
    """Force a run to render Chinese glyphs in the given font.

    python-docx's high-level `run.font.name = X` only sets w:ascii (and
    sometimes w:hAnsi). Word ignores those for CJK characters and falls back
    to whatever the theme says — usually 宋体 — so our 仿宋 setting gets
    dropped and the user sees ugly 宋体 in their export. We have to write the
    `w:eastAsia` font slot explicitly via lxml to bind CJK glyphs to the
    intended typeface.

    Optional `size_pt` sets the run size in points; the template spec asks
    for 小四 (12pt) for body text.
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt

    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)
    rFonts.set(qn("w:eastAsia"), name)
    rFonts.set(qn("w:cs"), name)
    if size_pt is not None:
        run.font.size = Pt(size_pt)


def _set_first_line_indent_chars(
    paragraph,
    chars: int = _ZHUKE_FIRST_LINE_INDENT_CHARS,
    font_size_pt: int = _ZHUKE_CJK_FONT_SIZE_PT,
) -> None:
    """Write w:firstLineChars + w:firstLine so Word/WPS indent N CJK chars."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    pPr = paragraph._p.get_or_add_pPr()
    ind = pPr.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind")
        pPr.append(ind)
    ind.set(qn("w:firstLineChars"), str(chars * 100))
    ind.set(qn("w:firstLine"), str(font_size_pt * 20 * chars))
    for k in ("w:hanging", "w:hangingChars"):
        if ind.get(qn(k)) is not None:
            del ind.attrib[qn(k)]


def _apply_zhuke_paragraph_format(paragraph) -> None:
    """1.25 line spacing + 2-char first-line indent per template's 编写说明."""
    pf = paragraph.paragraph_format
    pf.line_spacing = _ZHUKE_LINE_SPACING
    _set_first_line_indent_chars(paragraph)


def _apply_zhuke_reference_paragraph_format(paragraph) -> None:
    """Reference entries: line spacing only, no first-line indent (GB/T 7714)."""
    pf = paragraph.paragraph_format
    pf.line_spacing = _ZHUKE_LINE_SPACING
    pf.first_line_indent = None
    from docx.oxml.ns import qn

    pPr = paragraph._p.get_or_add_pPr()
    ind = pPr.find(qn("w:ind"))
    if ind is not None:
        for k in ("w:firstLine", "w:firstLineChars", "w:hanging", "w:hangingChars"):
            if ind.get(qn(k)) is not None:
                del ind.attrib[qn(k)]


def _set_cell_text(cell, text: str) -> None:
    """Replace cell content with paragraph(s) in 仿宋 + template body spacing.

    Tier 1 single-line short values (e.g. "2 学时", "第 1 周 星期一  第 3、4 节")
    skip the 2-char first-line indent — the indent looks visibly wrong when
    the value is the cell's only line, since it pushes a 4-character value
    to the middle of the cell.
    """
    text = text or ""
    lines = [ln.strip() for ln in text.split("\n")]
    lines = [ln for ln in lines if ln]
    if not lines:
        lines = [""]

    font_size = None
    for p in cell.paragraphs:
        for r in p.runs:
            if r.font.size:
                font_size = r.font.size
                break
        if font_size:
            break

    for p in list(cell.paragraphs):
        p._element.getparent().remove(p._element)

    for line in lines:
        p = cell.add_paragraph()
        _apply_body_line_format(p, line, reference=False)
        run = p.add_run(line)
        size_pt = None if font_size else _ZHUKE_CJK_FONT_SIZE_PT
        _set_run_font_cjk(run, size_pt=size_pt)
        if font_size:
            run.font.size = font_size


def _short_title(content: str) -> str:
    """Extract a short 授课题目 from a lesson content. Splits at the first
    Chinese punctuation that separates the topic from the description, falling
    back to the first 30 chars."""
    if not content:
        return ""
    text = str(content).strip()
    for sep in ("：", ":", "。", "！", "？", "；"):
        idx = text.find(sep)
        if 0 < idx < 30:
            return text[:idx].strip()
    return text[:30].strip()


def _next_logical_cell_in_row(table, target_cell):
    """Return the next logical cell (dedup'd by element id) after target_cell
    in the SAME row, or None. Used to navigate from a Tier 1 label cell to its
    sibling value cell."""
    target_eid = id(target_cell._element)
    for row in table.rows:
        seen, prev_was_target = set(), False
        for c in row.cells:
            eid = id(c._element)
            if eid in seen:
                continue
            seen.add(eid)
            if prev_was_target:
                return c
            if eid == target_eid:
                prev_was_target = True
    return None


def _set_value_after_label_colon(cell, value: str, *, reference: bool = False) -> None:
    """For Tier 2 cells (label+value combined in one cell): keep the LABEL
    paragraph(s) untouched (including the elaborate "（...）" description in the
    template) and replace everything AFTER the first paragraph that ends with
    "：" / ":" with the new value text.

    Algorithm:
      1) Find the first paragraph whose text contains "：" or ":".
      2) If that paragraph has sample value after the colon, truncate the
         paragraph to "label：" (preserving the first run's formatting).
      3) Remove every paragraph below the label paragraph.
      4) Append the value as one paragraph per non-empty line.
    """
    if not value:
        return  # leave the template untouched when no value supplied
    paras = list(cell.paragraphs)
    if not paras:
        cell.add_paragraph(str(value))
        return
    colon_para_idx = -1
    for i, p in enumerate(paras):
        if "：" in p.text or ":" in p.text:
            colon_para_idx = i
            break
    if colon_para_idx < 0:
        # No colon found anywhere — fall back to appending the value as a new
        # paragraph (label-only cells without colons are unusual but possible).
        cell.add_paragraph(str(value))
        return
    target_para = paras[colon_para_idx]
    text = target_para.text
    m = re.match(r"^(.*?[：:])\s*(.*)$", text, flags=re.DOTALL)
    if m and m.group(2):
        # Sample value present after the colon — truncate to "label：" only,
        # preserving the first run's formatting (font / size / weight).
        if target_para.runs:
            target_para.runs[0].text = m.group(1)
            for r in target_para.runs[1:]:
                r.text = ""
        else:
            target_para.add_run(m.group(1))
    # Remove ALL paragraphs after the label paragraph.
    for p in paras[colon_para_idx + 1 :]:
        p._element.getparent().remove(p._element)
    # Append the value as one paragraph per non-empty line (skip blank lines —
    # empty paragraphs in Word show up as visible gaps in the cell).
    for line in str(value).split("\n"):
        line = line.strip()
        if not line:
            continue
        new_p = cell.add_paragraph()
        _apply_body_line_format(new_p, line, reference=reference)
        run = new_p.add_run(line)
        _set_run_font_cjk(run, size_pt=_ZHUKE_CJK_FONT_SIZE_PT)


def _strip_tier2_cell_to_label_only(cell) -> None:
    """Keep the section label (…：) but remove template sample text and value paras.

    Used for 教学重点/教学难点 rows when the AI already folded them into
    主要教学内容 — avoids a whole empty table row from _clear_cell_paragraphs.
    """
    paras = list(cell.paragraphs)
    if not paras:
        return
    colon_para_idx = -1
    for i, p in enumerate(paras):
        if "：" in p.text or ":" in p.text:
            colon_para_idx = i
            break
    if colon_para_idx < 0:
        return
    target_para = paras[colon_para_idx]
    text = target_para.text
    m = re.match(r"^(.*?[：:])\s*(.*)$", text, flags=re.DOTALL)
    if m:
        label_only = m.group(1)
        if target_para.runs:
            target_para.runs[0].text = label_only
            for r in target_para.runs[1:]:
                r.text = ""
        else:
            target_para.add_run(label_only)
    for p in paras[colon_para_idx + 1 :]:
        p._element.getparent().remove(p._element)


def _fmt_section_value(value: Any, section_key: str = "") -> Optional[str]:
    """Return None to skip (leave template untouched), else cleaned value string."""
    if value is None:
        return None
    is_ref = section_key == "参考资料"
    text = format_section_for_docx(str(value), reference=is_ref)
    return text or None


def _clear_cell_paragraphs(cell) -> None:
    """Remove all paragraphs in a cell and leave one empty paragraph."""
    for p in list(cell.paragraphs):
        p._element.getparent().remove(p._element)
    cell.add_paragraph("")


def _row_for_cell(cell):
    """Return the ``<w:tr>`` element containing ``cell``."""
    tc = cell._element
    return tc.getparent() if tc is not None else None


def _remove_table_row(table, row) -> None:
    """Delete a row element from a table."""
    tr = row._element if hasattr(row, "_element") else row
    parent = tr.getparent()
    if parent is not None:
        parent.remove(tr)


def _remove_row_for_cell(table, cell) -> None:
    """Delete the entire table row that contains ``cell``."""
    tr = _row_for_cell(cell)
    if tr is None:
        return
    for row in list(table.rows):
        if row._element is tr:
            _remove_table_row(table, row)
            return


def _cell_body_text(cell) -> str:
    """Return user-visible body text in a cell (content after section label)."""
    paras = [p for p in cell.paragraphs if (p.text or "").strip()]
    if not paras:
        return ""
    for i, p in enumerate(paras):
        if "：" in p.text or ":" in p.text:
            m = re.match(r"^(.*?[：:])\s*(.*)$", p.text, flags=re.DOTALL)
            chunks: List[str] = []
            if m:
                tail = (m.group(2) or "").strip()
                if tail:
                    chunks.append(tail)
            else:
                tail = re.sub(r"^.*?[：:]", "", p.text, count=1).strip()
                if tail:
                    chunks.append(tail)
            chunks.extend((x.text or "").strip() for x in paras[i + 1 :])
            return "\n".join(x for x in chunks if x)
    return "\n".join(p.text.strip() for p in paras)


def _cell_is_effectively_empty(cell) -> bool:
    return not (_cell_body_text(cell).strip())


def _row_is_effectively_empty(row) -> bool:
    seen: Set[int] = set()
    for cell in row.cells:
        eid = id(cell._element)
        if eid in seen:
            continue
        seen.add(eid)
        if not _cell_is_effectively_empty(cell):
            return False
    return True


def _prune_empty_table_rows(table) -> None:
    """Remove rows whose logical cells contain no body text (bottom-up)."""
    for row in reversed(list(table.rows)):
        if _row_is_effectively_empty(row):
            _remove_table_row(table, row)


def _table_has_section_label(table, label: str) -> bool:
    return _find_cell_with(table, label) is not None


def _table_is_orphan_continuation(table) -> bool:
    """True for template T2-style tables: one row, no known section label."""
    if len(table.rows) != 1:
        return False
    all_labels = _TIER2_SECTION_LABELS + ("授课题目", "授课时间", "授课学时数")
    for key in all_labels:
        if _table_has_section_label(table, key):
            return False
    return True


def _clear_orphan_continuation_table(table) -> None:
    """Strip template continuation text from orphan single-row tables."""
    for row in table.rows:
        seen: Set[int] = set()
        for cell in row.cells:
            eid = id(cell._element)
            if eid in seen:
                continue
            seen.add(eid)
            _clear_cell_paragraphs(cell)


def _compact_lesson_tables(cloned: List[Any], doc, *, process_design_written: bool) -> List[Any]:
    """Remove orphan continuation tables, empty rows, and blank paragraphs."""
    from docx.table import Table  # type: ignore
    from docx.oxml.ns import qn

    remove_tbls: Set[Any] = set()
    for el in cloned:
        if el.tag != qn("w:tbl"):
            continue
        t = Table(el, doc)
        if _table_is_orphan_continuation(t):
            if process_design_written:
                remove_tbls.add(el)
                continue
            _clear_orphan_continuation_table(t)
        _prune_empty_table_rows(t)
        if len(t.rows) == 0:
            remove_tbls.add(el)

    compacted: List[Any] = []
    for el in cloned:
        if el.tag == qn("w:tbl") and el in remove_tbls:
            continue
        if el.tag == qn("w:p"):
            text = "".join(t.text or "" for t in el.iter(qn("w:t")))
            if not text.strip():
                continue
        compacted.append(el)
    return compacted


def _find_cell_with(table, key_substr: str):
    """Return the first cell whose FIRST-PARAGRAPH label-prefix contains
    `key_substr` (after normalization).

    "Label-prefix" = the part of the first paragraph BEFORE the first `（` /
    `(` / `：` / `:` — this matches how the 珠科 template formats every section
    as `LABEL[（DESCRIPTION）]：value`. Critically, it avoids two false-positive
    classes that broke the old substring-only logic:

      1) Long elaborate descriptions like "主要教学内容（课程教学内容需与课程
         教学目标相匹配...）：xxx" would otherwise match `教学目标`.
      2) Footnote-style cells like "复习思考及作业布置（目的是帮助学生达成
         教学目标）：..." would also match `教学目标`.

    We intentionally do NOT dedup by `id(cell._element)` because lxml proxy
    objects are short-lived and ids get recycled across rows, which silently
    drops the second half of a table from the search.
    """
    nkey = _norm(key_substr)
    if not nkey:
        return None
    for row in table.rows:
        for cell in row.cells:
            paras = cell.paragraphs
            if not paras:
                continue
            # Find the FIRST non-empty paragraph (template cells sometimes
            # start with a blank line, e.g. T1R1 holds an empty first <w:p>
            # then "教学重点：...").
            first_text = ""
            for p in paras:
                if (p.text or "").strip():
                    first_text = p.text
                    break
            if not first_text:
                continue
            # Trim from the first label-terminator we hit.
            cut = len(first_text)
            for sep in ("（", "(", "：", ":"):
                idx = first_text.find(sep)
                if idx >= 0 and idx < cut:
                    cut = idx
            prefix = _norm(first_text[:cut])
            if prefix and nkey in prefix:
                return cell
    return None


def _append_after(ref_element, new_element) -> None:
    ref_element.addnext(new_element)


def _normalize_tbl_width(table_el) -> None:
    """Force table to span full text width and remove left indent."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tblPr = table_el.find(qn("w:tblPr"))
    if tblPr is None:
        return
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:w"), "5000")
    tblW.set(qn("w:type"), "pct")
    tblInd = tblPr.find(qn("w:tblInd"))
    if tblInd is None:
        tblInd = OxmlElement("w:tblInd")
        tblPr.append(tblInd)
    tblInd.set(qn("w:w"), "0")
    tblInd.set(qn("w:type"), "dxa")


def build_docx(
    *,
    cover: Dict[str, str],
    lesson_contents: List[Dict[str, Any]],
    semester_label: str,
    template_bytes: Optional[bytes] = None,
) -> bytes:
    """Build the final docx.

    Args:
        cover: dict with keys college / major / class_name / course_type / course_name / teacher
        lesson_contents: list of per-lesson dicts; each item:
            { 'title': str, 'time_label': str, 'hours': str, 'sections': {学情分析: str, 教学目标: str, ...} }
        semester_label: e.g. "2025～2026 学年第 2 学期"
        template_bytes: optional in-memory override of the template file.
    """
    import docx as _docx
    from docx.oxml.ns import qn

    if template_bytes is None:
        if not template_exists():
            raise FileNotFoundError(f"模板缺失：{_TEMPLATE_PATH}")
        with open(_TEMPLATE_PATH, "rb") as f:
            template_bytes = f.read()

    doc = _docx.Document(io.BytesIO(template_bytes))

    # 1) Cover: replace the 6 cover label paragraphs.
    for para in doc.paragraphs:
        for label_seed, key in _COVER_PARA_LABELS:
            if _para_starts_with_cover_label(para.text, label_seed):
                _replace_cover_runs(para, label_seed, cover.get(key, ""))
                break

    # 2) Semester line: replace "2025～2026 学年第 2 学期" with the user-provided label
    #    (single occurrence near the top of the cover).
    if semester_label:
        for para in doc.paragraphs:
            n = re.sub(r"\s+", "", para.text)
            if "学年第" in n and "学期" in n:
                if para.runs:
                    para.runs[0].text = semester_label
                    for r in para.runs[1:]:
                        r.text = ""
                break

    # 3) Per-lesson blocks.
    #    The template currently contains heading + tables for sample chapters 1 / 2 / 3.
    #    We will:
    #      a) Find the FIRST occurrence of "珠海科技学院教学设计（教案）" heading.
    #      b) Capture that heading + the next 4 tables as the "lesson block template".
    #      c) Remove ALL subsequent body content from that point.
    #      d) For each user lesson, deepcopy the block, populate cells, append.

    body = doc.element.body
    children = list(body.iterchildren())
    heading_idx = -1
    for i, el in enumerate(children):
        if el.tag == qn("w:p"):
            text = "".join(t.text or "" for t in el.iter(qn("w:t")))
            if "珠海科技学院教学设计" in text:
                heading_idx = i
                break

    if heading_idx < 0:
        raise RuntimeError("模板异常：找不到『珠海科技学院教学设计（教案）』标题")

    # Collect the heading + next 4 tables (skip empty paragraphs in between).
    block_template: List[Any] = [children[heading_idx]]
    tables_collected = 0
    cursor = heading_idx + 1
    while cursor < len(children) and tables_collected < 4:
        el = children[cursor]
        block_template.append(el)
        if el.tag == qn("w:tbl"):
            tables_collected += 1
        cursor += 1

    # Remove everything from heading_idx onward (we'll re-append per lesson below).
    for el in children[heading_idx:]:
        body.remove(el)

    # The user explicitly requested that the 授课题目（项目或模块） cell
    # carries the FULL outline from the uploaded 教学日历 (lesson.content /
    # lesson.topic). The previous behaviour truncated via _short_title which
    # often dropped meaningful detail (e.g. lost the description after the
    # first 30 chars or first Chinese punctuation). Now we prefer the raw
    # outline verbatim; lesson.title only acts as a UI display fallback.
    def _resolve_topic(lesson: Dict[str, Any]) -> str:
        for k in ("topic", "content", "outline"):
            v = str(lesson.get(k) or "").strip()
            if v:
                return v
        v = str(lesson.get("title") or "").strip()
        if v:
            return v
        idx_label = str(lesson.get("idx_label") or "").strip()
        return idx_label or "本节课"

    # Helpers to populate one lesson block (list of XML elements).
    # `include_heading` controls whether the cloned chunk has the
    # "珠海科技学院教学设计（教案）" header paragraph in front of its 4 tables.
    # We pass False for the 2nd+ lesson within the SAME week so multiple
    # lessons share one heading per page; True for the first lesson of a new
    # week (which also gets a page break inserted before it).
    def _strip_cant_split(table_el) -> None:
        """Remove <w:cantSplit/> from every row's <w:trPr>.

        The template's cell rows are tagged ``cantSplit`` (do not split row
        across pages). With AI-generated paragraphs that often exceed one
        page, this forces the entire row onto the next page and produces
        a half-empty page above. Stripping it lets Word/LibreOffice flow
        content naturally.
        """
        for tr in table_el.iter(qn("w:tr")):
            for trPr in tr.iter(qn("w:trPr")):
                for cs in list(trPr.iter(qn("w:cantSplit"))):
                    cs.getparent().remove(cs)

    def _populate_and_clone(lesson: Dict[str, Any], *, include_heading: bool):
        cloned = [deepcopy(el) for el in block_template]
        # block_template = [heading_para, ...empty paras..., T0, ..., T3].
        # When include_heading is False we strip the leading paragraph(s)
        # before the FIRST table so the new lesson's tables hug the previous
        # lesson's tables without a stray heading line.
        if not include_heading:
            # Drop leading elements up to and INCLUDING the first paragraph
            # that contains the heading marker.
            keep_from = 0
            for i, el in enumerate(cloned):
                if el.tag == qn("w:p"):
                    text = "".join(t.text or "" for t in el.iter(qn("w:t")))
                    if "珠海科技学院教学设计" in text:
                        keep_from = i + 1
                        break
            cloned = cloned[keep_from:]
            while cloned and cloned[0].tag == qn("w:p"):
                para_text = "".join(t.text or "" for t in cloned[0].iter(qn("w:t")))
                if para_text.strip():
                    break
                cloned = cloned[1:]

        # Wrap cloned tables into python-docx Table objects so we can use cells API.
        # We do this by temporarily attaching the elements to a fresh doc.
        from docx.table import Table  # type: ignore

        title = _resolve_topic(lesson)
        time_label = str(lesson.get("time_label") or "").strip()
        hours = str(lesson.get("hours") or "").strip()
        sections = lesson.get("sections") or {}

        # Tier 1 — label/value 分两 cell：在 label cell 同行找下一 logical cell 写。
        tier1 = [
            ("授课题目", title),
            ("授课时间", time_label),
            ("授课学时数", hours),
        ]
        # Tier 2 — label+value 同居一 cell：保留模板 label paragraph，
        # 只替换「：」后的示例 value。空值时整段保留模板原样，方便手工补。
        tier2 = [
            ("学情分析", _fmt_section_value(sections.get("学情分析"), "学情分析")),
            ("教学目标", _fmt_section_value(sections.get("教学目标"), "教学目标")),
            ("主要教学内容", _fmt_section_value(sections.get("主要教学内容"), "主要教学内容")),
            ("教学方法", _fmt_section_value(sections.get("教学方法"), "教学方法")),
            ("教学媒体", _fmt_section_value(sections.get("教学媒体"), "教学媒体")),
            ("教学过程设计", _fmt_section_value(sections.get("教学过程设计"), "教学过程设计")),
            ("复习思考及作业布置", _fmt_section_value(sections.get("作业布置"), "作业布置")),
            ("参考资料", _fmt_section_value(sections.get("参考资料"), "参考资料")),
            ("评估与反馈", _fmt_section_value(sections.get("评估与反馈"), "评估与反馈")),
        ]
        process_design_written = tier2[5][1] is not None

        for el in cloned:
            if el.tag != qn("w:tbl"):
                continue
            _strip_cant_split(el)
            _normalize_tbl_width(el)
            t = Table(el, doc)

            for key, value in tier1:
                label_cell = _find_cell_with(t, key)
                if label_cell is None:
                    continue
                if not value:
                    value_cell = _next_logical_cell_in_row(t, label_cell)
                    if value_cell is not None and id(value_cell._element) != id(label_cell._element):
                        _clear_cell_paragraphs(value_cell)
                    elif key == "授课题目":
                        _remove_row_for_cell(t, label_cell)
                    continue
                value_cell = _next_logical_cell_in_row(t, label_cell)
                if value_cell is None or id(value_cell._element) == id(label_cell._element):
                    _set_cell_text(label_cell, value)
                else:
                    _set_cell_text(value_cell, value)

            for key, value in tier2:
                cell = _find_cell_with(t, key)
                if cell is None:
                    continue
                if value is None:
                    _remove_row_for_cell(t, cell)
                    continue
                _set_value_after_label_colon(
                    cell, value, reference=(key == "参考资料")
                )

            if sections.get("主要教学内容"):
                for leftover_key in ("教学重点", "教学难点"):
                    leftover = _find_cell_with(t, leftover_key)
                    if leftover is None:
                        continue
                    _remove_row_for_cell(t, leftover)

        return _compact_lesson_tables(
            cloned,
            doc,
            process_design_written=process_design_written,
        )

    # Per-week pagination: user requirement is that each new week opens a
    # fresh "珠海科技学院教学设计（教案）" page, and all lessons within the
    # SAME week stack continuously (one heading then N consecutive lesson
    # table blocks). We detect a new week when `lesson.week` (string or int)
    # changes between consecutive lessons. The very first lesson never needs
    # a page break (the doc cover already gives it a fresh page).
    def _week_key(lesson: Dict[str, Any]) -> str:
        # Normalize so "1" / "01" / 1 / " 1 " are all equal.
        w = str(lesson.get("week") or "").strip().lstrip("0")
        return w or "_"

    last_week_key = None
    for i, lesson in enumerate(lesson_contents):
        wk = _week_key(lesson)
        is_new_week = wk != last_week_key
        if is_new_week and i > 0:
            # Page break BEFORE this lesson's heading. We use a freshly
            # constructed paragraph carrying just a `<w:br w:type="page"/>`
            # run — most reliable cross-Word/LibreOffice page break.
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn as _qn
            pbreak_p = OxmlElement("w:p")
            r = OxmlElement("w:r")
            br = OxmlElement("w:br")
            br.set(_qn("w:type"), "page")
            r.append(br)
            pbreak_p.append(r)
            body.append(pbreak_p)
        for el in _populate_and_clone(lesson, include_heading=is_new_week):
            body.append(el)
        last_week_key = wk

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


# ─────────────────────────────────────────────────────────────────
# Convenience for callers
# ─────────────────────────────────────────────────────────────────

WEEKDAY_CN = ["星期一", "星期二", "星期三", "星期四", "星期五", "星期六", "星期日"]


def compose_time_label(week: str, weekday: str, periods: str, date: str = "") -> str:
    """Compose '第 1 周 星期一  第 3、4 节' style 授课时间 label.

    Priority:
      1) Legacy "第 N 周 星期X 第 P 节" — used whenever ANY of week / weekday /
         periods is filled. This is the format 珠科 expects in 教学设计（教案）.
      2) Date fallback — '2026-03-02（星期一）' if both date & weekday;
         else date alone. Only when all three legacy parts are empty.
      3) "未指定时间" sentinel when nothing useful is available.
    """
    wd = (weekday or "").strip()
    if wd.isdigit() and 1 <= int(wd) <= 7:
        wd = WEEKDAY_CN[int(wd) - 1]
    elif wd and "星期" not in wd:
        wd = f"星期{wd}"

    parts: List[str] = []
    if week:
        w = re.sub(r"[第周\s]", "", str(week)) or str(week)
        parts.append(f"第 {w} 周")
    if wd:
        parts.append(wd)
    if periods:
        p = str(periods).strip()
        if "节" not in p:
            p = f"第 {p} 节"
        parts.append(p)
    label = "  ".join(parts).strip()
    if label:
        return label

    d = (date or "").strip()
    if d and wd:
        return f"{d}（{wd}）"
    return d or "未指定时间"
